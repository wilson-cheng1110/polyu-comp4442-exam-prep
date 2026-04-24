"""
Generate coding_questions.docx — COMP4442 programming Q&A.
Style mirrors COMP3438: code listings with fill-in-blank, spot-the-bug,
line-by-line analysis, write-an-endpoint, trace-the-output.
60 marks worth of questions (matching ~40% of 100-mark paper weight).
Topics: Flask REST API, PySpark RDD, PySpark DataFrame, Flask+DynamoDB/MySQL,
        Docker Dockerfile, IoT/sensor pipeline.
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

for section in doc.sections:
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)

style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(10.5)


def H(text, lvl=1):
    p = doc.add_heading(text, level=lvl)
    return p


def Q(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(11)
    return p


def A(text):
    p = doc.add_paragraph(text)
    for r in p.runs:
        r.font.size = Pt(10.5)
    return p


def Code(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9.5)
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_after = Pt(3)
    return p


def Bullet(text, style='List Bullet'):
    p = doc.add_paragraph(text, style=style)
    return p


# ============================================================
# TITLE
# ============================================================
title = doc.add_heading('COMP4442 — Programming Questions & Model Answers', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub = doc.add_paragraph(
    'Style: COMP3438-inspired (line-by-line analysis, fill-in-blank, spot bug, write endpoint).\n'
    'Stack: Flask (Python) + PySpark. Suitable for 40% coding portion of final exam.')
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.runs[0].italic = True
doc.add_paragraph()

# ============================================================
# QUESTION 1 — Flask REST CRUD (20 marks)
# ============================================================
H('Question 1 (20 marks) -- Flask REST API: Student Records', lvl=1)
A('A developer wrote the following Flask application to manage student records via a REST API. '
  'Read the code carefully and answer the questions below.')

Code(
    'Line 01: from flask import Flask, request, jsonify\n'
    'Line 02: \n'
    'Line 03: app = Flask(__name__)\n'
    'Line 04: \n'
    'Line 05: students = {}\n'
    'Line 06: \n'
    'Line 07: @app.route(\'/students\', methods=[\'GET\'])\n'
    'Line 08: def list_students():\n'
    'Line 09:     return jsonify(list(students.values())), 200\n'
    'Line 10: \n'
    'Line 11: @app.route(\'/students/<sid>\', methods=[\'GET\'])\n'
    'Line 12: def get_student(sid):\n'
    'Line 13:     if sid not in students:\n'
    'Line 14:         return jsonify({\'error\': \'not found\'}), ____\n'
    'Line 15:     return jsonify(students[sid]), 200\n'
    'Line 16: \n'
    'Line 17: @app.route(\'/students\', methods=[\'POST\'])\n'
    'Line 18: def create_student():\n'
    'Line 19:     data = request.____\n'
    'Line 20:     if not data or \'id\' not in data:\n'
    'Line 21:         return jsonify({\'error\': \'id required\'}), 400\n'
    'Line 22:     sid = data[\'id\']\n'
    'Line 23:     if sid in students:\n'
    'Line 24:         return jsonify({\'error\': \'exists\'}), ____\n'
    'Line 25:     students[sid] = {\'name\': data[\'name\'], \'dept\': data.get(\'dept\', \'\')}\n'
    'Line 26:     return jsonify(students[sid]), ____\n'
    'Line 27: \n'
    'Line 28: @app.route(\'/students/<sid>\', methods=[\'DELETE\'])\n'
    'Line 29: def delete_student(sid):\n'
    'Line 30:     if sid not in students:\n'
    'Line 31:         return jsonify({\'error\': \'not found\'}), 404\n'
    'Line 32:     del students[sid]\n'
    'Line 33:     return ____\n'
    'Line 34: \n'
    'Line 35: if __name__ == \'__main__\':\n'
    'Line 36:     app.run(host=\'0.0.0.0\', port=5000, debug=True)'
)

Q('Q1(a) Fill in the four blanks on lines 14, 19, 24, 26, and 33. Choose from the options below. (5 marks)\n'
  'Options: A. get_json()  B. form  C. json  D. args  E. 201  F. 200  G. 204  H. 404  I. 409  '
  'J. \'\', 204  K. jsonify({}), 204  L. 400')

A('Answers:')
Bullet('Line 14: H — 404 (student not found).')
Bullet('Line 19: A — get_json() (reads JSON body from POST request; equivalent to request.json).')
Bullet('Line 24: I — 409 (conflict — resource already exists).')
Bullet('Line 26: E — 201 (created — resource successfully created).')
Bullet('Line 33: J — \'\', 204 (no content — successful DELETE returns empty body with status 204).')

Q('Q1(b) Explain what HTTP verb and status code convention each of the four CRUD endpoints follows. '
  'Why is the status code for a successful creation different from a successful read? (5 marks)')
A('REST HTTP verb conventions:')
Bullet('GET /students — returns 200 OK + list. GET is idempotent and safe (no side effects).')
Bullet('GET /students/<sid> — returns 200 OK + record, or 404 if absent.')
Bullet('POST /students — creates a new record; returns 201 CREATED. POST is not idempotent.')
Bullet('DELETE /students/<sid> — deletes; returns 204 NO CONTENT (no body). DELETE is idempotent.')
A('Why 201 vs 200: 200 means "the request succeeded and here is a representation"; '
  '201 specifically means "a new resource was created". Using 201 lets the client distinguish '
  '"I got data" from "I just created something", and the Location header could optionally '
  'carry the URL of the new resource.')

Q('Q1(c) Add a PUT /students/<sid> endpoint that replaces a student\'s record. '
  'If the student does not exist return 404. Return 200 with the updated record. (5 marks)')
A('Model answer:')
Code(
    '@app.route(\'/students/<sid>\', methods=[\'PUT\'])\n'
    'def update_student(sid):\n'
    '    if sid not in students:\n'
    '        return jsonify({\'error\': \'not found\'}), 404\n'
    '    data = request.get_json()\n'
    '    if not data:\n'
    '        return jsonify({\'error\': \'JSON body required\'}), 400\n'
    '    students[sid] = {\n'
    '        \'name\': data.get(\'name\', students[sid][\'name\']),\n'
    '        \'dept\': data.get(\'dept\', students[sid].get(\'dept\', \'\'))\n'
    '    }\n'
    '    return jsonify(students[sid]), 200'
)

Q('Q1(d) The code on line 05 uses a Python dictionary as an in-memory database. '
  'Describe TWO problems this causes in a production cloud deployment and propose solutions. (5 marks)')
A('Problems:')
Bullet('Statelessness violation — the dictionary lives in the process memory of one Lambda/EC2 instance. '
       'If multiple instances are running (e.g. behind an ELB), each has its own copy: one instance may '
       'have a record another does not, leading to inconsistent responses. Horizontal scaling breaks.')
Bullet('Persistence loss — data is lost on every restart, crash, or deployment. A REST API needs durable '
       'storage.')
A('Solutions:')
Bullet('Replace the dict with a DynamoDB table (NoSQL, PAY_PER_REQUEST). Lambda (stateless) queries '
       'DynamoDB on every request — all instances share the same table, and data persists across restarts.')
Bullet('Alternatively, use Amazon RDS (MySQL/PostgreSQL) behind a connection pool (RDS Proxy for Lambda) '
       'or Amazon ElastiCache (Redis) for a fast in-memory store that is shared and persistent.')

doc.add_page_break()

# ============================================================
# QUESTION 2 — PySpark RDD (20 marks)
# ============================================================
H('Question 2 (20 marks) — PySpark RDD: Word Frequency Filter', lvl=1)
A('A data engineer wrote the following PySpark RDD job to find all words appearing more than '
  '5 times in a log file. Read the code and answer the questions below.')

Code(
    'Line 01: from pyspark import SparkContext\n'
    'Line 02: \n'
    'Line 03: sc = SparkContext("local[*]", "WordFreq")\n'
    'Line 04: \n'
    'Line 05: lines = sc.textFile("hdfs:///logs/server.log")\n'
    'Line 06: words = lines.______(lambda line: line.lower().split())\n'
    'Line 07: pairs = words.map(lambda w: (w, 1))\n'
    'Line 08: counts = pairs.reduceByKey(lambda a, b: a + b)\n'
    'Line 09: frequent = counts.filter(lambda kv: kv[1] > 5)\n'
    'Line 10: result = frequent.______(lambda kv: -kv[1])\n'
    'Line 11: \n'
    'Line 12: for word, count in result.take(20):\n'
    'Line 13:     print(f"{word}: {count}")\n'
    'Line 14: \n'
    'Line 15: result.saveAsTextFile("hdfs:///output/frequent_words")\n'
    'Line 16: sc.stop()'
)

Q('Q2(a) Fill in the two blanks on lines 06 and 10. Choose from the options below. (4 marks)\n'
  'Options: A. map  B. flatMap  C. filter  D. reduce  E. sortBy  F. sortByKey  G. groupByKey  H. collect')

A('Answers:')
Bullet('Line 06: B — flatMap. Unlike map (which would produce an RDD of lists), flatMap flattens '
       'the output so each word becomes a separate element. Each line is split into multiple words; '
       'flatMap "explodes" the result into one element per word.')
Bullet('Line 10: E — sortBy. Sorts the RDD by the result of the given key function; '
       'lambda kv: -kv[1] sorts by descending count (negative → ascending sort → highest count first). '
       'sortByKey would only sort on the natural key (the word string alphabetically).')

Q('Q2(b) For each of lines 05–15, state whether it is a TRANSFORMATION or an ACTION '
  'and give a one-line reason. (5 marks)')
A('Line-by-line classification:')
Bullet('Line 05 — TRANSFORMATION. textFile returns an RDD of lines (lazy; file not read yet).')
Bullet('Line 06 — TRANSFORMATION. flatMap returns a new RDD; no data movement yet.')
Bullet('Line 07 — TRANSFORMATION. map returns a new pair-RDD (w, 1) per word.')
Bullet('Line 08 — TRANSFORMATION. reduceByKey records the aggregation intent; triggers a wide '
       'shuffle but is still lazy.')
Bullet('Line 09 — TRANSFORMATION. filter creates a new RDD with the predicate; lazy.')
Bullet('Line 10 — TRANSFORMATION. sortBy records a sort; lazy (wide — requires a range shuffle).')
Bullet('Line 12 — ACTION. take(20) triggers execution of the entire DAG (lines 05–10) and returns '
       'the first 20 elements to the driver.')
Bullet('Line 15 — ACTION. saveAsTextFile writes the result to HDFS; triggers a separate job that '
       're-executes the lineage from line 05.')

A('Note: lines 12 and 15 each trigger separate Spark jobs because the RDD is not cached. '
  'For efficiency one should call result.cache() before line 12.')

Q('Q2(c) What is the difference between map and flatMap? When must you use flatMap? '
  'Give an example from the code. (5 marks)')
A('map(f) transforms each element of RDD[T] into exactly ONE element of RDD[U] — one-to-one. '
  'flatMap(f) transforms each element into ZERO or MORE elements — one-to-many. '
  'flatMap flattens the resulting iterables, merging them into a single RDD.')
A('You must use flatMap when each input element produces a variable-length collection of output '
  'elements (e.g., splitting a line of text into words, exploding a JSON array, expanding a range).')
A('In the code: line 06 uses flatMap because each log line (a string) is split into multiple words '
  '(a list). If map were used instead, the result would be RDD[List[str]] — each element would '
  'be a list of words rather than an individual word, making the subsequent map(lambda w: (w, 1)) '
  'receive a list rather than a string.')

Q('Q2(d) The job calls saveAsTextFile on line 15 AFTER take(20) on line 12. '
  'What performance problem does this cause and how do you fix it? (6 marks)')
A('Problem — double computation: Because result is not cached (persisted), each action (take and '
  'saveAsTextFile) re-triggers the full DAG from the source file (sc.textFile → flatMap → map → '
  'reduceByKey → filter → sortBy). The full log file is read twice from HDFS, all transformations '
  'are re-executed twice, including the expensive shuffle in reduceByKey. For a large log file this '
  'wastes both time and cluster resources.')
A('Fix — call result.cache() (or result.persist()) before the first action:')
Code(
    '# After line 10, add:\n'
    'result.cache()          # stores RDD partitions in Executor memory\n'
    '\n'
    'for word, count in result.take(20):       # first action: reads from cache after first compute\n'
    '    print(f"{word}: {count}")\n'
    '\n'
    'result.saveAsTextFile("hdfs:///output/frequent_words")  # second action: reads from cache\n'
    '\n'
    'result.unpersist()      # release memory when done'
)
A('After caching, the first action (take) computes the full DAG once and stores the result in '
  'Executor memory. The second action (saveAsTextFile) reads from that cache, skipping all '
  'transformations. This reduces HDFS reads from 2 to 1 and eliminates the duplicate shuffle.')

doc.add_page_break()

# ============================================================
# QUESTION 3 — PySpark DataFrame (20 marks)
# ============================================================
H('Question 3 (20 marks) — PySpark DataFrame: IoT Sensor Analytics', lvl=1)
A('A cloud engineer is building an IoT analytics pipeline on AWS EMR. '
  'The sensor data is stored in S3 as a JSON file with the following schema:\n\n'
  '    sensor_id: string, room: string, timestamp: long (unix epoch),\n'
  '    temperature: double, humidity: double\n\n'
  'The engineer wrote the following PySpark job. Read it and answer the questions.')

Code(
    'Line 01: from pyspark.sql import SparkSession\n'
    'Line 02: from pyspark.sql.functions import col, avg, max, min, count, from_unixtime\n'
    'Line 03: \n'
    'Line 04: spark = SparkSession.builder \\\n'
    'Line 05:     .appName("IoT Analytics") \\\n'
    'Line 06:     .getOrCreate()\n'
    'Line 07: \n'
    'Line 08: df = spark.read.json("s3://iot-data/sensors/*.json")\n'
    'Line 09: \n'
    'Line 10: df = df.withColumn("readable_time",\n'
    'Line 11:                    from_unixtime(col("timestamp")))\n'
    'Line 12: \n'
    'Line 13: hot_rooms = (df\n'
    'Line 14:     .filter(col("temperature") > 28.0)\n'
    'Line 15:     .groupBy("room")\n'
    'Line 16:     .agg(\n'
    'Line 17:         avg("temperature").alias("avg_temp"),\n'
    'Line 18:         max("temperature").alias("max_temp"),\n'
    'Line 19:         count("*").alias("readings")\n'
    'Line 20:     )\n'
    'Line 21:     .orderBy(col("avg_temp").desc())\n'
    'Line 22: )\n'
    'Line 23: \n'
    'Line 24: hot_rooms.show()\n'
    'Line 25: \n'
    'Line 26: hot_rooms.write.mode("overwrite") \\\n'
    'Line 27:     .parquet("s3://iot-results/hot_rooms/")\n'
    'Line 28: \n'
    'Line 29: spark.stop()'
)

Q('Q3(a) Describe what each block of the code does. '
  'Identify which lines are transformations and which are actions. (5 marks)')
A('Line-by-line:')
Bullet('Lines 01-02 — imports. SparkSession is the entry point for the DataFrame API; '
       'functions provide column operations (avg, max, count, from_unixtime).')
Bullet('Lines 04-06 — create / reuse a SparkSession named "IoT Analytics".')
Bullet('Line 08 — TRANSFORMATION: reads all JSON files matching the wildcard into a DataFrame '
       'with the inferred schema. Lazy — no data is read yet.')
Bullet('Lines 10-11 — TRANSFORMATION: withColumn adds a new column "readable_time" by converting '
       'the unix epoch integer to a human-readable datetime string. Narrow transformation.')
Bullet('Lines 13-22 — TRANSFORMATION chain (all lazy, all wide except filter which is narrow):\n'
       '  • filter (line 14) — narrow: keep rows where temperature > 28°C.\n'
       '  • groupBy + agg (lines 15-20) — wide: shuffle by room; compute avg, max, count per group.\n'
       '  • orderBy (line 21) — wide: global sort by avg_temp descending.')
Bullet('Line 24 — ACTION: show() triggers execution of the full lineage and prints ≤20 rows.')
Bullet('Lines 26-27 — ACTION: write.parquet() triggers a second execution and writes results to S3 '
       'as Parquet files partitioned by the number of output tasks.')
Bullet('Line 29 — spark.stop() releases SparkSession and underlying SparkContext.')

Q('Q3(b) The job calls show() on line 24 and write.parquet() on lines 26-27. '
  'Both are actions on hot_rooms. Is the lineage executed once or twice? '
  'What change would you make to execute it only once? (5 marks)')
A('The lineage is executed TWICE — once for each action. hot_rooms is an RDD/DataFrame whose '
  'lineage is: read JSON → withColumn → filter → groupBy → agg → orderBy. Every action triggers '
  'a new Spark job that re-reads the source files from S3 and re-executes all transformations.')
A('Fix: call hot_rooms.cache() (or .persist()) after line 22, before the first action:')
Code(
    'hot_rooms = (...).orderBy(col("avg_temp").desc())\n'
    'hot_rooms.cache()          # materialise in Executor memory after the first action\n'
    '\n'
    'hot_rooms.show()           # first action: compute + store in cache\n'
    '\n'
    'hot_rooms.write.mode("overwrite").parquet("s3://iot-results/hot_rooms/")\n'
    '                           # second action: read from cache, no re-computation\n'
    '\n'
    'hot_rooms.unpersist()      # free memory\n'
    'spark.stop()'
)
A('Now S3 is read once, and the shuffle (groupBy+agg+orderBy) runs once. '
  'For datasets that fit in Executor memory, this halves the job time and EMR cost.')

Q('Q3(c) The output is written as Parquet (line 27). '
  'Explain two advantages of Parquet over CSV for analytical workloads. (4 marks)')
A('Advantage 1 — columnar storage: Parquet stores data column-by-column rather than row-by-row. '
  'For a query that only reads avg_temp and room, Parquet only reads those columns from disk; '
  'irrelevant columns (sensor_id, humidity, timestamp, readings) are skipped entirely. '
  'For wide tables this reduces I/O and speeds up analytical queries dramatically.')
A('Advantage 2 — schema embedding + compression: Parquet embeds the schema in the file footer '
  '(no need for a separate DDL or inferSchema scan) and applies column-level compression '
  '(Snappy or Gzip). Integer columns like readings compress much better column-wise '
  '(run-length encoding, dictionary encoding) than row-wise CSV. '
  'Result: 5-10× smaller files, faster S3 reads, lower EMR cost.')
A('Bonus: Parquet enables Spark predicate push-down (skip entire row-groups where avg_temp < 30) '
  'and Athena/Redshift Spectrum to query S3 data efficiently without loading it into Spark.')

Q('Q3(d) Write a NEW PySpark DataFrame query (using the same df variable) that finds '
  'the top 3 sensors (by sensor_id) with the highest TOTAL readings count, '
  'and returns their sensor_id and reading count. (6 marks)')
A('Model answer:')
Code(
    'from pyspark.sql.functions import col, count, desc\n'
    '\n'
    'top_sensors = (\n'
    '    df\n'
    '    .groupBy("sensor_id")\n'
    '    .agg(count("*").alias("total_readings"))\n'
    '    .orderBy(desc("total_readings"))\n'
    '    .limit(3)\n'
    ')\n'
    '\n'
    'top_sensors.show()'
)
A('Explanation:')
Bullet('groupBy("sensor_id") — wide transformation: shuffle all rows by sensor_id.')
Bullet('agg(count("*").alias("total_readings")) — count all rows per sensor_id.')
Bullet('orderBy(desc("total_readings")) — wide: global sort by count descending.')
Bullet('limit(3) — narrow: take the first 3 rows (Spark may push the limit down to avoid '
       'reading all partitions).')
Bullet('show() — action: triggers execution.')

doc.add_page_break()

# ============================================================
# QUESTION 4 — Flask + DynamoDB Lambda (15 marks)
# ============================================================
H('Question 4 (15 marks) — Flask / Lambda with AWS DynamoDB', lvl=1)
A('The following code is a Flask application that stores student records in AWS DynamoDB. '
  'Read the code and answer the questions.')

Code(
    'Line 01: from flask import Flask, request, jsonify\n'
    'Line 02: import boto3\n'
    'Line 03: \n'
    'Line 04: app = Flask(__name__)\n'
    'Line 05: dynamodb = boto3.resource(\'dynamodb\', region_name=\'us-east-1\')\n'
    'Line 06: table = dynamodb.Table(\'Students\')\n'
    'Line 07: \n'
    'Line 08: @app.route(\'/students/<sid>\', methods=[\'GET\'])\n'
    'Line 09: def get_student(sid):\n'
    'Line 10:     response = table.get_item(Key={\'StudentID\': sid})\n'
    'Line 11:     item = response.get(\'Item\')\n'
    'Line 12:     if not item:\n'
    'Line 13:         return jsonify({\'error\': \'not found\'}), 404\n'
    'Line 14:     return jsonify(item), 200\n'
    'Line 15: \n'
    'Line 16: @app.route(\'/students\', methods=[\'POST\'])\n'
    'Line 17: def create_student():\n'
    'Line 18:     data = request.get_json()\n'
    'Line 19:     table.put_item(Item={\n'
    'Line 20:         \'StudentID\': data[\'id\'],\n'
    'Line 21:         \'Name\':      data[\'name\'],\n'
    'Line 22:         \'Dept\':      data.get(\'dept\', \'Unknown\')\n'
    'Line 23:     })\n'
    'Line 24:     return jsonify({\'message\': \'created\'}), 201\n'
    'Line 25: \n'
    'Line 26: @app.route(\'/students\', methods=[\'GET\'])\n'
    'Line 27: def list_students():\n'
    'Line 28:     response = table.____(____)  # BLANK A and BLANK B\n'
    'Line 29:     return jsonify(response[\'Items\']), 200\n'
    'Line 30: \n'
    'Line 31: if __name__ == \'__main__\':\n'
    'Line 32:     app.run(debug=True)'
)

Q('Q4(a) Fill in BLANK A (the DynamoDB method) and BLANK B (any required arguments) on line 28. '
  'Explain the difference between the two approaches for listing all items. (4 marks)')
A('Fill-in:')
Code('response = table.scan()    # BLANK A = scan, BLANK B = empty (no args required)')
A('Explanation of the two approaches:')
Bullet('scan() — reads EVERY item in the table and optionally applies a filter. Simple but expensive: '
       'for a table with millions of items, scan reads all partitions even if you only want 10 items. '
       'Charge = all consumed RCUs. Avoid in production at scale.')
Bullet('query() — retrieves items using the partition key (and optional sort key condition). '
       'Requires a KeyConditionExpression; it is much faster and cheaper for targeted reads. '
       'In the SkyPulse CLAUDE.md: "Always query the GSI; never Scan." For listing ALL students '
       'a scan is technically correct if the intent is list-all, but for filtering by e.g. Dept '
       'you should use a GSI + query.')

Q('Q4(b) There is a missing validation bug in create_student() (lines 17-24) that will cause a '
  'KeyError exception at runtime. Identify the line, explain the bug, and fix it. (4 marks)')
A('Bug: line 20 — data[\'id\'] raises KeyError if the JSON body does not contain the key "id". '
  'Similarly line 21 — data[\'name\'] raises KeyError if "name" is missing.')
A('Fixed version:')
Code(
    '@app.route(\'/students\', methods=[\'POST\'])\n'
    'def create_student():\n'
    '    data = request.get_json()\n'
    '    if not data or \'id\' not in data or \'name\' not in data:   # validation guard\n'
    '        return jsonify({\'error\': \'id and name are required\'}), 400\n'
    '    table.put_item(Item={\n'
    '        \'StudentID\': data[\'id\'],\n'
    '        \'Name\':      data[\'name\'],\n'
    '        \'Dept\':      data.get(\'dept\', \'Unknown\')\n'
    '    })\n'
    '    return jsonify({\'message\': \'created\'}), 201'
)

Q('Q4(c) The application is deployed as an AWS Lambda function behind API Gateway. '
  'Draw the request flow and explain what happens when a client calls '
  'GET https://api.example.com/prod/students/001. (4 marks)')
A('Request flow:')
Code(
    'Client (curl / mobile app)\n'
    '    | HTTPS GET /prod/students/001\n'
    '    v\n'
    'API Gateway (edge)\n'
    '    | validates JWT (Cognito User Pool authoriser)\n'
    '    | maps HTTP method + path to Lambda integration\n'
    '    | builds event {httpMethod:"GET", pathParameters:{sid:"001"}, ...}\n'
    '    v\n'
    'AWS Lambda ("student-api" function, Python 3.11)\n'
    '    | Flask (via Zappa/Mangum adapter) routes GET /students/001\n'
    '    | calls get_student("001")\n'
    '    | table.get_item(Key={"StudentID":"001"})\n'
    '    v\n'
    'Amazon DynamoDB ("Students" table)\n'
    '    | lookup by partition key StudentID = "001"\n'
    '    | returns Item dict or empty\n'
    '    v\n'
    'Lambda returns {statusCode:200, body: JSON}\n'
    '    v\n'
    'API Gateway maps response → HTTP 200 + JSON body\n'
    '    v\n'
    'Client receives {StudentID:"001", Name:"Alice", Dept:"CS"}'
)
A('Key points:')
Bullet('Lambda is stateless — no persistent in-memory data between requests; DynamoDB is the durable store.')
Bullet('Cognito authoriser verifies the JWT from the Authorization header before Lambda is invoked.')
Bullet('Cold start: if Lambda has not been invoked recently, AWS first initialises a new execution '
       'environment (~100-500 ms overhead for Python). Subsequent calls within the container lifetime '
       'reuse the warm container.')

Q('Q4(d) Without changing any Flask code, describe how you would add rate-limiting so that a '
  'single user can call GET /students at most 100 times per second. (3 marks)')
A('Use API Gateway Usage Plans + API Keys:')
Bullet('Create a Usage Plan in API Gateway with throttle rate = 100 req/s and burst = 200.')
Bullet('Generate an API key and associate it with the usage plan and the API stage.')
Bullet('Client includes header x-api-key: <key> in every request. API Gateway enforces the rate '
       'limit before the request ever reaches Lambda.')
A('Alternative: API Gateway also supports per-route throttling by stage variables. '
  'For per-user limits (Cognito user), use a Lambda authoriser that checks a rate counter in '
  'DynamoDB or ElastiCache Redis with a TTL of 1 second.')

doc.add_page_break()

# ============================================================
# QUESTION 5 — Flask Debugging + tracing (10 marks)
# ============================================================
H('Question 5 (10 marks) — Flask: Spot the Bugs + Trace Output', lvl=1)
A('The following Flask application manages a counter service. '
  'There are EXACTLY THREE bugs in this code. Find them, explain each, and write the fixed line(s). '
  'Then trace what the client receives when calling POST /increment.')

Code(
    'Line 01: from flask import Flask, request, jsonify\n'
    'Line 02: \n'
    'Line 03: app = Flask(__name__)\n'
    'Line 04: counter = 0\n'
    'Line 05: \n'
    'Line 06: @app.route(\'/counter\', methods=[\'GET\'])\n'
    'Line 07: def get_counter():\n'
    'Line 08:     return jsonify(value=counter), 200\n'
    'Line 09: \n'
    'Line 10: @app.route(\'/increment\', methods=[\'GET\'])\n'
    'Line 11: def increment():\n'
    'Line 12:     counter += 1\n'
    'Line 13:     return jsonify(value=counter), 200\n'
    'Line 14: \n'
    'Line 15: @app.route(\'/reset\', methods=[\'POST\'])\n'
    'Line 16: def reset():\n'
    'Line 17:     amount = request.get_json().get(\'amount\', 0)\n'
    'Line 18:     counter = amount\n'
    'Line 19:     return jsonify(value=counter), 200\n'
    'Line 20: \n'
    'Line 21: if __name__ == \'__main__\':\n'
    'Line 22:     app.run(debug=False)'
)

Q('Q5(a) Identify the THREE bugs. For each: name the line, describe the bug category, '
  'and write the corrected code. (6 marks)')
A('Bug 1 — Line 10: Wrong HTTP method (semantic error)')
A('The /increment endpoint MODIFIES state (increments a counter) but is registered with '
  'methods=[\'GET\']. Mutating state should use POST (or PUT). '
  'GET requests are supposed to be safe (no side effects) and idempotent. '
  'Browsers and CDNs may cache GET responses, causing increment to appear to have no effect.')
Code("Corrected line 10:  @app.route('/increment', methods=['POST'])")

A('Bug 2 — Line 12: UnboundLocalError (Python scoping bug)')
A('counter on line 12 is assigned inside the function (counter += 1), which makes Python treat '
  'counter as a local variable for the entire function. But it was never assigned locally before '
  'the += — it only exists in the global scope (line 04). '
  'Python raises: UnboundLocalError: local variable \'counter\' referenced before assignment. '
  'Fix: declare global counter before using it.')
Code(
    'Corrected lines 11-13:\n'
    'def increment():\n'
    '    global counter\n'
    '    counter += 1\n'
    '    return jsonify(value=counter), 200'
)

A('Bug 3 — Line 18: Same UnboundLocalError in reset() (same Python scoping issue)')
A('Same problem as Bug 2: counter = amount on line 18 makes Python treat counter as local '
  'throughout reset(). Without global counter the assignment raises UnboundLocalError.')
Code(
    'Corrected lines 16-19:\n'
    'def reset():\n'
    '    global counter\n'
    '    amount = request.get_json().get(\'amount\', 0)\n'
    '    counter = amount\n'
    '    return jsonify(value=counter), 200'
)

Q('Q5(b) After all three fixes, trace exactly what the client receives when making '
  'POST /increment (assume counter starts at 0). (4 marks)')
A('After the fixes, POST /increment calls increment():')
Bullet('global counter — tells Python to use the module-level counter (currently 0).')
Bullet('counter += 1 — counter becomes 1.')
Bullet('return jsonify(value=counter), 200 — returns HTTP 200 with JSON body.')
A('Client receives:')
Code(
    'HTTP/1.1 200 OK\n'
    'Content-Type: application/json\n'
    '\n'
    '{"value": 1}'
)

doc.add_page_break()

# ============================================================
# QUESTION 6 — Spark DataFrame Write-it-yourself (10 marks)
# ============================================================
H('Question 6 (10 marks) — Write a PySpark Job from Scratch', lvl=1)
A('You are given a CSV file at s3://flights/data/*.csv with the following columns:\n\n'
  '    flight_id (string), carrier (string), origin (string), destination (string),\n'
  '    departure_date (string, format yyyy-MM-dd), delay_minutes (integer)\n\n'
  'Write a complete, runnable PySpark job that:')
Bullet('Reads the CSV from S3 with header and inferred schema.')
Bullet('Filters out rows where delay_minutes is null or < 0.')
Bullet('For each carrier, computes: count of flights, average delay, and maximum delay.')
Bullet('Orders the result by average delay descending.')
Bullet('Saves the result to s3://flights/results/carrier_stats/ as Parquet, overwriting any '
       'existing data.')
Bullet('Stops the SparkSession.')

Q('Write the complete PySpark job (all imports, session creation, transformations, actions). '
  'Comments are NOT required but any non-obvious line should be briefly annotated. (10 marks)')
A('Model answer (full runnable job):')
Code(
    'from pyspark.sql import SparkSession\n'
    'from pyspark.sql.functions import col, avg, max, count, desc\n'
    '\n'
    'spark = SparkSession.builder \\\n'
    '    .appName("CarrierDelayStats") \\\n'
    '    .getOrCreate()\n'
    '\n'
    'df = spark.read.csv(\n'
    '    "s3://flights/data/*.csv",\n'
    '    header=True,\n'
    '    inferSchema=True\n'
    ')\n'
    '\n'
    'clean = df.filter(\n'
    '    col("delay_minutes").isNotNull() &\n'
    '    (col("delay_minutes") >= 0)\n'
    ')\n'
    '\n'
    'stats = (\n'
    '    clean\n'
    '    .groupBy("carrier")\n'
    '    .agg(\n'
    '        count("*").alias("total_flights"),\n'
    '        avg("delay_minutes").alias("avg_delay"),\n'
    '        max("delay_minutes").alias("max_delay")\n'
    '    )\n'
    '    .orderBy(desc("avg_delay"))\n'
    ')\n'
    '\n'
    'stats.write \\\n'
    '    .mode("overwrite") \\\n'
    '    .parquet("s3://flights/results/carrier_stats/")\n'
    '\n'
    'spark.stop()'
)
A('Marking rubric:')
Bullet('[2 marks] Correct SparkSession creation with appName and getOrCreate.')
Bullet('[2 marks] Correct read.csv with header=True and inferSchema=True.')
Bullet('[2 marks] Correct filter — uses isNotNull() AND >= 0 (must handle both null and negative).')
Bullet('[2 marks] Correct groupBy/agg with count, avg, max and correct aliases.')
Bullet('[1 mark] orderBy(desc(...)) for descending sort.')
Bullet('[1 mark] write.mode("overwrite").parquet(...) and spark.stop().')

doc.add_page_break()

# ============================================================
# QUESTION 7 — Bonus: Flask + IoT + MySQL (extra practice)
# ============================================================
H('Question 7 (Bonus — 5 marks) — Flask Sensor Dashboard', lvl=1)
A('Complete the following Flask application that reads sensor readings from a MySQL database '
  'and returns them as JSON. Fill in ALL five blanks (A–E).')

Code(
    'from flask import Flask, jsonify\n'
    'import mysql.connector\n'
    '\n'
    'app = Flask(__name__)\n'
    '\n'
    'DB_CONFIG = {\n'
    '    "host": "sensor-db.xxx.us-east-1.rds.amazonaws.com",\n'
    '    "user": "admin",\n'
    '    "password": "secret",\n'
    '    "database": "iot",\n'
    '    "port": 3306\n'
    '}\n'
    '\n'
    '@app.route("/sensors", methods=["GET"])\n'
    'def list_readings():\n'
    '    con = mysql.connector.connect(____A____)\n'
    '    cur = con.cursor(____B____)\n'
    '    cur.execute("SELECT sensor_id, temperature, ctime FROM Monitor ORDER BY ctime DESC LIMIT 50")\n'
    '    rows = ____C____\n'
    '    cur.close()\n'
    '    con.close()\n'
    '    return ____D____(rows), ____E____\n'
    '\n'
    'if __name__ == "__main__":\n'
    '    app.run(host="0.0.0.0", port=5000)'
)

Q('Fill in blanks A through E. (5 marks)')
A('Answers:')
Bullet('Blank A: **DB_CONFIG (pass the config dict with `**`):\n'
       '`con = mysql.connector.connect(**DB_CONFIG)`\n'
       'Explanation: mysql.connector.connect() accepts keyword arguments for host, user, password, '
       'database, port. Unpacking the dict with ** passes them all.')
Bullet('Blank B: `dictionary=True`\n'
       'Explanation: with dictionary=True, each row is returned as a dict {column: value} instead '
       'of a tuple. This makes it directly serialisable to JSON.')
Bullet('Blank C: `cur.fetchall()`\n'
       'Explanation: fetchall() retrieves all remaining rows from the executed query as a list '
       'of dicts (since dictionary=True).')
Bullet('Blank D: `jsonify`\n'
       'Explanation: Flask\'s jsonify() serialises the Python list-of-dicts to an HTTP response '
       'with Content-Type: application/json.')
Bullet('Blank E: `200`\n'
       'Explanation: HTTP 200 OK is the status code for a successful GET response.')

# ============================================================
# QUESTION 8 — Flask Blueprints (10 marks)
# ============================================================
doc.add_page_break()
H('Question 8 (10 marks) — Flask Blueprints: Flight API', lvl=1)
A('A developer is restructuring a Flask application into blueprints. Examine the files below.')
Code(
    '# --- auth.py ---\n'
    'from flask import Blueprint, request, jsonify, session\n'
    'auth_bp = Blueprint(____A____, __name__)\n'
    '\n'
    '@auth_bp.route(\'/login\', methods=[\'POST\'])\n'
    'def login():\n'
    '    data = request.get_json()\n'
    '    if data.get(\'user\') == \'admin\' and data.get(\'pw\') == \'secret\':\n'
    '        session[\'user\'] = \'admin\'\n'
    '        return jsonify({\'message\': \'ok\'}), 200\n'
    '    return jsonify({\'error\': \'bad credentials\'}), 401\n'
    '\n'
    '@auth_bp.route(\'/logout\', methods=[\'POST\'])\n'
    'def logout():\n'
    '    session.pop(\'user\', None)\n'
    '    return \'\', 204\n'
    '\n'
    '# --- flights.py ---\n'
    'from flask import Blueprint, jsonify\n'
    'flights_bp = Blueprint(\'flights\', __name__)\n'
    'FLIGHTS = {\'AA1\': {\'origin\': \'HKG\', \'dest\': \'LHR\'}, \'CX2\': {\'origin\': \'HKG\', \'dest\': \'JFK\'}}\n'
    '\n'
    '@flights_bp.route(\'/flights\', methods=[\'GET\'])\n'
    'def list_flights(): return jsonify(list(FLIGHTS.values())), 200\n'
    '\n'
    '@flights_bp.route(\'/flights/<fid>\', methods=[\'GET\'])\n'
    'def get_flight(fid):\n'
    '    f = FLIGHTS.get(fid)\n'
    '    return (jsonify(f), 200) if f else (jsonify({\'error\': \'not found\'}), 404)\n'
    '\n'
    '# --- app.py ---\n'
    'from flask import Flask\n'
    'from auth import auth_bp\n'
    'from flights import flights_bp\n'
    '\n'
    'app = Flask(__name__)\n'
    'app.secret_key = \'change-me\'\n'
    'app.____B____(____C____, url_prefix=\'/auth\')\n'
    'app.____B____(flights_bp, url_prefix=\'/api/v1\')\n'
    '\n'
    'if __name__ == \'__main__\':\n'
    '    app.run(debug=True)'
)
Q('Q8(a) Fill in BLANK A, BLANK B, BLANK C. (3 marks)')
A('Answers:')
Bullet('BLANK A: \'auth\' — the first argument to Blueprint() is its name (a string), used for url_for() calls and endpoint namespacing. e.g. url_for(\'auth.login\') resolves to /auth/login.')
Bullet('BLANK B: register_blueprint — the Flask app method that registers a Blueprint.')
Bullet('BLANK C: auth_bp — the Blueprint object imported from auth.py.')

Q('Q8(b) What is a Flask Blueprint and what problem does it solve? (4 marks)')
A('A Blueprint is a way to organize a Flask application into reusable, modular components. Each Blueprint '
  'groups related routes, templates, and static files under a single logical unit. Problems solved:')
Bullet('Avoid circular imports — with a large app.py that imports models and models import app, you get circular import errors. Blueprints defer binding to the app until register_blueprint() is called.')
Bullet('Modularity / separation of concerns — auth routes live in auth.py, flight routes in flights.py. Each team can work on a different module independently.')
Bullet('Reusability — a Blueprint can be packaged and reused across multiple applications (e.g. an auth blueprint used by multiple microservices).')
Bullet('URL prefix isolation — register_blueprint(bp, url_prefix=\'/api/v1\') automatically prepends /api/v1 to all routes in that blueprint, making versioning trivial.')

Q('Q8(c) If both blueprints define a function named `index()`, what happens? How do you resolve it? (3 marks)')
A('Flask raises an AssertionError at register_blueprint() time: "View function mapping is overwriting an existing endpoint function". Flask endpoint names are global; by default the endpoint name equals the function name, so two `index` functions clash.')
A('Resolution: use the name_prefix (the Blueprint name itself) — Flask automatically namespaces endpoints as blueprint_name.function_name. So auth_bp\'s `login` is endpoint \'auth.login\' and flights_bp\'s routes are \'flights.list_flights\' etc. Two blueprints CAN both have a function named `index` as long as their Blueprint names differ — their endpoints will be \'auth.index\' and \'flights.index\' with no conflict. The issue only occurs if you accidentally define the same function name WITHIN THE SAME blueprint.')

doc.add_page_break()

# ============================================================
# QUESTION 9 — Flask Session + Login Flow (10 marks)
# ============================================================
H('Question 9 (10 marks) — Flask Session-Based Authentication', lvl=1)
A('The following Flask application implements a login/logout flow using server-side sessions. '
  'Read the code and answer the questions.')
Code(
    'from flask import Flask, request, jsonify, session\n'
    '\n'
    'app = Flask(__name__)\n'
    'app.____A____ = \'super-secret-key-change-in-prod\'\n'
    '\n'
    'USERS = {\'alice\': {\'pw\': \'pass123\', \'role\': \'admin\'},\n'
    '         \'bob\':   {\'pw\': \'pass456\', \'role\': \'user\'}}\n'
    '\n'
    '@app.route(\'/login\', methods=[\'POST\'])\n'
    'def login():\n'
    '    data = request.get_json()\n'
    '    user = USERS.get(data.get(\'username\', \'\'))\n'
    '    if not user or user[\'pw\'] != data.get(\'password\', \'\'):\n'
    '        return jsonify({\'error\': \'invalid credentials\'}), ____B____\n'
    '    session[\'username\'] = data[\'username\']\n'
    '    session[\'role\'] = user[\'role\']\n'
    '    return jsonify({\'message\': \'logged in\'}), 200\n'
    '\n'
    '@app.route(\'/dashboard\', methods=[\'GET\'])\n'
    'def dashboard():\n'
    '    if \'username\' not in ____C____:\n'
    '        return jsonify({\'error\': \'not logged in\'}), ____D____\n'
    '    return jsonify({\'user\': session[\'username\'], \'role\': session[\'role\']}), 200\n'
    '\n'
    '@app.route(\'/logout\', methods=[\'POST\'])\n'
    'def logout():\n'
    '    session.____E____\n'
    '    return jsonify({\'message\': \'logged out\'}), 200'
)
Q('Q9(a) Fill in blanks A–E. (5 marks)')
A('Answers:')
Bullet('BLANK A: secret_key — Flask uses this to cryptographically sign the session cookie (HMAC-SHA1). Without it, Flask raises RuntimeError: The session is unavailable because no secret key was set.')
Bullet('BLANK B: 401 — 401 Unauthorized means authentication failed (wrong credentials). (403 would mean authenticated but no permission.)')
Bullet('BLANK C: session — the Flask session proxy object. Checking \'username\' not in session verifies the user is logged in.')
Bullet('BLANK D: 403 — 403 Forbidden. The user is not authenticated; returning 403 indicates they cannot access this resource. (Some APIs return 401 here too — either is defensible; 403 is used when we know the identity is absent.)')
Bullet('BLANK E: clear() — session.clear() removes ALL keys from the session, effectively logging the user out. Alternative: session.pop(\'username\', None) to remove individual keys.')

Q('Q9(b) Describe how Flask\'s server-side session mechanism works. What is stored on the client vs. the server? (3 marks)')
A('Flask\'s default session (from flask.sessions.SecureCookieSessionInterface) is actually a CLIENT-SIDE session:')
Bullet('The entire session dict is serialised (JSON), signed with HMAC using app.secret_key, and stored in a browser cookie named "session". Nothing is stored on the server.')
Bullet('On each request, Flask reads the cookie, verifies the HMAC signature, and deserialises the data. If the signature is invalid (tampered), the session is treated as empty.')
Bullet('Implication: the data is visible (base64-encoded, not encrypted) but not modifiable by the client. To make it secret too, use Flask-Session with a server-side backend (Redis, DB) which stores only a session ID in the cookie.')

Q('Q9(c) A security reviewer flags that storing role in the session is risky. Explain the risk and mitigation. (2 marks)')
A('Risk: if the user\'s role changes in the database (e.g. downgraded from admin to user), the old role remains in the session cookie until logout. An attacker who stays logged in retains elevated privileges indefinitely.')
A('Mitigation: do NOT store mutable authorization data in the session. Instead, look up the role from the database on every protected request (or use short-lived JWTs with role embedded, expiry enforced server-side). Alternatively, store only the username in the session and always fetch current permissions from DB at request time.')

doc.add_page_break()

# ============================================================
# QUESTION 10 — Flask Error Handling (8 marks)
# ============================================================
H('Question 10 (8 marks) — Flask Error Handling', lvl=1)
Code(
    'from flask import Flask, jsonify, abort, request\n'
    'app = Flask(__name__)\n'
    '\n'
    '@app.errorhandler(400)\n'
    'def bad_request(e): return jsonify(error=\'bad request\', detail=str(e)), 400\n'
    '\n'
    '@app.errorhandler(404)\n'
    'def not_found(e): return jsonify(error=\'not found\'), 404\n'
    '\n'
    '@app.errorhandler(405)\n'
    'def method_not_allowed(e): return jsonify(error=\'method not allowed\'), 405\n'
    '\n'
    '@app.errorhandler(500)\n'
    'def server_error(e): return jsonify(error=\'internal server error\'), 500\n'
    '\n'
    '@app.route(\'/divide\', methods=[\'POST\'])\n'
    'def divide():\n'
    '    data = request.get_json()\n'
    '    if not data or \'a\' not in data or \'b\' not in data:\n'
    '        abort(400)\n'
    '    try:\n'
    '        result = int(data[\'a\']) / int(data[\'b\'])\n'
    '    except ZeroDivisionError:\n'
    '        return jsonify(error=\'division by zero\'), 422\n'
    '    except ValueError:\n'
    '        return jsonify(error=\'a and b must be integers\'), 422\n'
    '    return jsonify(result=result), 200'
)
Q('Q10(a) When does Flask automatically raise a 404 vs a 405 error? (3 marks)')
A('404 Not Found — Flask raises this automatically when the requested URL path does not match ANY registered route. e.g. GET /nonexistent triggers 404. Also raised explicitly via abort(404).')
A('405 Method Not Allowed — Flask raises this when the URL matches a registered route BUT the HTTP method used is not in the route\'s methods list. e.g. DELETE /divide → 405 because /divide only accepts POST. Flask also automatically adds a 405 handler that sets the Allow header listing the permitted methods.')
A('Key distinction: 404 = path unknown; 405 = path known but verb wrong.')

Q('Q10(b) What does abort(400) do in the divide() route? Trace what the client receives. (2 marks)')
A('abort(400) immediately raises an HTTPException with status 400, short-circuiting the rest of the view function. Flask catches this exception and looks for a registered @errorhandler(400). The registered handler returns jsonify(error=\'bad request\', detail=str(e)), 400.')
A('Client receives: HTTP 400 with JSON body {"error": "bad request", "detail": "400 Bad Request: ..."} and Content-Type: application/json.')

Q('Q10(c) Write an @errorhandler that catches ALL unhandled Python exceptions (not just HTTPExceptions) and returns HTTP 500 with the exception message. (3 marks)')
A('Model answer:')
Code(
    'from flask import jsonify\n'
    'import traceback\n'
    '\n'
    '@app.errorhandler(Exception)\n'
    'def handle_all_exceptions(e):\n'
    '    # In production, log the traceback but do NOT return it to the client\n'
    '    app.logger.error(traceback.format_exc())\n'
    '    return jsonify(error=\'internal server error\', message=str(e)), 500'
)
A('Note: @app.errorhandler(Exception) catches everything including HTTPExceptions. If you want HTTPExceptions to use their own handlers, add: `from werkzeug.exceptions import HTTPException; if isinstance(e, HTTPException): return e` as the first line.')

doc.add_page_break()

# ============================================================
# QUESTION 11 — Flask before_request JWT Validator (10 marks)
# ============================================================
H('Question 11 (10 marks) — Flask JWT Authentication Middleware', lvl=1)
A('The following before_request function validates a JWT Bearer token on every protected route. Fill in the blanks and answer the questions.')
Code(
    'import jwt   # PyJWT library\n'
    'from flask import Flask, request, jsonify, g\n'
    '\n'
    'app = Flask(__name__)\n'
    'JWT_SECRET = \'my-signing-secret\'\n'
    'PUBLIC_ENDPOINTS = {\'login\', \'health\'}\n'
    '\n'
    '@app.____A____\n'
    'def verify_token():\n'
    '    if request.endpoint in PUBLIC_ENDPOINTS:\n'
    '        return   # skip auth for public routes\n'
    '    auth_header = request.headers.get(____B____, \'\')\n'
    '    if not auth_header.startswith(\'Bearer \'):\n'
    '        return jsonify({\'error\': \'missing token\'}), 401\n'
    '    token = auth_header[____C____:]\n'
    '    try:\n'
    '        payload = jwt.decode(token, JWT_SECRET, algorithms=[____D____])\n'
    '        g.user = payload[\'sub\']\n'
    '    except jwt.ExpiredSignatureError:\n'
    '        return jsonify({\'error\': \'token expired\'}), 401\n'
    '    except jwt.InvalidTokenError:\n'
    '        return jsonify({\'error\': \'invalid token\'}), 401'
)
Q('Q11(a) Fill in blanks A–D. (4 marks)')
A('Answers:')
Bullet('BLANK A: before_request — decorator that registers the function to run before every request handler. Flask calls all before_request functions in registration order; if any returns a non-None value, that value is sent as the response and the route handler is skipped.')
Bullet('BLANK B: \'Authorization\' — the standard HTTP header for authentication credentials (RFC 7235). Bearer tokens are sent as: Authorization: Bearer <token>.')
Bullet('BLANK C: 7 — len(\'Bearer \') == 7. Slicing [7:] strips the "Bearer " prefix, leaving just the token string.')
Bullet('BLANK D: \'HS256\' — HMAC with SHA-256, the most common symmetric JWT signing algorithm. Alternatives: HS384, HS512 (symmetric) or RS256 (asymmetric, RSA). Always specify algorithms explicitly to prevent algorithm confusion attacks (CVE-2015-9235).')

Q('Q11(b) What is Flask\'s `g` object and why is it used here to store g.user? (3 marks)')
A('`g` is Flask\'s application context global — a namespace object that exists for the duration of a single request-response cycle and is reset for each new request. It is thread-safe (each worker thread/greenlet gets its own `g`).')
A('Using g.user here: the before_request function decodes the JWT and stores the username in g.user. The actual route handler can then read g.user without re-decoding the token. This avoids passing the user as a parameter and avoids re-running JWT validation inside every handler.')
A('Alternative patterns: use Flask-Login\'s current_user proxy, or use request context locals with a decorator.')

Q('Q11(c) What HTTP header convention does Bearer token authentication follow? Name the RFC and describe the flow. (3 marks)')
A('Bearer token authentication follows RFC 6750 (The OAuth 2.0 Authorization Framework: Bearer Token Usage).')
A('Flow:')
Bullet('Client authenticates (POST /login) and receives a JWT access token (and optionally a refresh token).')
Bullet('On every subsequent protected request, the client sends: Authorization: Bearer <jwt_token> in the HTTP header.')
Bullet('The server (our before_request function) extracts the token, verifies the signature and expiry, and identifies the user from the payload\'s "sub" (subject) claim.')
Bullet('If invalid/expired: server responds 401 Unauthorized with WWW-Authenticate: Bearer realm="api" header (RFC 6750 Section 3).')
A('Bearer = the holder (bearer) of this token is granted access. Unlike Basic Auth (username:password base64), the token is self-contained and time-limited.')

doc.add_page_break()

# ============================================================
# QUESTION 12 — PySpark Window Functions (8 marks)
# ============================================================
H('Question 12 (8 marks) — PySpark Window Functions', lvl=1)
A('A data engineer uses PySpark window functions to rank flights by delay within each carrier.')
Code(
    'from pyspark.sql import SparkSession\n'
    'from pyspark.sql.functions import col, rank, dense_rank, row_number, desc\n'
    'from pyspark.sql.window import Window\n'
    '\n'
    'spark = SparkSession.builder.appName("FlightRank").getOrCreate()\n'
    'df = spark.read.parquet("s3://flights/processed/*.parquet")\n'
    '\n'
    'w = Window.partitionBy("carrier").orderBy(desc("delay_minutes"))\n'
    '\n'
    'df = (df\n'
    '    .withColumn("rank",       rank().over(w))\n'
    '    .withColumn("dense_rank", dense_rank().over(w))\n'
    '    .withColumn("row_num",    row_number().over(w))\n'
    ')\n'
    '\n'
    'top3 = df.filter(col("rank") <= 3)\n'
    'top3.show()\n'
    'spark.stop()'
)
Q('Q12(a) What does Window.partitionBy("carrier").orderBy(desc("delay_minutes")) do? (3 marks)')
A('This creates a Window Specification defining how the window function should partition and order the data:')
Bullet('partitionBy("carrier") — divides all rows into groups (partitions) by carrier value. Each carrier is treated as an independent window — like GROUP BY but the rows are NOT collapsed into one output row; every input row retains its individual identity.')
Bullet('orderBy(desc("delay_minutes")) — within each carrier partition, rows are ordered by delay_minutes descending (highest delay first). This ordering determines rank assignments.')
A('The window spec is then passed to a window function (rank(), dense_rank(), row_number()) via .over(w) to compute a per-partition, order-sensitive value for each row.')

Q('Q12(b) Given this data for carrier "CX": delays [120, 90, 90, 45], what would rank(), dense_rank(), and row_number() return for each row? (3 marks)')
t = doc.add_table(rows=5, cols=4)
t.style = 'Light Grid Accent 1'
hdr = t.rows[0].cells
hdr[0].text = 'delay_minutes'; hdr[1].text = 'rank()'; hdr[2].text = 'dense_rank()'; hdr[3].text = 'row_number()'
data12 = [('120','1','1','1'),('90','2','2','2'),('90','2','2','3'),('45','4','3','4')]
for i,(a,b,c,d) in enumerate(data12, 1):
    t.rows[i].cells[0].text=a; t.rows[i].cells[1].text=b; t.rows[i].cells[2].text=c; t.rows[i].cells[3].text=d
A('Key differences:')
Bullet('rank() — tied rows get the same rank; the NEXT rank SKIPS (gap). 90,90 → both rank 2; next is rank 4.')
Bullet('dense_rank() — tied rows get the same rank; next rank does NOT skip. 90,90 → both rank 2; next is rank 3.')
Bullet('row_number() — NO ties; each row gets a unique sequential number. Tie-breaking is arbitrary (non-deterministic if no secondary sort).')

Q('Q12(c) Is df.withColumn("rank", rank().over(w)) a transformation or an action? Why? (2 marks)')
A('It is a TRANSFORMATION (specifically a wide transformation). withColumn() is always lazy — it records the window function computation in the logical plan but does NOT execute. Window functions require a shuffle (all rows for the same partition key must land on the same executor), making this a wide transformation. Execution is deferred until an action (show(), write(), collect()) triggers the Spark job.')

doc.add_page_break()

# ============================================================
# QUESTION 13 — PySpark UDF (10 marks)
# ============================================================
H('Question 13 (10 marks) — PySpark UDF vs Built-in Functions', lvl=1)
Code(
    'from pyspark.sql import SparkSession\n'
    'from pyspark.sql.functions import udf, col\n'
    'from pyspark.sql.types import StringType\n'
    '\n'
    'spark = SparkSession.builder.appName("TempClassify").getOrCreate()\n'
    'df = spark.read.parquet("s3://iot/sensors/*.parquet")\n'
    '\n'
    'def classify_temp(t):\n'
    '    if t is None: return "unknown"\n'
    '    if t < 15.0:  return "cold"\n'
    '    if t < 28.0:  return "warm"\n'
    '    return "hot"\n'
    '\n'
    'classify_udf = udf(classify_temp, StringType())\n'
    'df = df.withColumn("temp_class", classify_udf(col("temperature")))\n'
    'df.show()\n'
    'spark.stop()'
)
Q('Q13(a) What is a UDF and what is its performance cost compared to built-in Spark functions? (4 marks)')
A('A UDF (User-Defined Function) is a custom Python (or Java/Scala) function registered with Spark to run row-by-row inside transformations. It lets you express arbitrary logic not covered by Spark\'s built-in column functions.')
A('Performance cost of Python UDFs:')
Bullet('Serialization overhead — for each row, data must be serialized from JVM (Spark\'s Tungsten binary format) to Python pickle format, sent to the Python interpreter, and the result serialized back. This crossing of the JVM-Python boundary is the main bottleneck.')
Bullet('No Catalyst optimization — Spark\'s Catalyst optimizer treats a UDF as a black box: it cannot push predicates through it, reorder it, or apply constant folding. Built-in functions (when, col, avg, max…) are known to Catalyst and can be fully optimized.')
Bullet('No Tungsten code generation — built-in functions generate bytecode (whole-stage codegen); UDFs execute interpreted Python, losing the 10-100x speedup from code generation.')
A('Rule: always prefer built-in Spark SQL functions. Use UDFs only when no built-in equivalent exists.')

Q('Q13(b) Rewrite classify_temp using Spark\'s built-in when/otherwise — no UDF. (4 marks)')
A('Model answer:')
Code(
    'from pyspark.sql.functions import when, col\n'
    '\n'
    'df = df.withColumn("temp_class",\n'
    '    when(col("temperature").isNull(), "unknown")\n'
    '    .when(col("temperature") < 15.0, "cold")\n'
    '    .when(col("temperature") < 28.0, "warm")\n'
    '    .otherwise("hot")\n'
    ')'
)
A('This is functionally identical but: (1) stays in the JVM — no Python serialization, (2) Catalyst can optimize the expression tree, (3) Tungsten generates native bytecode. Typically 3-10x faster than the Python UDF version on large datasets.')

Q('Q13(c) What is a Pandas UDF (vectorized UDF) and when should you use it over a regular Python UDF? (2 marks)')
A('A Pandas UDF (also called vectorized UDF, defined with @pandas_udf decorator) passes data to Python as a pandas Series (or DataFrame) rather than one row at a time. Arrow IPC is used for efficient columnar serialization — far less overhead than row-by-row pickle.')
A('Use Pandas UDF when: (1) you need custom Python/NumPy/SciPy logic not in Spark built-ins AND (2) the operation is naturally vectorized (apply a numpy function to a whole column). It is typically 10-100x faster than a row-by-row Python UDF. Still slower than built-in functions, but the practical choice when UDF logic is unavoidable.')

doc.add_page_break()

# ============================================================
# QUESTION 14 — Write PySpark Job: Delay Buckets (10 marks)
# ============================================================
H('Question 14 (10 marks) — Write a PySpark Job: Flight Delay Buckets', lvl=1)
A('Write a complete, runnable PySpark job using the following specification:')
Bullet('Input: Parquet files at s3://flights/processed/*.parquet. Columns: flight_id (string), carrier (string), delay_minutes (integer, may be null or negative).')
Bullet('Create a new column "delay_bucket" with values: "early" (delay < -10), "on_time" (-10 ≤ delay ≤ 15), "minor" (16–60), "significant" (61–180), "severe" (>180), "unknown" (null).')
Bullet('Count flights per carrier per delay_bucket. Order by carrier ASC, total_flights DESC.')
Bullet('Show the first 20 rows, then write as Parquet to s3://flights/results/delay_buckets/ (overwrite). Use cache() to avoid double computation.')
Bullet('Stop the SparkSession.')
Q('Write the complete PySpark job with all imports. (10 marks)')
A('Model answer:')
Code(
    'from pyspark.sql import SparkSession\n'
    'from pyspark.sql.functions import col, when, count, desc, asc\n'
    '\n'
    'spark = SparkSession.builder.appName("DelayBuckets").getOrCreate()\n'
    '\n'
    'df = spark.read.parquet("s3://flights/processed/*.parquet")\n'
    '\n'
    'df = df.withColumn("delay_bucket",\n'
    '    when(col("delay_minutes").isNull(),      "unknown")\n'
    '    .when(col("delay_minutes") < -10,        "early")\n'
    '    .when(col("delay_minutes") <= 15,        "on_time")\n'
    '    .when(col("delay_minutes") <= 60,        "minor")\n'
    '    .when(col("delay_minutes") <= 180,       "significant")\n'
    '    .otherwise(                              "severe")\n'
    ')\n'
    '\n'
    'result = (\n'
    '    df.groupBy("carrier", "delay_bucket")\n'
    '    .agg(count("*").alias("total_flights"))\n'
    '    .orderBy(asc("carrier"), desc("total_flights"))\n'
    ')\n'
    '\n'
    'result.cache()\n'
    'result.show(20)\n'
    'result.write.mode("overwrite").parquet("s3://flights/results/delay_buckets/")\n'
    'result.unpersist()\n'
    'spark.stop()'
)
A('Marking rubric:')
Bullet('[1] SparkSession.builder.appName().getOrCreate()')
Bullet('[1] spark.read.parquet() with correct path')
Bullet('[3] when/otherwise chain: null check first, correct boundary conditions (< -10, <= 15, <= 60, <= 180, otherwise)')
Bullet('[2] groupBy("carrier","delay_bucket").agg(count("*").alias("total_flights"))')
Bullet('[1] orderBy(asc("carrier"), desc("total_flights"))')
Bullet('[1] result.cache() before two actions')
Bullet('[1] show(20) + write.mode("overwrite").parquet() + unpersist() + spark.stop()')

doc.add_page_break()

# ============================================================
# QUESTION 15 — AWS Lambda Handler + DynamoDB (10 marks)
# ============================================================
H('Question 15 (10 marks) — AWS Lambda Handler + DynamoDB', lvl=1)
A('The following is a Python AWS Lambda function that handles GET /flights/{flight_id} from API Gateway.')
Code(
    'import json\n'
    'import boto3\n'
    'from boto3.dynamodb.conditions import Key\n'
    '\n'
    '# Module-level initialization (runs at cold start only)\n'
    'dynamodb = boto3.resource(\'dynamodb\', region_name=\'ap-east-1\')\n'
    'table = dynamodb.Table(\'Flights\')\n'
    '\n'
    'def lambda_handler(event, context):\n'
    '    flight_id = event[\'pathParameters\'][\'flight_id\']\n'
    '    response = table.get_item(Key={\'FlightID\': flight_id})\n'
    '    item = response.get(\'Item\')\n'
    '    if not item:\n'
    '        return {\'statusCode\': 404,\n'
    '                \'body\': json.dumps({\'error\': \'not found\'})}\n'
    '    return {\'statusCode\': 200,\n'
    '            \'body\': json.dumps(item, default=str)}'
)
Q('Q15(a) What are the `event` and `context` parameters in a Lambda handler? (3 marks)')
A('event — a Python dict (parsed from JSON) containing all input data for the invocation. For an API Gateway proxy integration, it includes:')
Bullet('httpMethod, path, pathParameters, queryStringParameters, headers, body (JSON string), requestContext (identity, stage, etc.).')
A('The Lambda reads event["pathParameters"]["flight_id"] to get the URL path parameter {flight_id}.')
A('context — a LambdaContext object provided by the Lambda runtime with metadata about the invocation:')
Bullet('context.function_name, context.function_version, context.invoked_function_arn')
Bullet('context.aws_request_id (unique per invocation, useful for logging/tracing)')
Bullet('context.get_remaining_time_in_millis() — how many ms until the 15-min timeout')
Bullet('context.memory_limit_in_mb')

Q('Q15(b) Why is `dynamodb` initialized at module level (outside lambda_handler)? Explain cold start and warm container. (4 marks)')
A('Lambda execution model:')
Bullet('Cold start — when AWS receives an invocation and there is no warm (idle) execution environment, it: (1) downloads the deployment package, (2) starts a new execution environment (container), (3) imports your Python module (runs module-level code), (4) calls lambda_handler. This takes ~100-500 ms extra.')
Bullet('Warm container — after the first invocation, AWS keeps the execution environment alive for a period. Subsequent invocations reuse the same process — module-level code does NOT re-run. Only lambda_handler() is called again.')
A('By creating the boto3 client at module scope, the DynamoDB connection/session setup (network handshake, auth, session reuse) happens once at cold start. Warm invocations reuse the existing client, saving ~50-200 ms per request and avoiding redundant network overhead.')
A('Best practice: also fetch secrets (API keys) and parse config at module scope for the same reason — cache in module scope, not inside the handler.')

Q('Q15(c) What does `default=str` do in json.dumps(item, default=str)? Why is it needed for DynamoDB items? (3 marks)')
A('json.dumps(obj) raises TypeError if obj contains any non-JSON-serializable Python type. DynamoDB returns Python Decimal objects (e.g. Decimal(\'23.5\')) for numeric attributes, and datetime objects for timestamps — neither is natively JSON-serializable.')
A('default=str means: "if json.dumps encounters a type it cannot serialize, call str() on it and serialize the resulting string instead." So Decimal(\'23.5\') becomes "23.5" (string) in the JSON output.')
A('Alternative: use a custom encoder class (class DecimalEncoder(json.JSONEncoder): def default(self, o): if isinstance(o, Decimal): return float(o); ...) for better type fidelity, or use boto3\'s TypeDeserializer. For exam purposes, default=str is the accepted shortcut.')

doc.add_page_break()

# ============================================================
# QUESTION 16 — boto3 DynamoDB Operations (8 marks)
# ============================================================
H('Question 16 (8 marks) — boto3 DynamoDB CRUD + GSI Query', lvl=1)
A('Fill in the blanks A–H in the following boto3 DynamoDB code. The table "Students" has partition key StudentID (string) and a GSI "dept-index" on Dept (string).')
Code(
    'import boto3\n'
    'from boto3.dynamodb.conditions import ____A____, Attr\n'
    '\n'
    'table = boto3.resource(\'dynamodb\').Table(\'Students\')\n'
    '\n'
    '# 1. Create / overwrite item\n'
    'table.____B____(Item={\'StudentID\': \'S001\', \'Name\': \'Alice\', \'Dept\': \'CS\'})\n'
    '\n'
    '# 2. Read item by primary key\n'
    'resp = table.get_item(____C____={\'StudentID\': \'S001\'})\n'
    'item = resp.get(\'Item\')   # None if not found\n'
    '\n'
    '# 3. Query GSI by Dept\n'
    'resp = table.query(\n'
    '    ____D____=\'dept-index\',\n'
    '    KeyConditionExpression=____A____(\'Dept\').eq(\'CS\')\n'
    ')\n'
    'cs_students = resp[\'Items\']\n'
    '\n'
    '# 4. Partial update — change Name (\'Name\' is a reserved word)\n'
    'table.update_item(\n'
    '    Key={\'StudentID\': \'S001\'},\n'
    '    UpdateExpression=\'SET ____E____ = :name\',\n'
    '    ExpressionAttributeNames={\'____E____\': \'Name\'},\n'
    '    ExpressionAttributeValues={\'____F____\': \'Alicia\'}\n'
    ')\n'
    '\n'
    '# 5. Delete item\n'
    'table.____G____(Key={\'StudentID\': \'S001\'})'
)
Q('Fill in blanks A–G. (8 marks — 1 mark each, G counts 2)')
A('Answers:')
Bullet('BLANK A: Key — from boto3.dynamodb.conditions import Key. Used in KeyConditionExpression for queries on key attributes (partition/sort key and GSI keys). Attr is for non-key attributes in FilterExpression.')
Bullet('BLANK B: put_item — creates or fully replaces an item. If an item with the same primary key exists, it is overwritten entirely.')
Bullet('BLANK C: Key — table.get_item(Key={\'pk\': \'value\'}) specifies the primary key to look up.')
Bullet('BLANK D: IndexName — specifies which GSI to query. Without this, DynamoDB queries the base table.')
Bullet('BLANK E: #n — an ExpressionAttributeNames alias. \'Name\' is a DynamoDB reserved word; using it directly in UpdateExpression raises a ValidationException. #n is an arbitrary placeholder that maps to \'Name\' via ExpressionAttributeNames={\'#n\': \'Name\'}.')
Bullet('BLANK F: :name — ExpressionAttributeValues placeholders must start with a colon. {\'  :name\': \'Alicia\'} maps the placeholder in UpdateExpression (SET #n = :name) to the actual value.')
Bullet('BLANK G: delete_item — removes the item specified by the primary key. No error if item does not exist (idempotent).')

doc.add_page_break()

# ============================================================
# QUESTION 17 — Dockerfile (8 marks)
# ============================================================
H('Question 17 (8 marks) — Write a Dockerfile for a Flask Application', lvl=1)
A('Write a Dockerfile for a Flask application with these requirements:')
Bullet('Base image: Python 3.11 slim variant.')
Bullet('Working directory: /app.')
Bullet('Install Python dependencies from requirements.txt (copy requirements.txt FIRST for Docker layer caching).')
Bullet('Copy all source code into /app.')
Bullet('Expose port 5000.')
Bullet('Run with gunicorn, 4 workers, binding to 0.0.0.0:5000, app object is "app" in module "app".')
Q('Q17(a) Write the complete Dockerfile. (4 marks)')
A('Model answer:')
Code(
    'FROM python:3.11-slim\n'
    '\n'
    'WORKDIR /app\n'
    '\n'
    'COPY requirements.txt .\n'
    'RUN pip install --no-cache-dir -r requirements.txt\n'
    '\n'
    'COPY . .\n'
    '\n'
    'EXPOSE 5000\n'
    '\n'
    'CMD ["gunicorn", "--workers", "4", "--bind", "0.0.0.0:5000", "app:app"]'
)
Q('Q17(b) Why is COPY requirements.txt . placed BEFORE COPY . .? (2 marks)')
A('Docker builds images in layers — each instruction creates a new layer. Layers are cached: if the layer\'s input has not changed, Docker reuses the cached layer without re-running the instruction.')
A('By copying requirements.txt first and running pip install before copying source code: if you change app.py (but not requirements.txt), Docker reuses the cached pip install layer — no reinstallation needed. If COPY . . came first, any code change would invalidate the requirements layer and force a full pip install on every build, wasting minutes. This ordering is a fundamental Docker performance pattern.')

Q('Q17(c) Explain the difference between CMD and ENTRYPOINT. (2 marks)')
A('ENTRYPOINT — sets the fixed executable that always runs when the container starts. Cannot be overridden by docker run arguments (only by --entrypoint flag). Use for: "this container IS a gunicorn server".')
A('CMD — provides default arguments. If ENTRYPOINT is set, CMD provides default arguments to it (can be overridden by docker run args). If ENTRYPOINT is not set, CMD is the full command to run (also overridable).')
A('Together: ENTRYPOINT ["gunicorn"] + CMD ["--workers","4","app:app"] lets you override workers at run time with `docker run myimage --workers 8 app:app`. In practice, for simple apps, CMD alone (as above) is sufficient.')

doc.add_page_break()

# ============================================================
# QUESTION 18 — Spot 4 Bugs (8 marks)
# ============================================================
H('Question 18 (8 marks) — Spot the Bugs in a Flask Application', lvl=1)
A('The following Flask application has EXACTLY FOUR bugs. Identify each bug: state the line(s), the bug category, and write the corrected code.')
Code(
    'from flask import Flask, request, jsonify\n'
    'import boto3\n'
    '\n'
    'app = Flask(__name__)\n'
    '\n'
    '@app.route(\'/upload\', methods=[\'POST\'])\n'
    'def upload():\n'
    '    file = request.files[\'document\']\n'
    '    s3 = boto3.client(\'s3\')\n'
    '    s3.upload_fileobj(file, \'my-bucket\', file.filename)\n'
    '    return {\'url\': f\'s3://my-bucket/{file.filename}\'}   # LINE A\n'
    '\n'
    '@app.route(\'/data\', methods=[\'GET\', \'POST\'])\n'
    'def data():\n'
    '    if request.method == \'GET\':\n'
    '        return jsonify({\'data\': \'ok\'})\n'
    '    query = request.args.get(\'q\')                       # LINE B\n'
    '    return jsonify({\'query\': query}), 200\n'
    '\n'
    '@app.route(\'/items/<iid>\', methods=[\'DELETE\'])\n'
    'def delete_item(iid):\n'
    '    boto3.resource(\'dynamodb\').Table(\'Items\').delete_item(Key={\'id\': iid})\n'
    '    return jsonify({\'deleted\': True}), 200              # LINE C\n'
    '\n'
    '@app.route(\'/admin\')                                    # LINE D\n'
    'def admin():\n'
    '    return jsonify({\'panel\': \'admin-only-data\'})'
)
Q('Identify all 4 bugs, explain each, and write the corrected code. (8 marks — 2 marks per bug)')
A('Bug 1 — LINE A: Returning a plain dict instead of a Response object')
A('Flask < 2.2 does not auto-jsonify a plain dict return. This raises TypeError: The view function did not return a valid response. Even in Flask 2.2+ (which does auto-convert), the status code defaults to 200 but Content-Type may not be set correctly in all environments.')
Code('return jsonify({\'url\': f\'s3://my-bucket/{file.filename}\'}), 200  # FIXED')

A('Bug 2 — LINE B: Using request.args for a POST body parameter')
A('request.args reads from the URL query string (?q=value). For a POST request, the body data is in request.form (form-encoded) or request.get_json() (JSON body). Using args on POST body data will return None, making the query parameter silently missing.')
Code('query = request.get_json().get(\'q\')  # FIXED — reads from JSON body\n'
     '# OR: query = request.form.get(\'q\')  # if form-encoded POST')

A('Bug 3 — LINE C: DELETE endpoint returning 200 with a body instead of 204 No Content')
A('REST convention: a successful DELETE should return 204 No Content (no response body). Returning 200 with a body is not strictly wrong but violates REST semantics and may confuse clients that check status codes.')
Code('return \'\', 204  # FIXED — no content, correct REST convention')

A('Bug 4 — LINE D: /admin endpoint has no authentication')
A('Any unauthenticated user can call GET /admin and receive sensitive admin panel data. There is no JWT check, no session check, no role check.')
Code(
    '@app.route(\'/admin\')\n'
    'def admin():\n'
    '    if session.get(\'role\') != \'admin\':          # FIXED — role check\n'
    '        return jsonify({\'error\': \'forbidden\'}), 403\n'
    '    return jsonify({\'panel\': \'admin-only-data\'})'
)

doc.add_page_break()

# ============================================================
# QUESTION 19 — repartition vs coalesce (5 marks)
# ============================================================
H('Question 19 (5 marks) — PySpark: repartition vs coalesce', lvl=1)
Q('Q19(a) What is the difference between repartition(n) and coalesce(n) in Spark? When would you use each? (5 marks)')
A('repartition(n):')
Bullet('WIDE transformation — performs a full shuffle: all data is redistributed randomly across exactly n new partitions.')
Bullet('Can INCREASE or DECREASE partition count.')
Bullet('Result is evenly distributed (good for downstream joins/aggregations that benefit from balanced partitions).')
Bullet('Cost: full network shuffle — expensive.')
Bullet('Use when: increasing partitions before a join or aggregation on a large dataset; rebalancing skewed data; changing partition count significantly.')

A('coalesce(n):')
Bullet('NARROW transformation — merges existing partitions without a full shuffle (combines adjacent partitions, reduces partition count by moving data locally).')
Bullet('Can only DECREASE partition count (n must be ≤ current partition count).')
Bullet('Result may be unevenly distributed (some merged partitions are larger).')
Bullet('Cost: no shuffle — very cheap.')
Bullet('Use when: reducing partition count before writing output to S3/HDFS (fewer output files), or after a filter that dropped most rows (consolidate small partitions). df.coalesce(1).write.csv(...) writes a single file.')

A('Rule of thumb: coalesce to reduce (cheap, no shuffle); repartition to increase or rebalance (expensive but necessary).')
Code(
    '# Increase partitions before expensive join\n'
    'df_large.repartition(200).join(df_small, "key")\n'
    '\n'
    '# Reduce to 1 file for output\n'
    'result.coalesce(1).write.mode("overwrite").csv("s3://out/")\n'
    '\n'
    '# Rebalance skewed data\n'
    'df.repartition(100, col("carrier"))  # partition by key'
)

doc.add_page_break()

# ============================================================
# QUESTION 20 — Kubernetes Deployment + Service YAML (5 marks)
# ============================================================
H('Question 20 (5 marks) — Write Kubernetes Deployment + Service YAML', lvl=1)
A('Write a Kubernetes Deployment and Service YAML for a Flask application with:')
Bullet('Image: myapp:1.0, port 5000, 3 replicas.')
Bullet('Resource requests: CPU 100m, memory 128Mi. Limits: CPU 500m, memory 256Mi.')
Bullet('Service type LoadBalancer on port 80, routing to container port 5000.')
Q('Write the complete YAML for both resources. (5 marks)')
A('Model answer:')
Code(
    'apiVersion: apps/v1\n'
    'kind: Deployment\n'
    'metadata:\n'
    '  name: flask-app\n'
    'spec:\n'
    '  replicas: 3\n'
    '  selector:\n'
    '    matchLabels:\n'
    '      app: flask-app\n'
    '  template:\n'
    '    metadata:\n'
    '      labels:\n'
    '        app: flask-app\n'
    '    spec:\n'
    '      containers:\n'
    '      - name: flask-app\n'
    '        image: myapp:1.0\n'
    '        ports:\n'
    '        - containerPort: 5000\n'
    '        resources:\n'
    '          requests:\n'
    '            cpu: "100m"\n'
    '            memory: "128Mi"\n'
    '          limits:\n'
    '            cpu: "500m"\n'
    '            memory: "256Mi"\n'
    '---\n'
    'apiVersion: v1\n'
    'kind: Service\n'
    'metadata:\n'
    '  name: flask-app-svc\n'
    'spec:\n'
    '  type: LoadBalancer\n'
    '  selector:\n'
    '    app: flask-app\n'
    '  ports:\n'
    '  - port: 80\n'
    '    targetPort: 5000'
)
A('Key points:')
Bullet('spec.selector.matchLabels must match template.metadata.labels — this is how the Deployment knows which pods it owns.')
Bullet('resources.requests = scheduling guarantee (scheduler only places pod on node with this much available). resources.limits = hard cap (CPU throttled, memory OOM-killed if exceeded).')
Bullet('Service.spec.selector (app: flask-app) must match pod labels. The LoadBalancer Service provisions a cloud load balancer (AWS ELB) with a public IP, routing port 80 → pod port 5000.')

# ============================================================
# QUICK REFERENCE — key Flask patterns
# ============================================================
doc.add_page_break()
H('Quick Reference Card — Flask & PySpark Patterns', lvl=1)
A('Print this page and attach it to your exam notes.')

H('Flask Essentials', lvl=2)
Code(
    'from flask import Flask, request, jsonify, render_template,\n'
    '               redirect, url_for, session, abort, make_response\n'
    'app = Flask(__name__)\n'
    'app.secret_key = "change-me"\n'
    '\n'
    '# Route with variable + methods\n'
    '@app.route("/api/<int:id>", methods=["GET","POST","PUT","DELETE"])\n'
    'def resource(id): ...\n'
    '\n'
    '# Read request data\n'
    'request.args.get("key")          # ?key=val in URL\n'
    'request.form.get("field")        # form POST\n'
    'request.get_json()               # JSON body\n'
    'request.files["upload"]          # file upload\n'
    'request.headers.get("Auth")      # headers\n'
    '\n'
    '# Response helpers\n'
    'return jsonify(data), 200\n'
    'return jsonify(error="msg"), 404\n'
    'return "", 204\n'
    'return redirect(url_for("func"))\n'
    'abort(403)\n'
    '\n'
    '# Error handler\n'
    '@app.errorhandler(404)\n'
    'def not_found(e): return jsonify(error="not found"), 404\n'
    '\n'
    '# Template\n'
    'return render_template("page.html", name=name, items=items)\n'
    '# Jinja: {{ var }}, {% for x in items %}, {% if cond %}, {{ v|upper }}\n'
    '\n'
    '# Session\n'
    'session["user"] = "alice"\n'
    'session.get("user", "anon")\n'
    'session.pop("user", None)'
)

H('PySpark Essentials', lvl=2)
Code(
    '# Session setup\n'
    'from pyspark.sql import SparkSession\n'
    'spark = SparkSession.builder.appName("App").getOrCreate()\n'
    'sc = spark.sparkContext        # for RDD API\n'
    '\n'
    '# RDD\n'
    'rdd = sc.textFile("path")\n'
    '.flatMap(lambda l: l.split())\n'
    '.map(lambda w: (w, 1))\n'
    '.reduceByKey(lambda a,b: a+b)\n'
    '.filter(lambda kv: kv[1] > 5)\n'
    '.sortBy(lambda kv: -kv[1])\n'
    '.take(10)                      # ACTION\n'
    '.collect()                     # ACTION - small data only\n'
    '.saveAsTextFile("out/")        # ACTION\n'
    '\n'
    '# DataFrame\n'
    'from pyspark.sql.functions import col, avg, max, min, sum, count, desc\n'
    'df = spark.read.csv("path", header=True, inferSchema=True)\n'
    'df = spark.read.json("path")\n'
    'df = spark.read.parquet("s3://bucket/path/")\n'
    '\n'
    'df.filter(col("x") > 0)                    # transformation\n'
    'df.withColumn("new", col("a") * 2)         # transformation\n'
    'df.groupBy("col").agg(\n'
    '    avg("v").alias("avg_v"),\n'
    '    count("*").alias("n"))\n'
    '.orderBy(desc("avg_v"))                    # transformation\n'
    '.limit(10)                                 # transformation\n'
    '.show()                                    # ACTION\n'
    '.write.mode("overwrite").parquet("out/")   # ACTION\n'
    '\n'
    '# cache to avoid double-compute on two actions\n'
    'df.cache(); action1; action2; df.unpersist()\n'
    '\n'
    'spark.stop()'
)

H('HTTP Status Codes to Remember', lvl=2)
t = doc.add_table(rows=9, cols=2)
t.style = 'Light Grid Accent 1'
codes = [
    ('Code', 'Meaning'),
    ('200 OK', 'Successful GET / PUT / PATCH'),
    ('201 Created', 'Successful POST — resource created'),
    ('204 No Content', 'Successful DELETE — no body'),
    ('400 Bad Request', 'Invalid input / missing field'),
    ('401 Unauthorized', 'Not authenticated'),
    ('403 Forbidden', 'Authenticated but no permission'),
    ('404 Not Found', 'Resource does not exist'),
    ('409 Conflict', 'Resource already exists'),
]
for i, (code, meaning) in enumerate(codes):
    t.rows[i].cells[0].text = code
    t.rows[i].cells[1].text = meaning

H('boto3 DynamoDB Patterns', lvl=2)
Code(
    'import boto3\n'
    'from boto3.dynamodb.conditions import Key, Attr\n'
    'table = boto3.resource("dynamodb").Table("MyTable")\n'
    '\n'
    '# CREATE / OVERWRITE\n'
    'table.put_item(Item={"pk": "id1", "sk": "v1", "attr": "val"})\n'
    '\n'
    '# READ by primary key\n'
    'resp = table.get_item(Key={"pk": "id1"})\n'
    'item = resp.get("Item")           # None if not found\n'
    '\n'
    '# QUERY by GSI\n'
    'resp = table.query(\n'
    '    IndexName="gsi-name",\n'
    '    KeyConditionExpression=Key("gsi_pk").eq("value")\n'
    ')\n'
    'items = resp["Items"]\n'
    '\n'
    '# PARTIAL UPDATE  (# = reserved word alias, : = value placeholder)\n'
    'table.update_item(\n'
    '    Key={"pk": "id1"},\n'
    '    UpdateExpression="SET #n = :v, score = score + :inc",\n'
    '    ExpressionAttributeNames={"#n": "Name"},\n'
    '    ExpressionAttributeValues={":v": "Alice", ":inc": 1}\n'
    ')\n'
    '\n'
    '# DELETE item\n'
    'table.delete_item(Key={"pk": "id1"})\n'
    '\n'
    '# SCAN (avoid at scale!)\n'
    'resp = table.scan(FilterExpression=Attr("dept").eq("CS"))\n'
    '\n'
    '# BATCH WRITE\n'
    'with table.batch_writer() as batch:\n'
    '    for item in items:\n'
    '        batch.put_item(Item=item)'
)

H('Dockerfile Commands Quick Ref', lvl=2)
Code(
    'FROM python:3.11-slim          # base image\n'
    'WORKDIR /app                   # set working directory (creates if absent)\n'
    'COPY requirements.txt .        # copy file (src dst)\n'
    'RUN pip install --no-cache-dir -r requirements.txt  # build-time command\n'
    'COPY . .                       # copy everything else (after deps for cache)\n'
    'ENV FLASK_ENV=production       # set environment variable (build + runtime)\n'
    'ARG BUILD_VERSION              # build-time variable only (not in final image)\n'
    'EXPOSE 5000                    # declare port (informational, not a firewall rule)\n'
    'VOLUME /data                   # declare mount point\n'
    'USER nobody                    # run as non-root (security best practice)\n'
    'ENTRYPOINT ["gunicorn"]        # fixed executable (cannot override without --entrypoint)\n'
    'CMD ["--workers","4","app:app"]# default args to ENTRYPOINT (can override at run)\n'
    '\n'
    '# CMD alone (no ENTRYPOINT) — full command, overridable:\n'
    'CMD ["python", "app.py"]\n'
    '\n'
    '# Build & run:\n'
    'docker build -t myapp:1.0 .\n'
    'docker run -p 5000:5000 --env-file .env myapp:1.0\n'
    'docker run myapp:1.0 --workers 8 app:app  # overrides CMD args'
)

H('kubectl Cheat Sheet', lvl=2)
Code(
    '# GET resources\n'
    'kubectl get pods [-n namespace] [-o wide] [-o yaml] [-l app=flask]\n'
    'kubectl get nodes/svc/deploy/rs/pv/pvc/cm/secrets/ingress\n'
    'kubectl get all                          # everything in namespace\n'
    '\n'
    '# INSPECT\n'
    'kubectl describe pod <name>              # events, conditions, resource usage\n'
    'kubectl logs <pod> [-c container] [-f]   # -f = follow live\n'
    'kubectl exec -it <pod> -- /bin/sh        # shell into running container\n'
    'kubectl top pods/nodes                   # resource usage (metrics-server needed)\n'
    '\n'
    '# APPLY / DELETE\n'
    'kubectl apply -f deployment.yaml         # create or update\n'
    'kubectl delete -f deployment.yaml        # delete from file\n'
    'kubectl delete pod/deploy/svc <name>     # delete by resource type + name\n'
    '\n'
    '# SCALE / ROLLOUT\n'
    'kubectl scale deploy <name> --replicas=5\n'
    'kubectl rollout status deploy/<name>     # watch rollout progress\n'
    'kubectl rollout history deploy/<name>    # revision history\n'
    'kubectl rollout undo deploy/<name>       # rollback to previous revision\n'
    '\n'
    '# PORT FORWARD (local testing)\n'
    'kubectl port-forward pod/<name> 8080:5000\n'
    '\n'
    '# CONFIG\n'
    'kubectl config get-contexts              # list clusters\n'
    'kubectl config use-context <name>        # switch cluster\n'
    '\n'
    '# EVENTS (debugging)\n'
    'kubectl get events --sort-by=.metadata.creationTimestamp'
)

doc.add_paragraph()

end = doc.add_heading('End of Coding Questions', level=1)
end.alignment = WD_ALIGN_PARAGRAPH.CENTER
p = doc.add_paragraph('For the exam: read code LINE BY LINE before answering — identify transformations vs actions, '
                      'spot missing validations, note HTTP verb/status code mismatches.')
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

out = r"C:\Users\User\Downloads\4442\exam_output\coding_questions.docx"
doc.save(out)
print(f"Saved {out}")
