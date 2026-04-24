"""
Generate exam_notes.docx — COMP4442 Open-Book Exam Cheat Sheet.
11 sections, dense and scannable. Print and bring to exam.
Based on: all lecture slides 01-09, labs 01-06, tutorials 1-5, 3 past papers.
"""
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()
for section in doc.sections:
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)

style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(10)


def H(text, lvl=1):
    p = doc.add_heading(text, level=lvl)
    return p


def A(text):
    p = doc.add_paragraph(text)
    for r in p.runs:
        r.font.size = Pt(10)
    return p


def B(text):
    p = doc.add_paragraph(text, style='List Bullet')
    for r in p.runs:
        r.font.size = Pt(10)
    return p


def Code(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    p.paragraph_format.left_indent = Cm(0.4)
    p.paragraph_format.space_after = Pt(2)
    return p


def table2(headers, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Light Grid Accent 1'
    for i, h in enumerate(headers):
        t.rows[0].cells[i].text = h
    for ri, row in enumerate(rows, 1):
        for ci, val in enumerate(row):
            t.rows[ri].cells[ci].text = val
    return t


# ===================== TITLE =====================
title = doc.add_heading('COMP4442 — Open-Book Exam Notes', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub = doc.add_paragraph('Service and Cloud Computing | 100 marks | 2 hrs | Open Book')
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.runs[0].italic = True
doc.add_paragraph()

# ===========================================================
# SECTION 1: CLOUD FUNDAMENTALS
# ===========================================================
H('1. Cloud Computing Fundamentals', lvl=1)

H('NIST 5 Essential Characteristics', lvl=2)
B('On-demand self-service — provision compute/storage without human interaction with provider')
B('Broad network access — available over network, accessed by heterogeneous clients (mobile, laptop)')
B('Resource pooling — multi-tenant model, physical resources dynamically assigned, location independence')
B('Rapid elasticity — capabilities can be scaled out/in quickly (to user appears unlimited)')
B('Measured service — metered usage (pay-per-use); transparency for provider and consumer')

H('Service Models', lvl=2)
table2(
    ['Model', 'Customer Manages', 'Provider Manages', 'Examples'],
    [
        ('IaaS', 'OS, runtime, middleware, app, data', 'Hardware, network, virtualisation, storage', 'EC2, GCE, Azure VM'),
        ('PaaS', 'App code + data only', 'OS, runtime, middleware, capacity, patching', 'Elastic Beanstalk, Heroku, GAE, Cloud Run'),
        ('SaaS', 'User config only', 'Everything', 'Gmail, Salesforce, Office 365'),
    ]
)

H('Deployment Models', lvl=2)
B('Public cloud — shared infrastructure, owned by provider (AWS/Azure/GCP), internet-accessible')
B('Private cloud — dedicated to one org, on-premises or hosted; full control, higher cost')
B('Hybrid cloud — mix of public+private connected by VPN/Direct Connect; burst to public')
B('Community cloud — shared by several orgs with common concerns (gov, healthcare)')

H('Economics & Design Principles', lvl=2)
B('CAPEX → OPEX: no upfront hardware investment; pay monthly for consumption')
B('Economies of scale: providers buy in bulk → lower per-unit cost passed to customers')
B('Design for failure: assume any component can fail; build redundancy (multi-AZ, health checks)')
B('Decouple components: SQS/SNS between tiers so failures don\'t cascade')
B('Elasticity: Auto Scaling Groups, Lambda concurrency → scale on demand')
B('Think parallel: distribute work across many instances instead of one big server')

H('Cloud Application Types', lvl=2)
table2(
    ['Type', 'Description', 'AWS Services'],
    [
        ('Web Application', 'Client-server via browser; static + dynamic content; CDN + LB + DB', 'S3, CloudFront, ALB, EC2/Lambda, RDS'),
        ('Serverless', 'Event-driven; no server management; scales to zero', 'Lambda + API Gateway + DynamoDB + S3'),
        ('Data Processing Pipeline', 'Flow of data between processing elements (indexing, image processing)', 'Kinesis → Lambda → S3/DynamoDB'),
        ('Batch Processing', 'High-volume repetitive jobs on schedule; pay-per-use (machines idle = no charge)', 'EMR, S3, Redshift, EventBridge'),
        ('Transaction (OLTP)', 'Online transaction processing; immediate response; ACID consistency', 'RDS Aurora, DynamoDB'),
        ('Big Data / Analytics', 'Collect→Store→Process→Analyze→Visualize large datasets', 'Kinesis+S3+EMR+Redshift+QuickSight'),
        ('IoT Application', 'Device telemetry ingestion; rules-based routing; device shadow sync', 'IoT Core + Lambda + DynamoDB + SNS'),
        ('ML Application', 'Real-time predictions; model training on large datasets', 'Kinesis + Lambda + SageMaker/EMR + SNS'),
    ]
)

H('AWS Global Infrastructure', lvl=2)
B('Region — geographic area with 2+ AZs (e.g. ap-east-1 = Hong Kong). Choose for latency + compliance')
B('Availability Zone (AZ) — one or more discrete data centres, physically separate, connected by low-latency links')
B('Edge Location / PoP — CloudFront CDN nodes (400+); closer to users for cached content')
B('Shared Responsibility: AWS = security OF the cloud (hardware, AZs, network). Customer = security IN the cloud (OS, IAM, data, encryption)')

doc.add_page_break()

# ===========================================================
# SECTION 2: VIRTUALIZATION
# ===========================================================
H('2. Virtualization', lvl=1)

H('Types of Virtualization', lvl=2)
B('Hardware/Hypervisor — virtualises CPU/RAM/devices; multiple VMs share one physical machine')
B('OS-Level (Containers) — shares host kernel; cgroups+namespaces isolate processes (Docker, LXC)')
B('Storage virtualization — abstract physical storage into logical volumes (SAN, EBS)')
B('Network virtualization — virtual networks on shared physical fabric (VLAN, VXLAN, SDN)')
B('Desktop virtualization — VDI; remote desktop on centralised servers')

H('Hypervisor: Type-1 vs Type-2', lvl=2)
table2(
    ['Aspect', 'Type-1 (Bare-metal)', 'Type-2 (Hosted)'],
    [
        ('Runs on', 'Directly on hardware (IS the kernel)', 'On top of host OS as an application'),
        ('Examples', 'Xen, VMware ESXi, Hyper-V, KVM', 'VMware Workstation, VirtualBox, Parallels'),
        ('Performance', 'Better — no host OS overhead layer', 'Worse — extra OS layer adds latency'),
        ('Isolation', 'Stronger — smaller TCB, fewer attack surfaces', 'Weaker — host OS shared with VMs'),
        ('Use case', 'Production datacentres, public cloud', 'Dev/test on developer laptops'),
        ('Setup', 'Needs dedicated hardware', 'Install alongside existing OS'),
    ]
)
A('*KVM is a Linux kernel module making Linux itself a Type-1 hypervisor (sometimes called "Type-1.5").')

H('Xen Architecture', lvl=2)
B('Dom0 — privileged control domain; runs real NIC drivers, management tools, xenstore')
B('DomU — unprivileged guest VMs; cannot access hardware directly')
B('PV (Paravirtualisation) — guest OS modified to use hypercalls instead of privileged instructions; fastest')
B('HVM (Hardware-assisted) — unmodified guest OS using Intel VT-x/AMD-V; slower than PV but compatible')

H('Xen Network Virtualisation', lvl=2)
Code(
    'Physical NIC (eth0/peth0)\n'
    '        |\n'
    '    xenbr0 (software bridge in Dom0)\n'
    '   /        \\\n'
    'vif1.0      vif2.0      <- back-end virtual interfaces in Dom0\n'
    '  |              |\n'
    '[shared memory rings via event channels + grant tables]\n'
    '  |              |\n'
    'eth0(VM A)   eth0(VM B)  <- front-end PV interfaces in DomUs'
)
A('VM A → VM B data flow:')
B('VM A app → TCP/IP stack → eth0 front-end → shared-memory tx ring → event channel notifies Dom0')
B('Dom0 back-end reads from vif1.0 → hands frame to xenbr0 → xenbr0 looks up dst MAC → forwards to vif2.0')
B('vif2.0 → shared-memory rx ring → event channel notifies VM B → eth0 front-end → VM B TCP/IP → app')
A('Frame never leaves the host. All transfer is in-memory via grant tables — zero-copy, very fast.')

H('OS-Level Virtualization (Containers)', lvl=2)
A('Shares ONE host kernel. Isolation via two Linux kernel features:')
B('cgroups (Control Groups) — enforce resource limits/accounting per container group:')
Code(
    '  cpu (shares, CFS quota/period)   memory (limit, OOM policy)\n'
    '  blkio (disk IOPS/bandwidth)      pids (max processes)\n'
    '  net_cls/net_prio (traffic tag)   devices (allow/deny)\n'
    '  freezer (pause/resume)           cpuset (CPU affinity)'
)
B('namespaces — give each container its own VIEW of a resource:')
Code(
    '  pid   — separate process tree (container has its own PID 1)\n'
    '  net   — separate network interfaces, routing, iptables\n'
    '  mnt   — separate mount table / root filesystem\n'
    '  uts   — separate hostname and domain name\n'
    '  ipc   — separate System V IPC (semaphores, shared memory)\n'
    '  user  — separate UID/GID mapping (container root ≠ host root)\n'
    '  cgroup— hides cgroup hierarchy from container'
)
A('cgroups = HOW MUCH (resource limits); namespaces = WHAT YOU SEE (isolation)')

H('Docker Key Concepts', lvl=2)
B('Image — layered read-only filesystem (built from Dockerfile). Each RUN/COPY = one layer')
B('Container — running image + writable layer. Ephemeral by default')
B('Registry — Docker Hub, ECR, GCR — stores and distributes images')
B('Volume — persistent storage outside container filesystem (survives container removal)')
Code(
    'FROM python:3.11-slim          # base image\n'
    'WORKDIR /app                   # set working dir (creates if absent)\n'
    'COPY requirements.txt .        # copy file first (Docker layer cache)\n'
    'RUN pip install --no-cache-dir -r requirements.txt\n'
    'COPY . .                       # copy source after deps\n'
    'ENV FLASK_ENV=production       # environment variable\n'
    'EXPOSE 5000                    # declarative (not a firewall rule)\n'
    'USER nobody                    # non-root (security)\n'
    'ENTRYPOINT ["gunicorn"]        # fixed executable\n'
    'CMD ["--workers","4","app:app"]# default args (overridable at run)'
)
B('CMD vs ENTRYPOINT: ENTRYPOINT = fixed binary; CMD = default args. CMD overridable with docker run args.')
B('Why COPY requirements.txt first? Each layer is cached. Code changes don\'t invalidate the pip install layer.')
B('Docker CPU shares: CPU_i = (weight_i / Σweights) × total_CPUs. Weights: 1024+256+512+128=1920; 3 CPUs → 1.6, 0.4, 0.8, 0.2')

doc.add_page_break()

# ===========================================================
# SECTION 3: KUBERNETES
# ===========================================================
H('3. Kubernetes (K8s)', lvl=1)

H('Architecture', lvl=2)
A('Control Plane (Master) — brains:')
B('kube-apiserver — REST front-door; all components communicate through it; validates + persists state')
B('etcd — distributed KV store (Raft consensus); authoritative cluster state (desired state, secrets, configs)')
B('kube-scheduler — watches Pending pods, picks node based on: requests/capacity, affinity/anti-affinity, taints/tolerations, topology spread')
B('kube-controller-manager — runs reconciliation loops: ReplicaSet, Deployment, Node, Endpoint, ServiceAccount, Namespace, Job...')
B('cloud-controller-manager — integrates with cloud LBs, nodes, volumes (AWS/GCP/Azure specific)')

A('Data Plane (Worker Node) — brawn:')
B('kubelet — node agent; pulls pod specs from apiserver; tells container runtime to start/stop containers; reports health')
B('container runtime — containerd / CRI-O; actually runs containers via runc; OCI-compliant')
B('kube-proxy — programs iptables/IPVS rules to implement Service virtual IPs → pod IPs')
B('CNI plugin — pod networking: Flannel (overlay), Calico (BGP+policy), Weave, Cilium')

H('K8s Workflow (kubectl apply → running pod)', lvl=2)
Code(
    '[kubectl apply -f deploy.yaml]\n'
    '        |\n'
    '        v\n'
    '[kube-apiserver] ——validates——> [etcd]  (stores desired state)\n'
    '        |\n'
    '        v  (controllers watch apiserver via list/watch)\n'
    '[Deployment controller] → creates/updates ReplicaSet object\n'
    '[ReplicaSet controller] → creates Pod objects  (status: Pending)\n'
    '        |\n'
    '        v\n'
    '[kube-scheduler] → picks Node (checks requests vs available)\n'
    '                 → writes spec.nodeName to pod via apiserver\n'
    '        |\n'
    '        v\n'
    '[kubelet @ target Node] → watches pods assigned to this node\n'
    '  → invokes CRI (containerd): pull image, create container\n'
    '  → runc: sets up cgroups + namespaces\n'
    '  → CNI: wire pod network (IP, routes)\n'
    '  → runs liveness/readiness probes\n'
    '  → reports pod status → Running → apiserver → etcd\n'
    '        |\n'
    '        v\n'
    '[kube-proxy] → programs iptables/IPVS for any Service selecting this pod\n'
    '\n'
    'kubectl get pods → apiserver → returns status from etcd'
)

H('Key K8s Objects', lvl=2)
table2(
    ['Object', 'Purpose'],
    [
        ('Pod', 'Smallest deployable unit; 1+ containers share netns + volumes; ephemeral'),
        ('Deployment', 'Manages ReplicaSets; rolling-update, rollback, desired replica count'),
        ('ReplicaSet', 'Ensures N pod replicas running; self-healing (replaces crashed pods)'),
        ('Service', 'Stable VIP + DNS for a set of pods. Types: ClusterIP, NodePort, LoadBalancer, ExternalName'),
        ('Ingress', 'L7 HTTP routing rules → Services; TLS termination; requires Ingress Controller'),
        ('ConfigMap', 'Non-sensitive config key-value pairs; mounted as env vars or files'),
        ('Secret', 'Base64-encoded sensitive data (passwords, tokens); same mount options as ConfigMap'),
        ('PersistentVolume', 'Cluster-level storage resource (NFS, EBS, GCE PD)'),
        ('PVC', 'PersistentVolumeClaim — pod\'s request for storage; bound to a PV'),
        ('HPA', 'HorizontalPodAutoscaler — scales replicas based on CPU/memory/custom metrics'),
        ('DaemonSet', 'One pod per node (logging, monitoring agents)'),
        ('StatefulSet', 'Ordered startup/shutdown, stable network identity, persistent storage per pod'),
        ('Job / CronJob', 'Run-to-completion task / scheduled recurring task'),
    ]
)

H('Resource Scheduling', lvl=2)
B('Scheduler uses REQUESTS for scheduling decisions (not actual usage). Requests = "I need at least this much".')
B('Limits = hard cap: CPU throttled if exceeded; memory → OOM-killed if exceeded.')
B('available_for_new_pod = node_capacity − Σrequests_of_existing_pods')
B('Pod D can be scheduled ONLY IF Pod D\'s CPU-request ≤ CPU-unallocated AND memory-request ≤ memory-unallocated')
B('Currently-Unused ≠ Unallocated. Scheduler ignores actual usage; only looks at committed requests.')

H('kubectl Cheat Sheet', lvl=2)
Code(
    'kubectl get pods/nodes/svc/deploy/rs/pvc/cm/secrets [-n ns] [-o wide/yaml]\n'
    'kubectl describe pod <name>           # events + conditions + resource info\n'
    'kubectl logs <pod> [-c container] -f  # follow live logs\n'
    'kubectl exec -it <pod> -- /bin/sh     # shell into container\n'
    'kubectl apply -f file.yaml            # create or update\n'
    'kubectl delete pod/deploy/svc <name>\n'
    'kubectl scale deploy <name> --replicas=5\n'
    'kubectl rollout status/history/undo deploy/<name>\n'
    'kubectl top pods/nodes                # resource usage\n'
    'kubectl get events --sort-by=.metadata.creationTimestamp'
)

doc.add_page_break()

# ===========================================================
# SECTION 4: DATA CENTER NETWORKING
# ===========================================================
H('4. Data Center Networking & SDN', lvl=1)

H('Fat-Tree Topology', lvl=2)
B('Built from identical k-port commodity switches in 3 layers: Edge, Aggregation, Core')
B('k-port switches → k pods; each pod = (k/2) edge + (k/2) aggregation switches')
B('Core layer = (k/2)² switches; total servers = k³/4')
B('Full bisection bandwidth — any host can send at full line rate to any other host')
B('ECMP (Equal-Cost Multi-Path) routing — multiple equal-cost paths; load-balanced')
B('Strengths: uniform, cheap commodity switches, high bisection bandwidth, fault-tolerant multi-path')
B('Weaknesses: wiring complexity grows with k; scaling requires replacing all switches')

H('DCell', lvl=2)
B('Server-centric recursive architecture (servers do packet forwarding, not just switching)')
B('DCell_0: n servers + 1 mini-switch; DCell_k: (t_{k-1}+1) copies of DCell_{k-1} with one extra NIC per level')
B('With n=6: DCell_0=6 servers, DCell_1=42, DCell_2=1806... millions of servers with 4 levels')
B('Strengths: extreme scalability, high fault tolerance (many disjoint paths), low cost')
B('Weaknesses: servers must route (CPU overhead + complexity); longer paths than Fat-Tree')

H('Other Topologies', lvl=2)
B('BCube — server-centric for shipping-container DCs; multiple layers of switches + multi-port servers')
B('VL2 — Valiant Load Balancing; "Agility" (any server in any service pool); two-level Clos topology')

H('SDN (Software-Defined Networking)', lvl=2)
B('Separates: Control Plane (software, centralised SDN controller) from Data Plane (hardware switches — just forward)')
B('OpenFlow protocol — controller programs flow tables in switches (match header fields → action: forward/drop/modify)')
B('Benefits: programmable network, vendor-neutral APIs, global traffic engineering, easy policy changes')
B('Examples: OpenDaylight, ONOS, Google B4 WAN, VMware NSX')

H('Network Overlay', lvl=2)
B('VXLAN — Layer 2 frames encapsulated in UDP/IP packets. 24-bit VNI = 16M virtual segments. Used in AWS VPC, OpenStack Neutron')
B('GRE — Generic Routing Encapsulation; tunnels any network-layer protocol over IP')
B('NVGRE — Microsoft\'s alternative to VXLAN; uses GRE key for tenant isolation')
B('Purpose: multi-tenant isolation in cloud; VMs on different physical hosts appear on same L2 segment')

doc.add_page_break()

# ===========================================================
# SECTION 5: HADOOP
# ===========================================================
H('5. Hadoop Ecosystem', lvl=1)

H('HDFS Architecture', lvl=2)
Code(
    'Client\n'
    '  |\n'
    '  |--metadata ops (create/open/delete/rename)-->\n'
    '  |                                        NameNode\n'
    '  |                                        (namespace in RAM)\n'
    '  |                                        fsimage + edits log on disk\n'
    '  |\n'
    '  |--block read/write-->\n'
    '  |         DataNode 1  ——replication——>  DataNode 2  ——>  DataNode 3\n'
    '  |         (block 128MB, 3 replicas default)\n'
    '  |         (rack-aware: 2 same rack + 1 different rack)'
)
B('NameNode — stores filesystem namespace (directory tree + block→DataNode map) ENTIRELY IN RAM. Persistent: fsimage (snapshot) + edits log (journal). NEVER stores actual data. SPOF in Hadoop 1.')
B('Secondary NameNode — periodic checkpointer: merges fsimage+edits to prevent edits log growth. NOT a hot standby (confusingly named).')
B('Hadoop 2 HA: Active NameNode + Standby NameNode share edits via Quorum Journal Manager (QJM, 3+ JournalNodes). ZooKeeper ZKFC for automatic failover. Eliminates SPOF.')
B('DataNode — stores blocks on local disk. Heartbeat every 3s + block report every hour to NameNode.')

H('MapReduce Phases', lvl=2)
Code(
    '1. CLIENT submits job to ResourceManager (jar + input/output paths + config)\n'
    '2. RM allocates container, launches ApplicationMaster (AM)\n'
    '3. AM asks NameNode for input splits (1 split ≈ 1 HDFS block = 128MB)\n'
    '4. AM requests containers for Map tasks (prefer data-local: map on same node as block)\n'
    '\n'
    '5. MAP PHASE\n'
    '   RecordReader → key/value pairs → map() → intermediate (K,V)\n'
    '   Output buffered in circular memory buffer (80% → spill)\n'
    '   Spill: sort by partition+key → optional Combiner → write to local disk\n'
    '   At end: merge all spill files → one sorted partitioned file per mapper\n'
    '\n'
    '6. SHUFFLE & SORT\n'
    '   Reducers HTTP-fetch their partition from every mapper\'s local disk\n'
    '   Merge-sort fetched segments → sorted stream per reducer\n'
    '\n'
    '7. REDUCE PHASE\n'
    '   reduce(K, values_iterator) → output (K\',V\')\n'
    '   OutputFormat writes to HDFS (with replication)\n'
    '\n'
    '8. AM reports completion to RM → client notified → container released'
)
B('Combiner — optional mini-reducer after map-side. Reduces network traffic. ONLY valid if function is commutative+associative (sum/max/min — NOT average).')
B('Fault tolerance: failed task retried on another node (max 4 attempts). AM failure → RM restarts AM.')

H('YARN Components', lvl=2)
B('ResourceManager (RM) — master daemon. Scheduler (allocates containers: Capacity/Fair scheduler) + ApplicationsManager (accepts submissions, launches first AM container, restarts on failure). One per cluster.')
B('NodeManager (NM) — slave daemon on every worker. Launches containers, monitors resources, enforces cgroup limits, runs shuffle handler. Reports node health to RM.')
B('ApplicationMaster (AM) — per-job process. Negotiates containers from RM, launches tasks on NMs, retries failed tasks, reports progress, tells RM when done. Lives only for job duration.')

H('Big Data 4Vs', lvl=2)
table2(
    ['V', 'Meaning', 'Definition', 'Challenge'],
    [
        ('Volume', 'Data at Rest', 'Terabytes to Exabytes of existing data to process', 'Storage + parallel processing'),
        ('Velocity', 'Data in Motion', 'Streaming data; milliseconds to seconds to respond', 'Real-time ingestion + processing'),
        ('Variety', 'Data in Many Forms', 'Structured, unstructured, text, multimedia', 'Schema flexibility + ETL'),
        ('Veracity', 'Data in Doubt', 'Uncertainty due to inconsistency, incompleteness, ambiguity', 'Data quality + cleansing'),
    ]
)
A('AWS services map to 4Vs: S3=Volume, Kinesis=Velocity, Glue=Variety, DynamoDB/RDS=Veracity handling.')

H('Google Origins → Hadoop Equivalents', lvl=2)
table2(
    ['Google calls it', 'Hadoop equivalent', 'Purpose'],
    [
        ('MapReduce', 'Hadoop (YARN+MR)', 'Distributed batch computation framework'),
        ('GFS (Google File System)', 'HDFS', 'Distributed fault-tolerant file storage'),
        ('Bigtable', 'HBase', 'NoSQL wide-column store on top of HDFS'),
        ('Chubby', 'ZooKeeper', 'Distributed coordination / lock service'),
    ]
)

H('Hadoop Ecosystem Components', lvl=2)
table2(
    ['Component', 'Category', 'Purpose'],
    [
        ('HDFS', 'Storage', 'Distributed file system — block storage across DataNodes'),
        ('YARN', 'Resource Mgmt', 'Cluster resource manager — allocates CPU/RAM to applications'),
        ('MapReduce', 'Batch Compute', 'Two-phase (map+reduce) parallel batch processing'),
        ('Apache Spark', 'General Compute', 'In-memory DAG engine; replaces MR for most workloads'),
        ('Hive', 'SQL on Hadoop', 'Data warehouse; HiveQL (SQL-like) queries compile to MR/Spark jobs'),
        ('HBase', 'NoSQL DB', 'Wide-column store on HDFS; random read/write access to big data'),
        ('Pig', 'Scripting', 'Pig Latin scripting language; compiles to MapReduce'),
        ('Mahout', 'Machine Learning', 'Scalable ML algorithms (classification, clustering, CF) on Hadoop'),
        ('Sqoop', 'Data Integration', 'Import/export between HDFS and RDBMS (MySQL, Oracle)'),
        ('Flume', 'Ingestion', 'Collect and move large volumes of log data into HDFS'),
        ('ZooKeeper', 'Coordination', 'Distributed synchronisation, config management, leader election'),
        ('Storm/Flink', 'Streaming', 'Real-time stream processing (Storm=at-least-once, Flink=exactly-once)'),
        ('Tez', 'Workflow DAG', 'DAG execution engine; faster than MR for multi-stage Hive queries'),
    ]
)
A('Stack shorthand: Storage(HDFS/S3) → Resource(YARN) → Batch(MapReduce) | InMemory(Spark) → Apps(Hive/Pig/Spark SQL/Streaming)')

H('MapReduce Limitations (Why Spark was Needed)', lvl=2)
B('Good for: one-pass computation (one input, two-phase acyclic data flow) — batch, read-only datasets.')
B('BAD for iterative algorithms (e.g. ML, PageRank): state between steps written to HDFS → disk read next iteration → spends ~90% of time on I/O.')
B('BAD for interactive queries: each query = full MR job startup overhead.')
B('Rigid model: only Map + Reduce steps; complex multi-stage pipelines require chaining jobs with HDFS writes between them.')
B('Output of Reduce() must be replicated (3 copies) to HDFS → extra storage + network for each stage.')
B('Solution → Spark: keeps data in RAM between iterations (RDD cache); lineage-based fault tolerance without replication.')

H('Hadoop vs Spark', lvl=2)
table2(
    ['Aspect', 'Hadoop MapReduce', 'Apache Spark'],
    [
        ('Storage', 'Disk-based (intermediate writes HDFS)', 'In-memory (RDD partitions in RAM)'),
        ('Speed', 'Slower (disk I/O per stage)', '10-100x faster for iterative/interactive'),
        ('Model', '2-stage (Map + Reduce only)', 'DAG with many operators'),
        ('Fault tolerance', 'Re-read from HDFS', 'Lineage recomputation from source'),
        ('Workloads', 'Batch only', 'Batch, SQL, Streaming, ML, Graph'),
        ('Relationship', 'HDFS (storage) + YARN (resources)', 'Spark reads HDFS, runs on YARN'),
    ]
)
A('They are COMPLEMENTARY: Hadoop provides storage (HDFS) and resource management (YARN); Spark replaces MapReduce as the computation engine. Spark can also run standalone or on Kubernetes without Hadoop.')

doc.add_page_break()

# ===========================================================
# SECTION 6: APACHE SPARK
# ===========================================================
H('6. Apache Spark', lvl=1)

H('Architecture', lvl=2)
Code(
    '[Driver Program]\n'
    '  SparkContext / SparkSession\n'
    '  Builds DAG from transformations\n'
    '  Submits jobs → stages → tasks to Executors\n'
    '  Collects results (collect/show)\n'
    '        |\n'
    '        | (requests containers)\n'
    '        v\n'
    '[Cluster Manager]  YARN / Kubernetes / Standalone / Mesos\n'
    '        |\n'
    '        | (allocates)\n'
    '        v\n'
    '[Worker Node 1]           [Worker Node 2]\n'
    '  [Executor JVM]            [Executor JVM]\n'
    '    BlockManager (cache)      BlockManager (cache)\n'
    '    Thread pool               Thread pool\n'
    '    Task1 Task2 Task3         Task4 Task5 Task6'
)
B('Job — triggered by an Action. One job per action call.')
B('Stage — DAG split at shuffle boundaries. Tasks within a stage can be pipelined.')
B('Task — smallest unit of work. One task per partition per stage. Runs on one Executor thread.')
B('Executor — JVM on worker; runs tasks, caches partitions (BlockManager), reports metrics. Lives for app lifetime.')
B('Catalyst (DataFrame/SQL) — logical plan → analysed → optimised (predicate pushdown, constant folding, join reorder) → physical plan → Tungsten code generation')

H('Transformations vs Actions', lvl=2)
table2(
    ['TRANSFORMATIONS (lazy — return new RDD/DF)', 'ACTIONS (eager — trigger execution)'],
    [
        ('map, flatMap, filter, select, drop', 'collect, count, take(n), first'),
        ('withColumn, alias, cast, lit', 'reduce, foreach, toPandas'),
        ('groupBy, agg, reduceByKey, groupByKey', 'show, write.parquet, write.csv'),
        ('join, union, distinct, repartition, coalesce', 'saveAsTextFile, saveAsSequenceFile'),
        ('orderBy, sortBy, sortByKey, limit', ''),
        ('withWatermark, window (Streaming)', ''),
    ]
)

H('Narrow vs Wide (Shuffle Boundary)', lvl=2)
Code(
    'NARROW (no shuffle — pipeline-able, fast)    WIDE (shuffle — stage boundary, expensive)\n'
    '-------------------------------------------  -----------------------------------------\n'
    'map, flatMap, filter                         groupBy, reduceByKey, groupByKey\n'
    'select, withColumn, drop, alias              join (sort-merge or broadcast)\n'
    'union, mapPartitions                         distinct, orderBy, sortBy, sortByKey\n'
    'coalesce (decrease only, no full shuffle)    repartition (full shuffle)\n'
    '                                             cogroup, intersection'
)

H('RDD vs DataFrame', lvl=2)
table2(
    ['Aspect', 'RDD', 'DataFrame'],
    [
        ('Schema', 'No schema (arbitrary Python objects)', 'Schema-aware (column names + types)'),
        ('Optimisation', 'None — manual lambda chains', 'Catalyst optimizer + Tungsten codegen'),
        ('Performance', 'Lower (JVM↔Python serialisation in PySpark)', 'Higher (off-heap binary Tungsten format)'),
        ('API', 'map/filter/flatMap/reduceByKey/...', 'select/filter/groupBy/agg/join/...'),
        ('Use RDD when', 'Unstructured data, custom serialisation, low-level algorithms', 'Never — use DataFrame for structured data'),
    ]
)

H('Key RDD Operations (Syntax)', lvl=2)
Code(
    'sc = SparkContext("local[*]", "AppName")\n'
    '\n'
    'rdd = sc.textFile("hdfs:///path/file.txt")         # create from file\n'
    'rdd = sc.parallelize([1, 2, 3, 4, 5])              # create from collection\n'
    '\n'
    'rdd.map(lambda x: x * 2)                           # 1-to-1 transform\n'
    'rdd.flatMap(lambda line: line.split())             # 1-to-many (flatten)\n'
    'rdd.filter(lambda x: x > 0)                        # keep matching\n'
    'rdd.map(lambda w: (w, 1)).reduceByKey(lambda a,b: a+b)  # word count\n'
    'rdd.sortBy(lambda kv: -kv[1])                      # sort by custom key (desc)\n'
    'rdd.sortByKey()                                     # sort by natural key\n'
    'rdd.join(other_rdd)                                # inner join on key\n'
    'rdd.cache()  / rdd.persist()                       # store for reuse\n'
    'rdd.unpersist()                                    # release memory\n'
    '\n'
    '# ACTIONS:\n'
    'rdd.take(10)          # first 10 elements to driver\n'
    'rdd.collect()         # ALL elements to driver (small data only!)\n'
    'rdd.count()           # number of elements\n'
    'rdd.first()           # first element\n'
    'rdd.saveAsTextFile("hdfs:///output/")'
)

H('Key DataFrame Operations (Syntax)', lvl=2)
Code(
    'from pyspark.sql import SparkSession\n'
    'from pyspark.sql.functions import col, avg, max, min, sum, count\n'
    'from pyspark.sql.functions import desc, asc, when, lit, from_unixtime\n'
    'from pyspark.sql.functions import to_date, split, explode, lower, upper, trim\n'
    'from pyspark.sql.functions import udf, rank, dense_rank, row_number\n'
    'from pyspark.sql.window import Window\n'
    '\n'
    'spark = SparkSession.builder.appName("App").getOrCreate()\n'
    '\n'
    '# Read\n'
    'df = spark.read.csv("path", header=True, inferSchema=True)\n'
    'df = spark.read.json("path/*.json")\n'
    'df = spark.read.parquet("s3://bucket/path/")\n'
    '\n'
    '# Transform (all LAZY)\n'
    'df.filter(col("x") > 0)\n'
    'df.withColumn("new", col("a") * 2)\n'
    'df.select("a", "b").drop("c")\n'
    'df.groupBy("cat").agg(count("*").alias("n"), avg("v").alias("avg_v"))\n'
    'df.orderBy(desc("score"))\n'
    'df.limit(10)\n'
    'df.join(other, "key")                  # inner join\n'
    'df.join(other, "key", "left")          # left outer join\n'
    'df.withColumn("label",\n'
    '    when(col("x") < 0, "neg").when(col("x") == 0, "zero").otherwise("pos"))\n'
    '\n'
    '# Cache\n'
    'df.cache()                             # MEMORY_ONLY\n'
    'df.unpersist()                         # release\n'
    '\n'
    '# Actions\n'
    'df.show(20)                            # prints 20 rows\n'
    'df.count()\n'
    'df.collect()                           # returns list of Rows (small data only)\n'
    'df.write.mode("overwrite").parquet("s3://output/")  # write Parquet\n'
    'df.write.mode("append").csv("output/")\n'
    '\n'
    'spark.stop()'
)

H('Spark Ecosystem Components', lvl=2)
table2(
    ['Component', 'Purpose'],
    [
        ('Spark Core', 'Base engine: task scheduling, memory management, fault recovery, I/O. Foundation for all other components.'),
        ('Spark SQL / DataFrame', 'Structured data; SQL queries; Catalyst optimizer; Hive/JDBC/Parquet/JSON/CSV sources.'),
        ('Spark Streaming', 'Real-time micro-batch processing; DStream API; integrates with Kafka/Kinesis/Flume.'),
        ('MLlib', 'Scalable ML library on Spark. 10-100x faster than Hadoop Mahout for iterative algorithms.'),
        ('GraphX', 'Graph computation on RDDs; PageRank, connected components, triangle counting.'),
    ]
)
A('Language APIs: Python (PySpark), Scala (native), Java, R. PySpark = most common in labs/exams.')

H('Spark MLlib — Supported Algorithms', lvl=2)
B('Classification: logistic regression, linear SVM, naïve Bayes, classification/decision trees')
B('Regression: GLMs (generalised linear models), regression trees')
B('Collaborative Filtering: ALS (alternating least squares), NMF (non-negative matrix factorisation) — for recommendation systems')
B('Clustering: k-means, Gaussian mixture, LDA (topic modelling)')
B('Dimensionality reduction: SVD (singular value decomposition), PCA')
B('Optimisation: stochastic gradient descent (SGD), L-BFGS')

H('Amazon EMR (Elastic MapReduce)', lvl=2)
B('Managed Hadoop/Spark framework on AWS. Runs MapReduce + Spark + Hive + Pig on YARN.')
B('Tight integration: reads from S3 (instead of HDFS), writes to S3/DynamoDB/Redshift.')
Code(
    'EMR Stack (bottom to top):\n'
    '  Storage:    S3 (object) or HDFS (local cluster disk)\n'
    '  Resource:   YARN (cluster resource management)\n'
    '  Compute:    MapReduce (batch) | Spark (in-memory)\n'
    '  Apps:       Pig, Hive, Cascading, Spark Streaming, Spark SQL\n'
    '\n'
    'Typical batch pipeline:\n'
    '  Data Sources → S3 (staging) → EMR (ETL/transform) → S3 → Redshift → QuickSight (BI)\n'
    '\n'
    'Typical streaming pipeline:\n'
    '  Apps/Devices → Kinesis Streams → KCL App + Lambda → DynamoDB → Real-time dashboard'
)
B('Cost model: pay per EC2 instance-hour. Use spot instances for workers; auto-terminate when done to avoid idle charges.')
B('AWS data type mapping: File data → S3; Stream data → Kinesis; Transactional data → DynamoDB / RDS Aurora.')

H('Critical Spark Rules', lvl=2)
B('map vs flatMap: map=1-to-1; flatMap=1-to-many (flattens). Use flatMap when splitting lines into words.')
B('reduceByKey vs groupByKey: reduceByKey runs map-side combiner (less network traffic); groupByKey shuffles ALL values then aggregates (much worse for large value sets). ALWAYS prefer reduceByKey.')
B('cache() vs persist(): cache()=MEMORY_ONLY; persist(StorageLevel.MEMORY_AND_DISK) spills to disk. Use when same DF/RDD used in 2+ actions.')
B('Double computation: two actions on un-cached DF re-execute the full lineage TWICE. Fix: cache() between actions.')
B('repartition(n) = full shuffle (any n); coalesce(n) = narrow (can only decrease). Coalesce to reduce files before writing; repartition to rebalance before joins.')
B('Parquet vs CSV: columnar (skip unused columns), compressed (5-10x smaller), schema embedded, predicate pushdown. Always prefer Parquet for analytics.')

doc.add_page_break()

# ===========================================================
# SECTION 7: SOA & MICROSERVICES
# ===========================================================
H('7. SOA, Microservices & Cloud-Native', lvl=1)

H('SOA Principles (8)', lvl=2)
B('Standardized service contract — services expose well-defined interfaces (WSDL/REST/OpenAPI)')
B('Loose coupling — services depend on each other only through published APIs; no shared DB')
B('Service abstraction — implementation details hidden from consumers')
B('Service reusability — designed to be reused across multiple applications/consumers')
B('Service autonomy — services control their own logic and data; independently deployable')
B('Service statelessness — services do not hold client state between invocations (state in DB/cache)')
B('Service discoverability — services can be found (service registry, API Gateway, DNS)')
B('Service composability — services can be orchestrated to build higher-level services')

H('Service Composition & ESB', lvl=2)
B('Service Composition — combining multiple component services to fulfil a higher-level request. E.g., travel plan service invokes flight + hotel + cab services independently.')
B('Service consumer — entity that invokes a service. Service provider — entity that exposes its functions as services.')
B('Two composition topologies:')
Code(
    'Point-to-Point (direct):          Hub-and-Spoke (ESB):\n'
    '  App A ──────► App B               App A ──┐\n'
    '  App A ──────► App C               App B ──┤──► ESB ──► App D\n'
    '  (consumer knows endpoint URL)     App C ──┘      (routing + transform)\n'
    '  Simple but brittle at scale       ESB = Enterprise Service Bus'
)
B('ESB (Enterprise Service Bus) — middleware broker: routes messages between services, transforms data formats, manages a service directory. Used in classical SOA (not microservices).')
B('SLA (Service Level Agreement) — contract between provider and consumer defining availability, latency, throughput guarantees.')

H('Stateful vs Stateless Servers', lvl=2)
table2(
    ['Aspect', 'Stateful', 'Stateless'],
    [
        ('Server stores', 'Session state (login, basket, context)', 'Nothing — no client state between requests'),
        ('Request dependency', 'Request depends on prior state', 'Each request is independent transaction'),
        ('Examples', 'E-commerce session, banking login, gaming', 'REST APIs, static content, search'),
        ('Scalability', 'Harder — session affinity or shared store', 'Easy — any server can handle any request'),
        ('Cloud fit', 'Requires Redis/DynamoDB for state storage', 'Perfect for serverless/Lambda'),
    ]
)
A('REST is stateless by design — each request carries all info needed (JWT token, params). Sessions = server-side state = stateful. Flask session cookie = client-side state (signed cookie, not server-stored).')

H('SOA vs Microservices vs Monolith', lvl=2)
table2(
    ['Aspect', 'Monolith', 'SOA', 'Microservices'],
    [
        ('Size', 'One large app', 'Large services (bounded by domain)', 'Fine-grained (one responsibility)'),
        ('Communication', 'In-process calls', 'ESB (Enterprise Service Bus)', 'REST/gRPC/async messaging'),
        ('Database', 'One shared DB', 'Sometimes shared DB', 'Each service owns its own DB'),
        ('Deploy', 'Deploy everything together', 'Per-service (but heavyweight)', 'Independent CI/CD per service'),
        ('Tech stack', 'One language/framework', 'Often one language', 'Polyglot (each service differs)'),
    ]
)

H('SOAP vs REST', lvl=2)
table2(
    ['Aspect', 'SOAP', 'REST'],
    [
        ('Protocol', 'HTTP/SMTP/TCP — any transport', 'HTTP/HTTPS only'),
        ('Format', 'XML only, strict schema', 'JSON/XML/YAML/plain text'),
        ('Contract', 'WSDL (mandatory, machine-readable)', 'OpenAPI/Swagger (optional)'),
        ('State', 'Can be stateful (WS-*)', 'Stateless (each request self-contained)'),
        ('Security', 'WS-Security built-in spec', 'OAuth2 / JWT / TLS'),
        ('Overhead', 'Heavy (XML parsing, envelope)', 'Light (JSON, minimal overhead)'),
    ]
)

H('Event-Driven Architecture', lvl=2)
B('SNS (pub/sub) — topic; publishers push messages; subscribers (SQS/Lambda/HTTP/email/SMS) receive. Fan-out: 1 SNS → many SQS queues.')
B('SQS — queue; consumers PULL; decouples producers from consumers. Visibility timeout (default 30s) — message hidden from other consumers while being processed. DLQ (Dead Letter Queue) for failed messages after max retries.')
table2(
    ['Feature', 'Standard Queue', 'FIFO Queue'],
    [
        ('Throughput', 'Unlimited TPS (nearly)', '300 TPS (3000 with batching)'),
        ('Ordering', 'Best-effort (not guaranteed)', 'Strictly FIFO — preserved'),
        ('Delivery', 'At-least-once (duplicates possible)', 'Exactly-once processing'),
        ('Use case', 'High-throughput, order-tolerant jobs', 'Financial transactions, order processing'),
        ('Example', 'Schedule DB inserts (order irrelevant)', 'Prevent double-enrolment in course'),
    ]
)
B('Kinesis Data Streams — ordered, replayable, shards (1MB/s in, 2MB/s out). Consumers can replay. Retention 24h-365d.')
B('Kinesis Firehose — managed delivery to S3/Redshift/OpenSearch; auto-scale; optional Lambda transform.')
B('EventBridge — serverless event bus; cron rules; event pattern matching; targets: Lambda/SQS/SNS/Step Functions.')

H('12-Factor App (Cloud-Native)', lvl=2)
B('Config in env vars (not code). Stateless processes (state in backing services). Port binding. Dev/prod parity.')
B('Service mesh (Istio): sidecar proxy (Envoy) beside each pod → handles mTLS, circuit breaking, retries, observability without app changes.')

doc.add_page_break()

# ===========================================================
# SECTION 8: AWS SERVICES REFERENCE
# ===========================================================
H('8. AWS Services Quick Reference', lvl=1)

H('Compute', lvl=2)
B('Lambda — event-driven, max 15min, 10GB RAM/6 vCPU. Cold start ~100-500ms (module-level code runs once). Warm container reused. handler(event, context). Concurrency = requests × duration.')
B('EC2 — IaaS VMs. Instance families: t3 (burstable), m5 (general), c5 (compute), r5 (memory), p3 (GPU). EBS-backed. Security Groups (stateful). Key pairs for SSH.')
B('ECS — AWS managed container orchestration (Fargate = serverless containers; EC2 launch type = you manage nodes)')
B('EKS — managed Kubernetes control plane; you manage worker nodes (or use Fargate for serverless pods)')
B('Elastic Beanstalk — PaaS; upload app zip/Docker; AWS provisions EC2+ELB+ASG+RDS. Good for Flask/Django/Node/Java.')

H('Storage', lvl=2)
B('S3 — object store; 11 nines durability; classes: Standard, IA, One-Zone-IA, Glacier Instant/Flexible/Deep Archive. Presigned URLs (temporary access). Event notifications → Lambda/SQS/SNS. Static website hosting.')
B('EBS — block storage attached to ONE EC2. Types: gp3 (SSD general), io2 (high IOPS), st1 (throughput HDD), sc1 (cold HDD). Snapshots to S3.')
B('EFS — managed NFS; shared across multiple EC2; POSIX-compliant; auto-scale.')

H('Database', lvl=2)
B('DynamoDB — serverless NoSQL KV+document. Partition key (required) + optional sort key. GSI (alternate PK, eventual consistency, separate throughput). LSI (same PK, different SK, strongly consistent). PAY_PER_REQUEST or provisioned. TTL (auto-delete items). Max item 400KB. Streams → Lambda. DAX (in-memory cache μs latency).')
B('RDS — managed relational (MySQL/PostgreSQL/Oracle/MSSQL/MariaDB). Multi-AZ (sync standby, auto-failover). Read Replicas (async, cross-region, up to 15 for Aurora). RDS Proxy for Lambda connection pooling.')
B('ElastiCache — Redis (persistence, pub/sub, sorted sets, cluster mode) or Memcached (pure cache, multi-thread, no persistence).')

H('Networking', lvl=2)
B('VPC — isolated virtual network; CIDR block; public subnets (route to IGW); private subnets (no IGW, outbound via NAT GW). Security Groups (stateful, instance-level). NACLs (stateless, subnet-level, allow+deny rules).')
B('Route 53 — DNS; routing policies: Simple, Weighted (A/B testing), Latency (closest region), Geolocation, Failover (health check), Multivalue.')
B('CloudFront — CDN; 400+ edge locations; origin: S3/ALB/API GW/EC2. HTTPS termination, WAF integration, OAC for S3.')
B('ELB — ALB (L7 HTTP/HTTPS, path+host routing, WebSocket, Lambda target); NLB (L4 TCP/UDP, static IP, ultra-low latency); Classic (legacy L4/L7, avoid).')
B('API Gateway — REST/HTTP/WebSocket APIs. Stages (dev/prod/v1). Cognito/Lambda authorizer. Lambda integration (proxy). Usage Plans + API Keys for rate limiting. Throttling (rate+burst).')

H('Auth & Security', lvl=2)
B('Cognito User Pools — user directory; sign-up/sign-in; MFA; social federation (Google/FB); issues JWT (id_token, access_token, refresh_token). Integrate with API GW as authorizer.')
B('Cognito Identity Pools — federated identities; temporary AWS credentials (STS AssumeRoleWithWebIdentity) for direct AWS access (S3, DynamoDB).')
B('IAM — users, groups, roles, policies (Effect/Action/Resource/Condition JSON). Least privilege. Never use root. Service roles for Lambda (execution role)/EC2 (instance profile).')
B('KMS — managed encryption keys; CMK; envelope encryption. Used by S3/DynamoDB/EBS/SM.')
B('Secrets Manager — stores DB passwords/API keys; auto-rotation; fetch at Lambda cold start; cache in module scope.')

H('Monitoring', lvl=2)
B('CloudWatch — metrics (5min default, 1min detailed), Logs (log groups/streams, Insights query), Alarms (threshold → SNS/ASG action), Dashboards.')
B('X-Ray — distributed tracing; service map; segments/subsegments; sampling. Useful for Lambda+APIGW+DynamoDB chains.')

H('IoT', lvl=2)
B('IoT Core — MQTT broker (TLS, X.509 mutual auth, QoS 0/1). Device Registry (manage device identities). Device Shadow (last-known state; delta updates; offline devices sync on reconnect). Rules Engine (SQL-like; route MQTT → S3/DynamoDB/Lambda/Kinesis/SNS).')

doc.add_page_break()

# ===========================================================
# SECTION 9: FLASK & CLOUD APPS (Lab content)
# ===========================================================
H('9. Flask Web Framework & Cloud App Patterns', lvl=1)

H('Flask Core Patterns', lvl=2)
Code(
    'from flask import Flask, request, jsonify, render_template\n'
    'from flask import redirect, url_for, session, abort, g\n'
    '\n'
    'app = Flask(__name__)\n'
    'app.secret_key = "change-me"          # required for sessions\n'
    '\n'
    '# Routes\n'
    '@app.route("/api/<int:id>", methods=["GET","POST","PUT","DELETE"])\n'
    'def resource(id): ...\n'
    '@app.route("/path/<string:name>")     # converters: int, float, string, path, uuid\n'
    '\n'
    '# Request\n'
    'request.method                         # GET/POST/PUT/DELETE\n'
    'request.args.get("key", default)       # URL query string ?key=val\n'
    'request.form.get("field")              # form POST (application/x-www-form-urlencoded)\n'
    'request.get_json(force=True)           # JSON body (force ignores Content-Type header)\n'
    'request.files["upload"]               # file upload (.filename, .save(path))\n'
    'request.headers.get("Authorization")  # HTTP headers\n'
    'request.remote_addr                   # client IP\n'
    '\n'
    '# Response\n'
    'return jsonify({"key": "val"}), 200\n'
    'return "", 204                          # no content (DELETE)\n'
    'return jsonify(error="msg"), 404\n'
    'return redirect(url_for("func_name", param=val))\n'
    'abort(403)                             # raises HTTPException immediately\n'
    '\n'
    '# Error handlers\n'
    '@app.errorhandler(404)\n'
    'def not_found(e): return jsonify(error="not found"), 404\n'
    '@app.errorhandler(Exception)\n'
    'def handle_all(e): return jsonify(error=str(e)), 500\n'
    '\n'
    '# Before/after request\n'
    '@app.before_request\n'
    'def check_auth():\n'
    '    if request.endpoint in {"login","health"}: return\n'
    '    if not session.get("user"): return jsonify(error="unauthorized"), 401\n'
    '\n'
    '# Session\n'
    'session["user"] = "alice"             # stored in signed cookie (client-side)\n'
    'session.get("user")                   # read\n'
    'session.clear()                       # logout\n'
    '\n'
    '# Blueprint\n'
    'from flask import Blueprint\n'
    'bp = Blueprint("flights", __name__)\n'
    '@bp.route("/flights")\n'
    'def list_flights(): ...\n'
    '# In app.py:\n'
    'app.register_blueprint(bp, url_prefix="/api/v1")\n'
    '\n'
    '# Templates (Jinja2)\n'
    'return render_template("page.html", name=name, items=items)\n'
    '# {{ var }}  {{ val|upper|trim }}  {{ val|default("N/A") }}\n'
    '# {% for x in items %}...{% endfor %}\n'
    '# {% if cond %}...{% elif %}...{% else %}...{% endif %}\n'
    '# {% extends "base.html" %} {% block content %}...{% endblock %}'
)

H('Flask + AWS boto3 Patterns', lvl=2)
Code(
    'import boto3\n'
    'from boto3.dynamodb.conditions import Key, Attr\n'
    '\n'
    '# DynamoDB (module-scope for warm Lambda container reuse)\n'
    'table = boto3.resource("dynamodb").Table("MyTable")\n'
    'table.put_item(Item={"pk":"id1","attr":"val"})\n'
    'resp = table.get_item(Key={"pk":"id1"}); item = resp.get("Item")\n'
    'resp = table.query(IndexName="gsi", KeyConditionExpression=Key("gsi_pk").eq("v"))\n'
    'table.update_item(Key={"pk":"id1"},\n'
    '    UpdateExpression="SET #n = :v",\n'
    '    ExpressionAttributeNames={"#n":"Name"},  # # = reserved word alias\n'
    '    ExpressionAttributeValues={":v":"Alice"}) # : = value placeholder\n'
    'table.delete_item(Key={"pk":"id1"})\n'
    'resp = table.scan(FilterExpression=Attr("dept").eq("CS"))  # avoid at scale\n'
    '\n'
    '# S3\n'
    's3 = boto3.client("s3")\n'
    's3.upload_fileobj(file_obj, "bucket", "key")\n'
    'url = s3.generate_presigned_url("get_object",\n'
    '    Params={"Bucket":"bucket","Key":"key"}, ExpiresIn=3600)\n'
    'resp = s3.list_objects_v2(Bucket="bucket", Prefix="folder/")\n'
    '\n'
    '# Lambda handler pattern\n'
    'import json\n'
    'def lambda_handler(event, context):\n'
    '    # event = API GW proxy event {httpMethod, pathParameters, body, headers}\n'
    '    # context = {function_name, aws_request_id, get_remaining_time_in_millis()}\n'
    '    item_id = event["pathParameters"]["id"]\n'
    '    resp = table.get_item(Key={"ID": item_id})\n'
    '    item = resp.get("Item")\n'
    '    if not item:\n'
    '        return {"statusCode": 404, "body": json.dumps({"error":"not found"})}\n'
    '    return {"statusCode": 200, "body": json.dumps(item, default=str)}\n'
    '    # default=str: handles Decimal, datetime (DynamoDB types not JSON-serializable)'
)

H('Flask + MySQL (Lab pattern)', lvl=2)
Code(
    'import mysql.connector\n'
    '\n'
    'DB_CONFIG = {"host":"...rds.amazonaws.com","user":"admin",\n'
    '             "password":"secret","database":"mydb","port":3306}\n'
    '\n'
    'def get_db():\n'
    '    return mysql.connector.connect(**DB_CONFIG)\n'
    '\n'
    '@app.route("/data")\n'
    'def get_data():\n'
    '    con = get_db()\n'
    '    cur = con.cursor(dictionary=True)   # rows as dicts\n'
    '    cur.execute("SELECT * FROM t WHERE id=%s", (id,))  # parameterized!\n'
    '    rows = cur.fetchall()\n'
    '    cur.close(); con.close()\n'
    '    return jsonify(rows), 200'
)

H('Flask + Elastic Beanstalk Deployment (Lab-04 pattern)', lvl=2)
B('application.py (NOT app.py) — EB looks for "application" object')
B('requirements.txt — EB pip installs automatically')
B('eb init → eb create → eb deploy → eb open')
B('Environment variables set in EB console → os.environ.get("DB_HOST")')
B('EB auto-creates: EC2 + ELB (ALB) + Auto Scaling Group + Security Groups + CloudWatch alarms')

doc.add_page_break()

# ===========================================================
# SECTION 10: PAST PAPER PATTERN ANALYSIS
# ===========================================================
H('10. Past Paper Pattern Analysis', lvl=1)

H('Paper Structure (Confirmed from 3 papers)', lvl=2)
table2(
    ['Question', 'Marks', '2023 Topic', '2024 Topic', '2025 Topic'],
    [
        ('Q1', '10', '(omitted in PDF)', '(omitted in PDF)', '(omitted in PDF)'),
        ('Q2', '15', '5 short Qs: MapReduce/Spark/NameNode/Executors/Hadoop+Spark', '5 short: Virtualization/DC networks/Hadoop+Spark/Transform+Action', '5 short: IaaS/PaaS/Hypervisor/K8s components/Hadoop+Spark/Transform+Action'),
        ('Q3', '15', 'SOA design diagram + AWS services for social media app', 'Xen network virtualisation + VM A→B data flow (with figure)', 'K8s resource scheduling (Pod D) + Docker CPU shares (4 containers, 3 CPUs)'),
        ('Q4', '20', 'Xen Type-1/2 + Type-1 vs Type-2 + OS-Level virtualization', 'Same as 2023 Q4 (identical)', 'Hadoop MapReduce workflow + HDFS/MR/YARN + AM/RM/NM/NN/DN roles'),
        ('Q5', '20', 'Healthcare IoT on AWS (architecture + ML extension)', 'K8s: cgroups+namespaces + architecture + workflow', 'K8s: cgroups+namespaces + architecture + workflow (identical to 2024)'),
        ('Q6', '20', 'PySpark RDD: MaxTempByCity (line desc/map+reduceByKey/RDD vs DF)', 'PySpark DataFrame: Sales Analysis (line desc/Transform/Action/architecture)', 'PySpark DataFrame: E-commerce (line desc/Transform/Action/Spark components)'),
    ]
)

H('High-Frequency Topics (Every Paper)', lvl=2)
B('Hadoop vs Spark relationship — answer: HDFS for storage, YARN for resources, Spark replaces MR compute, complementary not competitors')
B('Transformation vs Action — ALWAYS asked. Transformations=lazy/DAG building; Actions=trigger execution/return result')
B('PySpark code analysis — describe each line, classify T vs A, Q6 always has sub-parts: (a) line descriptions (b) what is Transformation (c) what is Action (d) Spark architecture')
B('Hypervisor Type-1 vs Type-2 — appeared in 2023 Q4 and 2024 Q4 (IDENTICAL question)')
B('Kubernetes — appeared in 2024 Q5 and 2025 Q5 (IDENTICAL question: cgroups+namespaces, architecture, workflow)')

H('High-Confidence 2026 Predictions', lvl=2)
B('Q2: Almost certain — transformation vs action + Hadoop vs Spark. Other 3 from: cloud service models, SOA principles, DC network topology, streaming (Kinesis vs SQS), serverless')
B('Q3: Docker+K8s (scheduling/YAML/networking) OR SOA/microservices design OR AWS Lambda serverless architecture')
B('Q4: Either new Hadoop/Spark topic (streaming, optimization, DataFrames) OR AWS ML pipeline (SageMaker workflow)')
B('Q5: AWS IoT architecture (hasn\'t appeared since 2023) OR AWS event-driven design (SNS+SQS+Lambda+DynamoDB)')
B('Q6: PySpark DataFrame — expect: read parquet/json → withColumn (when/otherwise) → groupBy → agg → orderBy → write. Sub-questions same pattern.')

H('Answer Templates', lvl=2)
A('3-mark short answer: "[Term] is [definition, 1 sentence]. [Key feature 1 that distinguishes it]. [Key feature 2 or contrast with related concept]."')
A('8-mark architecture: "[ASCII diagram of system]. AWS services: [bullet list with each service\'s purpose]. Orchestration flow: [data flow, step by step]."')
A('PySpark line description: "Line N: [TRANSFORMATION/ACTION]. [What it does to the data]. [Why it is T or A — lazy/triggers execution, narrow/wide]."')
A('Spark component description: "[Name] — [role]. [Key responsibility]. [How it interacts with other components]."')

doc.add_page_break()

# ===========================================================
# SECTION 11: CODING QUICK REFERENCE CARDS
# ===========================================================
H('11. Coding Quick Reference Cards', lvl=1)

H('PySpark DataFrame Job Skeleton', lvl=2)
Code(
    'from pyspark.sql import SparkSession\n'
    'from pyspark.sql.functions import col, avg, max, min, sum, count, desc, asc, when\n'
    '\n'
    'spark = SparkSession.builder.appName("JobName").getOrCreate()\n'
    '\n'
    '# Read\n'
    'df = spark.read.csv("s3://bucket/*.csv", header=True, inferSchema=True)\n'
    '# df = spark.read.json("path")  /  spark.read.parquet("s3://bucket/path/")\n'
    '\n'
    '# Clean (narrow transformations)\n'
    'clean = df.filter(col("field").isNotNull() & (col("field") >= 0))\n'
    '\n'
    '# Classify with when/otherwise\n'
    'clean = clean.withColumn("label",\n'
    '    when(col("val").isNull(), "unknown")\n'
    '    .when(col("val") < 0,   "negative")\n'
    '    .when(col("val") == 0,  "zero")\n'
    '    .otherwise(             "positive")\n'
    ')\n'
    '\n'
    '# Aggregate (wide transformations)\n'
    'result = (\n'
    '    clean\n'
    '    .groupBy("group_col")\n'
    '    .agg(\n'
    '        count("*").alias("n"),\n'
    '        avg("metric").alias("avg_metric"),\n'
    '        max("metric").alias("max_metric")\n'
    '    )\n'
    '    .orderBy(desc("avg_metric"))\n'
    ')\n'
    '\n'
    '# Cache if 2+ actions on same DF\n'
    'result.cache()\n'
    'result.show(20)                                        # ACTION 1\n'
    'result.write.mode("overwrite").parquet("s3://out/")   # ACTION 2\n'
    'result.unpersist()\n'
    'spark.stop()'
)

H('PySpark RDD Job Skeleton', lvl=2)
Code(
    'from pyspark import SparkContext\n'
    'sc = SparkContext("local[*]", "AppName")\n'
    '\n'
    'rdd = (\n'
    '    sc.textFile("hdfs:///path/file.txt")\n'
    '    .flatMap(lambda line: line.lower().split())    # 1-to-many\n'
    '    .map(lambda w: (w, 1))                        # (word, 1) pairs\n'
    '    .reduceByKey(lambda a, b: a + b)              # sum per word (map-side combine)\n'
    '    .filter(lambda kv: kv[1] > 5)                # keep frequent words\n'
    '    .sortBy(lambda kv: -kv[1])                   # sort desc by count\n'
    ')\n'
    'rdd.cache()\n'
    'for word, cnt in rdd.take(20): print(f"{word}: {cnt}")   # ACTION 1\n'
    'rdd.saveAsTextFile("hdfs:///output/")                    # ACTION 2\n'
    'rdd.unpersist()\n'
    'sc.stop()'
)

H('HTTP Status Codes', lvl=2)
Code(
    '200 OK              — GET/PUT/PATCH success, response body present\n'
    '201 Created         — POST success, new resource created\n'
    '204 No Content      — DELETE success, no response body\n'
    '400 Bad Request     — invalid input, missing required field\n'
    '401 Unauthorized    — not authenticated (no/bad token/credentials)\n'
    '403 Forbidden       — authenticated but insufficient permissions\n'
    '404 Not Found       — resource does not exist at this URL\n'
    '405 Method Not All. — URL exists but HTTP verb not in methods=[...]\n'
    '409 Conflict        — resource already exists (duplicate create)\n'
    '422 Unprocessable   — valid JSON but fails business validation\n'
    '429 Too Many Req    — rate limit exceeded (API Gateway throttle)\n'
    '500 Server Error    — unhandled exception in the server\n'
    '502 Bad Gateway     — upstream service (Lambda) failed\n'
    '503 Unavailable     — service overloaded or down'
)

H('Key Numbers to Remember', lvl=2)
Code(
    'HDFS block size:       128 MB (default)\n'
    'HDFS replication:      3 (default: 2 same rack + 1 different rack)\n'
    'MapReduce task retries:4 (mapreduce.task.maxattempts)\n'
    'YARN heartbeat:        every 3s (DataNode → NameNode)\n'
    'DynamoDB max item:     400 KB\n'
    'Lambda max timeout:    15 minutes\n'
    'Lambda max RAM:        10 GB\n'
    'Lambda cold start:     ~100-500 ms (Python)\n'
    'SQS visibility timeout:30s (default); max 12 hours\n'
    'S3 durability:         11 nines (99.999999999%)\n'
    'Fat-Tree (k-port):     k^3/4 servers, (k/2)^2 core switches\n'
    'Docker CPU shares:     CPU_i = (weight_i / sum_weights) × total_CPUs\n'
    '  Example: 1024+256+512+128=1920 weights, 3 CPUs:\n'
    '  C1=1024/1920×3=1.6  C2=256/1920×3=0.4\n'
    '  C3=512/1920×3=0.8   C4=128/1920×3=0.2  (sum=3.0 ✓)\n'
    'SLA 99.9% = 8.76 hrs downtime/year\n'
    'SLA 99.99% = 52.6 min downtime/year'
)

H('Common Exam Pitfalls', lvl=2)
B('Scheduler uses REQUESTS not actual usage. Pod scheduled based on unallocated requests, not "Currently Unused".')
B('Secondary NameNode is NOT a hot standby — it\'s a checkpointer. Hot standby = Standby NameNode (HA mode).')
B('groupByKey shuffles ALL values; reduceByKey pre-aggregates. Always prefer reduceByKey for commutative+associative ops.')
B('Both actions on an un-cached DF trigger full lineage re-execution. Always cache() between two actions.')
B('DynamoDB scan is expensive (reads entire table). Always use query() with GSI for filtered lookups.')
B('Type-1 hypervisor runs ON hardware. Type-2 runs ON TOP OF a host OS. Xen = Type-1.')
B('cgroups = resource LIMITS (how much); namespaces = resource ISOLATION (what you see).')
B('ENTRYPOINT is fixed; CMD provides default args and is overridable. CMD alone is fine for simple containers.')
B('HDFS NameNode stores METADATA in RAM (not actual data). DataNodes store actual block data on disk.')
B('reduceByKey = wide transformation (shuffle). map/filter/flatMap = narrow (no shuffle).')

# ===================== END =====================
doc.add_page_break()
end = doc.add_heading('End of COMP4442 Exam Notes', level=1)
end.alignment = WD_ALIGN_PARAGRAPH.CENTER
closing = doc.add_paragraph(
    'Sections: 1-Cloud Fundamentals | 2-Virtualization | 3-Kubernetes | 4-DC Networking | '
    '5-Hadoop | 6-Spark | 7-SOA/Microservices | 8-AWS Services | 9-Flask & Labs | '
    '10-Past Paper Patterns | 11-Coding Reference Cards'
)
closing.alignment = WD_ALIGN_PARAGRAPH.CENTER
closing.runs[0].italic = True

out = r"C:\Users\User\Downloads\4442\exam_output\exam_notes.docx"
doc.save(out)
print(f"Saved {out}")
