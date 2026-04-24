"""
Generate glossary.docx — COMP4442 full glossary.
All acronyms, abbreviations, and key technical terms from lectures 01-09.
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()
for section in doc.sections:
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)

style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(10)


def H(text, lvl=1):
    return doc.add_heading(text, level=lvl)


def table2(headers, rows, col_widths=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Light Grid Accent 1'
    hdr = t.rows[0]
    for i, h in enumerate(headers):
        hdr.cells[i].text = h
        for run in hdr.cells[i].paragraphs[0].runs:
            run.bold = True
            run.font.size = Pt(10)
    for ri, row in enumerate(rows, 1):
        for ci, val in enumerate(row):
            cell = t.rows[ri].cells[ci]
            cell.text = val
            for run in cell.paragraphs[0].runs:
                run.font.size = Pt(10)
    if col_widths:
        for row in t.rows:
            for ci, w in enumerate(col_widths):
                row.cells[ci].width = Cm(w)
    doc.add_paragraph()


# ── TITLE ──────────────────────────────────────────────────────────────────
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('COMP4442 — Complete Glossary & Abbreviations')
run.bold = True
run.font.size = Pt(14)
doc.add_paragraph('All acronyms, short forms, and key technical terms. Covers Cloud, Virtualization, K8s, Networking, Hadoop, Spark, SOA, AWS, Flask.')

# ══════════════════════════════════════════════════════════════════════════
H('1. Acronyms A–Z')
# ══════════════════════════════════════════════════════════════════════════
table2(
    ['Acronym', 'Stands For', 'Context / Meaning'],
    [
        # A
        ('ACL',     'Access Control List',                    'List of permissions for a resource'),
        ('ALS',     'Alternating Least Squares',              'Spark MLlib collaborative filtering algorithm'),
        ('AMI',     'Amazon Machine Image',                   'EC2 instance template (OS + software)'),
        ('API',     'Application Programming Interface',      'Contract for how software components communicate'),
        ('AUC',     'Area Under Curve',                       'ML metric — area under ROC curve, higher = better'),
        # B
        ('BLOB',    'Binary Large Object',                    'Unstructured binary data stored in S3/Azure Blob'),
        # C
        ('CDN',     'Content Delivery Network',               'Edge servers cache content close to users (CloudFront)'),
        ('CF',      'Collaborative Filtering',                'Spark MLlib recommender technique'),
        ('CPU',     'Central Processing Unit',                'Core compute resource; VMs share physical CPUs'),
        ('CRUD',    'Create Read Update Delete',              'Four basic database/REST operations'),
        # D
        ('DAG',     'Directed Acyclic Graph',                 'Spark job execution plan — nodes=RDDs, edges=transforms; no cycles'),
        ('DLQ',     'Dead Letter Queue',                      'SQS queue for messages that failed processing N times'),
        ('DNS',     'Domain Name System',                     'Resolves hostnames to IPs (AWS: Route 53)'),
        ('DVM',     'Domain-0 VM',                            'Xen privileged VM that controls hardware access'),
        # E
        ('EB',      'Elastic Beanstalk',                      'AWS PaaS — auto-deploys Flask/Java apps on EC2+ALB'),
        ('ECMP',    'Equal-Cost Multi-Path',                  'Fat-Tree routing — spread traffic across equal-cost paths'),
        ('ECS',     'Elastic Container Service',              'AWS managed Docker container orchestration'),
        ('EKS',     'Elastic Kubernetes Service',             'AWS managed Kubernetes control plane'),
        ('EMR',     'Elastic MapReduce',                      'AWS managed Hadoop/Spark cluster service'),
        ('ESB',     'Enterprise Service Bus',                 'SOA middleware: routing, transformation, service directory'),
        # F
        ('FIFO',    'First In First Out',                     'SQS queue type: strict order, exactly-once, 300 TPS max'),
        ('FaaS',    'Function as a Service',                  'Serverless compute — AWS Lambda; billed per invocation'),
        # G
        ('GBT',     'Gradient Boosted Trees',                 'Spark MLlib ensemble classifier/regressor'),
        ('GFS',     'Google File System',                     'Google\'s distributed FS; Hadoop\'s HDFS is based on it'),
        ('GSI',     'Global Secondary Index',                 'DynamoDB alternate query key — always use GSI, never Scan'),
        # H
        ('HDFS',    'Hadoop Distributed File System',         'Block-replicated (3×) distributed storage; NameNode + DataNodes'),
        ('HTTP',    'HyperText Transfer Protocol',            'Application-layer protocol for web/REST APIs'),
        ('HTTPS',   'HTTP Secure',                            'HTTP over TLS — encrypted; required for prod APIs'),
        # I
        ('IaaS',    'Infrastructure as a Service',            'Rent VMs/storage/network — EC2, S3 (NIST model)'),
        ('IAM',     'Identity and Access Management',         'AWS: users, roles, policies; least-privilege principle'),
        ('IoT',     'Internet of Things',                     'Networked physical devices; AWS IoT Core = MQTT broker'),
        # J
        ('JSON',    'JavaScript Object Notation',             'Lightweight data format used in REST APIs, DynamoDB items'),
        ('JWT',     'JSON Web Token',                         'Signed token for stateless auth; Cognito issues JWTs'),
        # K
        ('K8s',     'Kubernetes',                             'Container orchestration; \'8\' = 8 letters between K and s'),
        # L
        ('LSI',     'Local Secondary Index',                  'DynamoDB index sharing partition key, alternate sort key'),
        ('LXC',     'Linux Containers',                       'OS-level virtualisation using cgroups + namespaces'),
        # M
        ('MLlib',   'Machine Learning Library',               'Spark\'s built-in ML: classification, regression, CF, clustering'),
        ('MQTT',    'Message Queuing Telemetry Transport',    'Lightweight pub/sub protocol for IoT devices'),
        ('MR',      'MapReduce',                              'Google/Hadoop batch processing paradigm: map → shuffle → reduce'),
        # N
        ('NAT',     'Network Address Translation',            'Maps private IPs to public — Lambda in VPC uses NAT GW'),
        ('NIST',    'National Institute of Standards & Technology', 'Defines cloud: 5 characteristics, 3 service models, 4 deploy models'),
        ('NLP',     'Natural Language Processing',            'ML subfield; Spark MLlib includes basic NLP features'),
        # O
        ('OLAP',    'Online Analytical Processing',           'Complex queries on large historical data (Spark, Hive)'),
        ('OLTP',    'Online Transaction Processing',          'Fast short transactions — DynamoDB, RDS'),
        ('ORM',     'Object-Relational Mapping',              'Maps Python classes to DB tables (SQLAlchemy for Flask)'),
        # P
        ('PaaS',    'Platform as a Service',                  'Managed runtime — Elastic Beanstalk, Heroku (NIST model)'),
        # R
        ('RAM',     'Random Access Memory',                   'In-memory storage; Spark caches RDDs in RAM for speed'),
        ('RDD',     'Resilient Distributed Dataset',          'Spark\'s core abstraction: immutable, partitioned, fault-tolerant collection'),
        ('REST',    'Representational State Transfer',        'Stateless HTTP API style: GET/POST/PUT/DELETE on resources'),
        ('RMSE',    'Root Mean Square Error',                 'Regression ML metric: sqrt(mean(predicted-actual)²); lower=better'),
        ('ROC',     'Receiver Operating Characteristic',      'ML curve: TPR vs FPR at various thresholds; AUC summarises it'),
        # S
        ('S3',      'Simple Storage Service',                 'AWS object store: bucket/key, 11 9s durability, unlimited scale'),
        ('SaaS',    'Software as a Service',                  'Full app delivered over internet — Gmail, Salesforce (NIST)'),
        ('SDN',     'Software-Defined Networking',            'Separates control plane (controller) from data plane (switch)'),
        ('SLA',     'Service Level Agreement',                'Contract: uptime%, latency, throughput guarantees'),
        ('SNS',     'Simple Notification Service',            'AWS pub/sub fan-out: one message → many subscribers (SQS/Lambda/email)'),
        ('SOA',     'Service-Oriented Architecture',          'Design: business functions as loosely-coupled interoperable services'),
        ('SOAP',    'Simple Object Access Protocol',          'XML-based web service protocol; heavier than REST'),
        ('SQS',     'Simple Queue Service',                   'AWS managed message queue; Standard (unlimited, ≥1) or FIFO (300 TPS, =1)'),
        ('SSH',     'Secure Shell',                           'Encrypted remote login protocol; used to connect to EC2'),
        ('SVD',     'Singular Value Decomposition',           'Matrix factorisation; MLlib uses it for dimensionality reduction'),
        # T
        ('TCP',     'Transmission Control Protocol',          'Reliable ordered transport; HTTP/HTTPS run over TCP'),
        ('TLS',     'Transport Layer Security',               'Encryption protocol for HTTPS; replaces SSL'),
        ('ToR',     'Top of Rack',                            'Network switch at top of server rack in Fat-Tree topology'),
        ('TTL',     'Time to Live',                           'DynamoDB: auto-delete expired items; also DNS record lifespan'),
        # U
        ('URI',     'Uniform Resource Identifier',            'General identifier for a resource (URL is a type of URI)'),
        ('URL',     'Uniform Resource Locator',               'Web address; in REST = resource path e.g. /api/flights/{id}'),
        # V
        ('VM',      'Virtual Machine',                        'Isolated OS instance on shared hardware via hypervisor'),
        ('VMM',     'Virtual Machine Monitor',                'Hypervisor; Type-1 runs on hardware, Type-2 runs on host OS'),
        ('VNI',     'VXLAN Network Identifier',               '24-bit ID in VXLAN header; supports 16M virtual networks'),
        ('VPC',     'Virtual Private Cloud',                  'Isolated AWS network; Lambdas and RDS placed inside VPC'),
        ('VXLAN',   'Virtual Extensible LAN',                 'Overlay network: encapsulates L2 frames in UDP — extends VLAN to 16M'),
        # W
        ('WSDL',    'Web Services Description Language',      'XML contract describing a SOAP service\'s operations'),
        # X
        ('XML',     'Extensible Markup Language',             'Document format used in SOAP; JSON has largely replaced it for REST'),
        # Y
        ('YARN',    'Yet Another Resource Negotiator',        'Hadoop cluster resource manager: ResourceManager + NodeManagers'),
    ],
    col_widths=[2.0, 5.5, 9.5]
)

# ══════════════════════════════════════════════════════════════════════════
H('2. Spark: Transformations vs Actions')
# ══════════════════════════════════════════════════════════════════════════
doc.add_paragraph(
    'Transformation — lazy; returns a new RDD/DataFrame; nothing runs until an Action is called.\n'
    'Action — triggers the DAG execution; returns a result to the driver or writes to storage.'
)

H('2.1 RDD Transformations (lazy)', lvl=2)
table2(
    ['Function', 'Returns', 'What it does'],
    [
        ('map(f)',              'RDD',  'Apply f to every element; 1-to-1'),
        ('flatMap(f)',          'RDD',  'Apply f, then flatten nested iterables; 1-to-many'),
        ('filter(f)',           'RDD',  'Keep elements where f returns True'),
        ('distinct()',          'RDD',  'Remove duplicate elements (triggers shuffle)'),
        ('union(other)',        'RDD',  'Concatenate two RDDs (no dedup)'),
        ('intersection(other)', 'RDD',  'Keep elements in both RDDs (shuffle)'),
        ('subtract(other)',     'RDD',  'Elements in this RDD not in other'),
        ('cartesian(other)',    'RDD',  'All (a,b) pairs — expensive O(m×n)'),
        ('sample(replace,frac)','RDD',  'Random fraction of elements'),
        ('repartition(n)',      'RDD',  'Shuffle to exactly n partitions (expensive)'),
        ('coalesce(n)',         'RDD',  'Reduce to n partitions without full shuffle'),
        ('mapPartitions(f)',    'RDD',  'Like map but f receives whole partition iterator'),
        ('zip(other)',          'RDD',  'Pair elements by index → (a, b) tuples'),
        # pair RDDs
        ('groupByKey()',        'RDD[(K,Iter[V])]', 'Group all values for each key — high memory'),
        ('reduceByKey(f)',      'RDD[(K,V)]',       'Merge values per key with f — preferred over groupByKey'),
        ('aggregateByKey(z,f,g)','RDD[(K,V)]',      'Per-key aggregation with different combine/merge functions'),
        ('combineByKey(…)',     'RDD[(K,C)]',       'General per-key combiner — lowest-level, flexible'),
        ('sortByKey(asc)',      'RDD[(K,V)]',       'Sort pair RDD by key'),
        ('keys()',              'RDD[K]',           'Extract keys'),
        ('values()',            'RDD[V]',           'Extract values'),
        ('join(other)',         'RDD[(K,(V,W))]',   'Inner join on key'),
        ('leftOuterJoin(other)','RDD[(K,(V,Opt[W]))]', 'Left outer join; missing right = None'),
        ('rightOuterJoin(other)','RDD',             'Right outer join'),
        ('cogroup(other)',      'RDD[(K,(Iter,Iter))]', 'Group both RDDs by key simultaneously'),
        ('partitionBy(n,func)', 'RDD',              'Hash-partition by key with custom partitioner'),
    ],
    col_widths=[4.5, 4.0, 8.5]
)

H('2.2 RDD Actions (trigger execution)', lvl=2)
table2(
    ['Function', 'Returns', 'What it does'],
    [
        ('collect()',         'list',   'Fetch all elements to driver — careful with large RDDs'),
        ('count()',           'int',    'Count total elements'),
        ('first()',           'elem',   'Return first element'),
        ('take(n)',           'list',   'Return first n elements'),
        ('takeSample(rep,n)', 'list',   'Random sample of n elements'),
        ('top(n)',            'list',   'Top n elements (uses natural order)'),
        ('reduce(f)',         'value',  'Aggregate all elements with associative/commutative f'),
        ('fold(zero, f)',     'value',  'Like reduce but with a zero/identity value'),
        ('aggregate(z,f,g)',  'value',  'Reduce within partitions (f), then merge partitions (g)'),
        ('foreach(f)',        'None',   'Apply f for side effects (e.g. write to DB)'),
        ('foreachPartition(f)','None',  'Like foreach but f receives whole partition'),
        ('countByKey()',      'dict',   'Count occurrences of each key → {key: count}'),
        ('countByValue()',    'dict',   'Count occurrences of each value'),
        ('collectAsMap()',    'dict',   'Return pair RDD as Python dict'),
        ('lookup(key)',       'list',   'Return all values for given key'),
        ('saveAsTextFile(p)', 'None',   'Write one text file per partition to path p'),
        ('saveAsPickleFile(p)','None',  'Write binary pickle files'),
    ],
    col_widths=[4.5, 3.0, 9.5]
)

H('2.3 DataFrame / SparkSQL Operations', lvl=2)
table2(
    ['Operation', 'Type', 'Example / Notes'],
    [
        # transformations
        ('select(cols…)',       'Transform', 'df.select("name", "age")  — pick columns'),
        ('filter(cond) / where','Transform', 'df.filter(df.age > 30)'),
        ('withColumn(name,expr)','Transform','df.withColumn("tax", df.salary * 0.2)  — add/replace col'),
        ('withColumnRenamed(o,n)','Transform','Rename a column'),
        ('drop(col)',           'Transform', 'Remove a column'),
        ('groupBy(cols…)',      'Transform', 'Returns GroupedData — must follow with agg()'),
        ('agg(exprs…)',         'Transform', '.agg(sum("sales"), avg("qty"))'),
        ('orderBy / sort',      'Transform', 'df.orderBy("price", ascending=False)'),
        ('join(other, on, how)','Transform', 'how = inner|left|right|outer|cross|semi|anti'),
        ('union(other)',        'Transform', 'Stack two DataFrames with same schema'),
        ('distinct()',          'Transform', 'Remove duplicate rows'),
        ('dropDuplicates(cols)','Transform', 'Dedup on specific columns'),
        ('limit(n)',            'Transform', 'Keep first n rows'),
        ('explode(col)',        'Transform', 'One row per array element (like RDD flatMap)'),
        ('pivot(col,vals,agg)', 'Transform', 'Pivot rows to columns'),
        ('cache() / persist()', 'Transform', 'Cache in memory for reuse — still lazy'),
        # actions
        ('show(n)',             'Action',    'Print first n rows (default 20) to console'),
        ('collect()',           'Action',    'Return all rows as list of Row objects to driver'),
        ('count()',             'Action',    'Count rows'),
        ('first() / head()',    'Action',    'Return first Row'),
        ('take(n)',             'Action',    'Return first n rows as list'),
        ('toPandas()',          'Action',    'Convert to Pandas DataFrame — driver must have enough RAM'),
        ('write.parquet(path)', 'Action',    'Save as Parquet files to path'),
        ('write.csv(path)',     'Action',    'Save as CSV files'),
        ('write.mode(m)',       'Action',    'm = overwrite|append|ignore|errorIfExists'),
        ('createOrReplaceTempView(name)','Action','Register as SQL temp table for spark.sql(…)'),
    ],
    col_widths=[4.5, 2.5, 10.0]
)

# ══════════════════════════════════════════════════════════════════════════
H('3. Spark Core Concepts')
# ══════════════════════════════════════════════════════════════════════════
table2(
    ['Term', 'Definition'],
    [
        ('RDD\n(Resilient Distributed Dataset)',
         'Core Spark abstraction. Immutable, partitioned collection of records distributed across a cluster. '
         'Fault-tolerant via lineage (recompute lost partitions). Created from HDFS/S3, parallelise(), or by transforming existing RDDs.'),
        ('DAG\n(Directed Acyclic Graph)',
         'Spark\'s execution plan. Each node is an RDD; each edge is a transformation. '
         'Submitted to the DAGScheduler which splits it into stages at shuffle boundaries.'),
        ('Lazy Evaluation',
         'Transformations are NOT executed when called — they build the DAG. '
         'Execution only starts when an Action is called. Allows Spark to optimise the full plan before running.'),
        ('Lineage',
         'The chain of transformations that produced an RDD. '
         'Used for fault recovery: if a partition is lost, Spark re-executes only the lineage steps for that partition.'),
        ('Transformation',
         'Lazy operation on an RDD/DataFrame that returns a new RDD/DataFrame. '
         'Nothing executes. Examples: map, filter, groupByKey, join.'),
        ('Action',
         'Triggers DAG execution and returns a value to the driver or writes to storage. '
         'Examples: collect, count, reduce, saveAsTextFile.'),
        ('Stage',
         'Set of transformations that can run without a shuffle. '
         'DAGScheduler splits the DAG into stages at wide dependency boundaries.'),
        ('Narrow Dependency',
         'Each parent partition feeds at most one child partition. '
         'No shuffle needed. Examples: map, filter, union. Fast.'),
        ('Wide Dependency\n(Shuffle Dependency)',
         'A child partition needs data from multiple parent partitions. '
         'Requires a shuffle (network transfer). Stage boundary. Examples: groupByKey, reduceByKey, join.'),
        ('Shuffle',
         'Redistribution of data across the cluster by key. '
         'Expensive: disk I/O + serialisation + network. Marks a stage boundary. Avoid unnecessary shuffles.'),
        ('Partition',
         'Unit of parallelism in Spark. One task per partition. '
         'Default = number of HDFS blocks. Use repartition(n) / coalesce(n) to adjust.'),
        ('Task',
         'A unit of work sent to one executor: processes one partition for one stage.'),
        ('Executor',
         'JVM process on a worker node. Runs tasks and caches RDD data. '
         'Multiple executors per cluster; configured with --executor-memory, --executor-cores.'),
        ('Driver',
         'The process running the main() function. Hosts SparkContext, builds the DAG, schedules tasks. '
         'collect() sends data back to the driver — can OOM if RDD is large.'),
        ('SparkContext (sc)',
         'Entry point for RDD API. Created once per application. '
         'sc.textFile(), sc.parallelize(), sc.broadcast(), sc.accumulator().'),
        ('SparkSession (spark)',
         'Unified entry point (Spark 2.0+). Encompasses SparkContext + SQLContext + HiveContext. '
         'Use for DataFrame/Dataset/SQL API.'),
        ('Broadcast Variable',
         'Read-only variable sent efficiently to all executors once (not re-sent per task). '
         'Use for large lookup tables: b = sc.broadcast(lookup); b.value.'),
        ('Accumulator',
         'Write-only shared counter/aggregator; only the driver can read. '
         'Typically used for counters: acc = sc.accumulator(0); rdd.foreach(lambda x: acc.add(1)).'),
        ('Persistence / Cache',
         'rdd.cache() = persist(MEMORY_ONLY). Keeps RDD in RAM across actions. '
         'Levels: MEMORY_ONLY, MEMORY_AND_DISK, DISK_ONLY, _SER (serialised), _2 (replicated).'),
        ('DataFrame',
         'Distributed table with named, typed columns. Built on RDD of Row objects. '
         'Optimised by Catalyst query optimiser and Tungsten execution engine. Faster than raw RDD for structured data.'),
        ('Catalyst Optimiser',
         'Spark SQL\'s query optimiser. Rewrites the logical plan (predicate pushdown, column pruning, join reordering) '
         'before generating physical execution plan.'),
        ('Tungsten',
         'Spark\'s memory and CPU optimisation engine. Off-heap memory management, code generation, cache-aware computation.'),
        ('Pair RDD',
         'RDD of (key, value) tuples. Enables key-based operations: groupByKey, reduceByKey, join, etc.'),
        ('reduceByKey vs groupByKey',
         'Prefer reduceByKey: combines values per key on each partition before shuffle → less data transferred. '
         'groupByKey sends all values across network then aggregates — much slower.'),
    ],
    col_widths=[4.5, 12.5]
)

# ══════════════════════════════════════════════════════════════════════════
H('4. Hadoop / HDFS / YARN Concepts')
# ══════════════════════════════════════════════════════════════════════════
table2(
    ['Term', 'Definition'],
    [
        ('HDFS',      'Hadoop Distributed File System. Splits files into blocks (default 128MB), replicates each 3× across DataNodes.'),
        ('NameNode',  'HDFS master. Stores filesystem metadata (file→block→DataNode mapping). Single point of failure — use HA pair.'),
        ('DataNode',  'HDFS worker. Stores actual data blocks. Sends heartbeat + block report to NameNode every 3s.'),
        ('Block',     'HDFS unit of storage. Default 128MB (was 64MB in older versions). Large blocks reduce metadata overhead.'),
        ('Replication Factor', 'Default 3. One block on local rack, two on different rack for fault tolerance.'),
        ('YARN',      'Yet Another Resource Negotiator. Hadoop 2 cluster resource manager. Decouples resource management from MR.'),
        ('ResourceManager', 'YARN master. Allocates cluster resources. Has Scheduler + ApplicationManager.'),
        ('NodeManager', 'YARN per-node agent. Manages containers, reports resources to ResourceManager.'),
        ('Container', 'YARN unit of resource allocation: fixed CPU + memory. Tasks run inside containers.'),
        ('ApplicationMaster', 'Per-job process negotiating containers from ResourceManager and tracking task progress.'),
        ('MapReduce',  'Batch processing paradigm: Map (emit k,v pairs) → Shuffle/Sort (group by key) → Reduce (aggregate per key).'),
        ('Mapper',    'Reads input splits, applies user function, emits (key, value) pairs.'),
        ('Reducer',   'Receives sorted (key, list[values]) per key, applies user aggregation, writes output.'),
        ('Combiner',  'Optional mini-reducer run on Map output before shuffle. Reduces network I/O. Must be associative+commutative.'),
        ('Partitioner', 'Decides which Reducer gets each (key,value). Default: hashCode(key) % numReducers.'),
        ('InputSplit', 'Logical chunk of input data assigned to one Mapper. Usually equals one HDFS block.'),
        ('Job Tracker / Task Tracker', 'Hadoop 1 names for what YARN ResourceManager / NodeManager replaced.'),
        ('Hive',      'SQL-on-Hadoop. Translates HiveQL to MapReduce/Tez/Spark jobs. Schema-on-read.'),
        ('HBase',     'Bigtable-inspired NoSQL on HDFS. Low-latency random reads/writes. Column-family model.'),
        ('Pig',       'Data flow scripting language. Pig Latin scripts compile to MapReduce.'),
        ('Sqoop',     'SQL-to-Hadoop. Bulk import/export between RDBMS and HDFS.'),
        ('Flume',     'Log/event ingestion agent. Collects, aggregates, and moves streaming data into HDFS.'),
        ('ZooKeeper', 'Distributed coordination service. Leader election, distributed locks. Hadoop equivalent of Google\'s Chubby.'),
        ('Tez',       'DAG execution engine replacing MapReduce under Hive/Pig. Avoids MR\'s write-to-HDFS-between-steps overhead.'),
    ],
    col_widths=[4.0, 13.0]
)

# ══════════════════════════════════════════════════════════════════════════
H('5. Kubernetes (K8s) Terms')
# ══════════════════════════════════════════════════════════════════════════
table2(
    ['Term', 'Definition'],
    [
        ('Pod',          'Smallest deployable unit. One or more containers sharing network + storage. Ephemeral.'),
        ('Node',         'A worker machine (VM or physical) running Pods. Has kubelet + kube-proxy.'),
        ('Cluster',      'Set of Nodes managed by a control plane.'),
        ('Control Plane','Master components: API Server, etcd, Scheduler, Controller Manager.'),
        ('API Server',   'Front-end for K8s control plane. All kubectl commands go here via REST.'),
        ('etcd',         'Distributed key-value store. K8s\'s source of truth for all cluster state.'),
        ('Scheduler',    'Watches for unscheduled Pods and assigns them to Nodes based on resource requirements.'),
        ('Controller Manager', 'Runs control loops: ReplicaSet controller, Deployment controller, etc.'),
        ('kubelet',      'Agent on each Node. Ensures containers described in PodSpecs are running and healthy.'),
        ('kube-proxy',   'Network proxy on each Node. Maintains network rules for Service routing.'),
        ('Deployment',   'Declarative update manager for Pods. Manages rollout, rollback, scaling.'),
        ('ReplicaSet',   'Ensures N copies of a Pod are running. Deployment creates/manages ReplicaSets.'),
        ('Service',      'Stable virtual IP + DNS name for a set of Pods. Types: ClusterIP, NodePort, LoadBalancer.'),
        ('ConfigMap',    'Non-secret configuration data injected into Pods as env vars or files.'),
        ('Secret',       'Like ConfigMap but base64-encoded. For passwords, tokens, keys.'),
        ('Namespace',    'Virtual cluster within a cluster. Isolates resources for different teams/envs.'),
        ('PersistentVolume (PV)', 'Cluster-level storage resource. Decouples storage lifecycle from Pod lifecycle.'),
        ('Ingress',      'HTTP/HTTPS routing rules into the cluster. Layer-7 load balancing.'),
        ('DaemonSet',    'Ensures one Pod per Node (e.g. log collector, monitoring agent).'),
        ('StatefulSet',  'Like Deployment but for stateful apps — stable Pod identity + persistent storage.'),
        ('Horizontal Pod Autoscaler (HPA)', 'Automatically scales Deployment replica count based on CPU/memory metrics.'),
        ('kubectl',      'K8s CLI. kubectl apply -f, get pods, describe pod, logs, exec, delete.'),
        ('YAML',         'Config language for K8s manifests. Key-value, indent-sensitive. Defines all K8s objects.'),
    ],
    col_widths=[4.5, 12.5]
)

# ══════════════════════════════════════════════════════════════════════════
H('6. Virtualization Terms')
# ══════════════════════════════════════════════════════════════════════════
table2(
    ['Term', 'Definition'],
    [
        ('Hypervisor / VMM', 'Software layer that creates and runs VMs. Type-1 = bare-metal (Xen, ESXi, Hyper-V). Type-2 = hosted (VirtualBox, VMware Workstation).'),
        ('Type-1 Hypervisor', 'Runs directly on hardware. No host OS. Lower overhead. Used in cloud (Xen on AWS, ESXi).'),
        ('Type-2 Hypervisor', 'Runs on top of a host OS. Higher overhead. Used for dev/test (VirtualBox).'),
        ('Guest OS',         'OS running inside a VM.'),
        ('Host OS',          'OS on the physical machine running a Type-2 hypervisor.'),
        ('Xen',              'Open-source Type-1 hypervisor. AWS used Xen historically. Has Dom0 (privileged) + DomU (guest) VMs.'),
        ('Dom0',             'Xen Domain 0. Privileged VM with direct hardware access. Runs device drivers; other VMs go through it.'),
        ('DomU',             'Xen unprivileged guest VM. Cannot access hardware directly — requests go via Dom0.'),
        ('Para-virtualisation', 'Guest OS is modified to make hypercalls instead of privileged instructions. Faster than full virtualisation.'),
        ('Full Virtualisation', 'Unmodified guest OS. Hypervisor traps & emulates privileged instructions. Slower but compatible.'),
        ('Hardware-assisted Virtualisation', 'CPU features (Intel VT-x, AMD-V) allow unmodified guests to run efficiently. Best of both.'),
        ('cgroups',          'Linux Control Groups. Limit and account for CPU, memory, I/O per process group. Used by Docker/K8s.'),
        ('namespaces',       'Linux kernel feature. Isolates process view of: PID, network, mount, UTS, IPC, user. Foundation of containers.'),
        ('Container',        'Lightweight process isolation using namespaces + cgroups. Shares host kernel. Faster than VMs; less isolated.'),
        ('Docker',           'Container runtime. Dockerfile → image → container. Uses namespaces + cgroups. Docker Hub = image registry.'),
        ('Docker Image',     'Read-only template: layers of filesystem changes. Built from Dockerfile. Stored in registry.'),
        ('Docker Container', 'Running instance of an image. Ephemeral by default.'),
        ('Dockerfile',       'Script of instructions to build a Docker image: FROM, RUN, COPY, EXPOSE, CMD.'),
        ('Docker Compose',   'Tool to define and run multi-container apps with a YAML file.'),
        ('VM vs Container',  'VM: full OS, minutes to start, strong isolation, GB image. Container: shared kernel, seconds to start, weaker isolation, MB image.'),
    ],
    col_widths=[4.5, 12.5]
)

# ══════════════════════════════════════════════════════════════════════════
H('7. Networking Terms')
# ══════════════════════════════════════════════════════════════════════════
table2(
    ['Term', 'Definition'],
    [
        ('Fat-Tree Topology', 'Data-center network using k-port switches. 3 layers: Core, Aggregation, Edge. Provides k³/4 servers, k²/4 redundant paths.'),
        ('k (Fat-Tree)',     'Port count of each switch. Determines: (k/2)² core switches, k aggregation pods, k/2 edge switches per pod.'),
        ('Over-subscription','Ratio of potential to actual bandwidth. Fat-Tree achieves 1:1 (non-blocking) at full bisection.'),
        ('DCell',            'Alternative data-center topology. Recursive structure: DCell₀ = n servers + 1 switch. Higher DCell levels link lower DCells.'),
        ('SDN',              'Software-Defined Networking. Separates control plane (where to send) from data plane (how to forward). Controller programs switches via OpenFlow.'),
        ('Control Plane',    'Network intelligence: routing decisions, topology discovery. In SDN: centralised in SDN Controller.'),
        ('Data Plane',       'Actual packet forwarding per the flow table. In SDN: distributed across physical switches.'),
        ('OpenFlow',         'Protocol between SDN controller and switches. Controller installs flow rules; switch matches packets and acts (forward/drop/modify).'),
        ('Flow Table',       'Table in an OpenFlow switch. Each entry: match fields + actions. Controller installs entries.'),
        ('VXLAN',            'Virtual Extensible LAN. Encapsulates Ethernet frames (L2) in UDP packets (L4). 24-bit VNI → 16M virtual networks. Used in cloud overlays.'),
        ('VNI',              'VXLAN Network Identifier. 24-bit tag in VXLAN header. Identifies which virtual network a packet belongs to.'),
        ('VTEP',             'VXLAN Tunnel Endpoint. Encapsulates/decapsulates VXLAN packets. Can be a hypervisor, switch, or router.'),
        ('Overlay Network',  'Virtual network built on top of physical network. VXLAN is an overlay. Decouples VM addressing from physical topology.'),
        ('Underlay Network', 'Physical IP network that carries overlay traffic. VXLAN tunnels run through the underlay.'),
        ('BGP',              'Border Gateway Protocol. Exterior routing protocol. Used between data-centers and ISPs.'),
        ('ECMP',             'Equal-Cost Multi-Path. Distribute traffic across multiple equal-cost links. Used in Fat-Tree to avoid bottlenecks.'),
        ('ARP',              'Address Resolution Protocol. Maps IP address to MAC address on a LAN.'),
        ('VLAN',             'Virtual LAN. L2 network segmentation. Limited to 4094 IDs — VXLAN extends this to 16M.'),
    ],
    col_widths=[4.0, 13.0]
)

# ══════════════════════════════════════════════════════════════════════════
H('8. AWS Services Reference')
# ══════════════════════════════════════════════════════════════════════════
table2(
    ['Service', 'Full Name / Category', 'One-line Description'],
    [
        ('EC2',            'Elastic Compute Cloud — IaaS',     'Rentable virtual machines; choose instance type, OS, storage'),
        ('S3',             'Simple Storage Service — Object',  'Unlimited object store; bucket+key, 11 9s durability, static website hosting'),
        ('Lambda',         'Lambda — FaaS/Serverless',         'Run code without servers; triggered by events; billed per 100ms'),
        ('DynamoDB',       'DynamoDB — NoSQL DB',              'Fully managed key-value + document DB; single-digit ms; auto-scale'),
        ('RDS',            'Relational Database Service',      'Managed MySQL/PostgreSQL/Aurora; handles patching, backups, failover'),
        ('API Gateway',    'API Gateway — API Management',     'Create/manage REST & WebSocket APIs; JWT auth, throttling, caching'),
        ('Cognito',        'Cognito — Auth',                   'User pools (auth) + identity pools (AWS creds); issues JWTs'),
        ('SQS',            'Simple Queue Service — Messaging', 'Managed message queue; Standard (unlimited,≥1) or FIFO (300TPS,=1)'),
        ('SNS',            'Simple Notification Service — Pub/Sub', 'Fan-out: one publish → many subscribers (SQS, Lambda, email, SMS)'),
        ('Kinesis',        'Kinesis — Streaming',              'Real-time data streaming; shards; Kinesis Analytics for SQL on stream'),
        ('EMR',            'Elastic MapReduce — Big Data',     'Managed Hadoop/Spark cluster; use --auto-terminate; spot workers'),
        ('Glue',           'Glue — ETL',                       'Serverless ETL service; Glue Catalog = central metadata store'),
        ('Athena',         'Athena — Query',                   'Serverless SQL on S3 data; pay per query; uses Glue Catalog'),
        ('CloudWatch',     'CloudWatch — Monitoring',          'Metrics, logs, alarms, dashboards; trigger Lambda on alarm'),
        ('EventBridge',    'EventBridge — Event Bus',          'Serverless event bus; cron schedules; route events between services'),
        ('Step Functions', 'Step Functions — Workflow',        'Serverless state machine orchestrating Lambda functions'),
        ('IAM',            'Identity and Access Management',   'Users, groups, roles, policies; principle of least privilege'),
        ('VPC',            'Virtual Private Cloud — Network',  'Isolated virtual network; subnets, route tables, security groups, NAT GW'),
        ('CloudFront',     'CloudFront — CDN',                 'Global CDN with edge locations; cache S3/API responses close to users'),
        ('Route 53',       'Route 53 — DNS',                   'Managed DNS; routing policies: simple, weighted, latency, failover, geolocation'),
        ('ECS',            'Elastic Container Service',        'Managed Docker orchestration; Fargate = serverless containers'),
        ('EKS',            'Elastic Kubernetes Service',       'Managed K8s control plane; worker nodes on EC2 or Fargate'),
        ('Elastic Beanstalk','Elastic Beanstalk — PaaS',       'Upload app code; EB auto-provisions EC2+ALB+ASG; supports Flask/Java/Node'),
        ('Secrets Manager','Secrets Manager — Secrets',        'Store and rotate API keys/DB passwords; fetch at Lambda cold start'),
        ('IoT Core',       'IoT Core — IoT',                   'Managed MQTT broker; rules engine routes device messages to other services'),
        ('Lightsail',      'Lightsail — VPS',                  'Simple VPS for small apps; fixed monthly price; easier than EC2'),
    ],
    col_widths=[3.5, 5.0, 8.5]
)

# ══════════════════════════════════════════════════════════════════════════
H('9. SOA / Microservices / REST Terms')
# ══════════════════════════════════════════════════════════════════════════
table2(
    ['Term', 'Definition'],
    [
        ('SOA',             'Service-Oriented Architecture. Design style: business functionality exposed as discrete, loosely-coupled, interoperable services.'),
        ('Microservices',   'Evolution of SOA: each service is independently deployable, has its own DB, communicates via lightweight APIs (REST/gRPC/events).'),
        ('Monolith',        'Single deployable unit. All components tightly coupled. Simple to develop initially; hard to scale/change independently.'),
        ('ESB',             'Enterprise Service Bus. SOA middleware for routing, protocol transformation, message enrichment, and service discovery. Hub-and-spoke pattern.'),
        ('Point-to-Point',  'Service composition where consumer knows provider\'s endpoint directly. Simple but creates tight coupling and spaghetti dependencies.'),
        ('Hub-and-Spoke',   'All services communicate through ESB. ESB handles routing + transformation. Decouples services; ESB can become bottleneck.'),
        ('Service Registry','Directory where services register themselves. Clients look up endpoints dynamically (like DNS for services). ESB includes one.'),
        ('REST',            'Representational State Transfer. HTTP-based API style. Stateless, uniform interface, resources identified by URIs, standard verbs (GET/POST/PUT/DELETE).'),
        ('Stateless',       'Server stores no client session state between requests. Each request is self-contained. REST is stateless. Enables horizontal scaling.'),
        ('Stateful',        'Server remembers client state (session, basket, login). Hard to scale — requests must reach the same server. Traditional web apps.'),
        ('SOAP',            'Simple Object Access Protocol. XML-based, heavier than REST, with WSDL contract. Used in legacy enterprise integrations.'),
        ('WSDL',            'Web Services Description Language. XML contract describing a SOAP service: operations, message formats, endpoint URL.'),
        ('Idempotent',      'Operation that produces same result if called once or multiple times. GET/PUT/DELETE are idempotent; POST is not.'),
        ('SLA',             'Service Level Agreement. Contractual commitments: uptime % (e.g. 99.9% = 8.7h downtime/yr), latency p99, error rate.'),
        ('Circuit Breaker', 'Microservices resilience pattern. Stops calling a failing downstream service; returns fallback. Prevents cascade failures.'),
        ('Service Mesh',    'Infrastructure layer for service-to-service communication. mTLS, load balancing, observability. Examples: Istio, Linkerd.'),
        ('API Gateway',     'Single entry point for all clients. Handles auth, rate limiting, routing to backend microservices, request/response transformation.'),
        ('Choreography',    'Services react to events independently — no central coordinator. Decoupled but hard to trace. SNS/SQS enables this.'),
        ('Orchestration',   'Central coordinator calls services in sequence. Easier to trace; coordinator is a potential bottleneck. Step Functions / ESB.'),
        ('CAP Theorem',     'Distributed system can guarantee at most 2 of: Consistency, Availability, Partition tolerance. NoSQL (DynamoDB) favours AP.'),
        ('BASE',            'Basically Available, Soft state, Eventually consistent. NoSQL design philosophy. Contrast with ACID (relational DBs).'),
    ],
    col_widths=[4.0, 13.0]
)

# ══════════════════════════════════════════════════════════════════════════
H('10. Flask & Web Development Terms')
# ══════════════════════════════════════════════════════════════════════════
table2(
    ['Term', 'Definition'],
    [
        ('Flask',          'Lightweight Python web framework. Routes map URLs to functions. WSGI-based. No built-in ORM.'),
        ('Route',          'URL pattern bound to a function via @app.route("/path", methods=["GET","POST"]).'),
        ('Blueprint',      'Flask component for organising routes into reusable modules. Registered on the app with app.register_blueprint().'),
        ('WSGI',           'Web Server Gateway Interface. Python standard for web server ↔ web app communication. Flask is WSGI.'),
        ('Jinja2',         'Flask\'s default HTML template engine. {{ variable }}, {% for %}, {% if %}. Auto-escapes HTML.'),
        ('request',        'Flask global: access current HTTP request data: request.args (query string), request.form (POST body), request.json.'),
        ('jsonify()',      'Flask helper: converts Python dict/list to JSON Response with Content-Type: application/json.'),
        ('boto3',          'AWS SDK for Python. Used in Flask/Lambda to call DynamoDB, S3, SQS, etc. boto3.resource() or boto3.client().'),
        ('ORM',            'Object-Relational Mapping. Maps Python classes to DB tables. SQLAlchemy is the common choice with Flask.'),
        ('SQLAlchemy',     'Python ORM. db.Model for table definitions. db.session.add/commit/query for operations.'),
        ('Elastic Beanstalk','AWS PaaS for Flask apps. Upload ZIP; EB manages EC2, ALB, Auto Scaling, RDS. application.py must define `application`.'),
        ('JWT',            'JSON Web Token. Header.Payload.Signature (base64). Cognito issues JWTs; API Gateway validates them.'),
        ('CORS',           'Cross-Origin Resource Sharing. Browser security policy. Flask-CORS extension adds Access-Control-Allow-Origin headers.'),
        ('Middleware',     'Code that runs between request receipt and route handler (or between handler and response). E.g. auth checks, logging.'),
        ('HTTP Status Codes', '200 OK, 201 Created, 204 No Content, 400 Bad Request, 401 Unauthorised, 403 Forbidden, 404 Not Found, 500 Server Error'),
    ],
    col_widths=[4.0, 13.0]
)

# ══════════════════════════════════════════════════════════════════════════
H('11. Cloud Computing Fundamentals')
# ══════════════════════════════════════════════════════════════════════════
H('NIST 5 Essential Characteristics', lvl=2)
table2(
    ['Characteristic', 'Meaning'],
    [
        ('On-demand self-service',   'Provision resources without human interaction with provider'),
        ('Broad network access',     'Accessible over network via standard mechanisms'),
        ('Resource pooling',         'Multi-tenant; provider pools resources for dynamic assignment'),
        ('Rapid elasticity',         'Scale out/in quickly; appears unlimited to user'),
        ('Measured service',         'Usage monitored, controlled, reported — pay-per-use'),
    ],
    col_widths=[5.5, 11.5]
)

H('Deployment Models', lvl=2)
table2(
    ['Model', 'Description'],
    [
        ('Public Cloud',   'Owned by provider (AWS/Azure/GCP). Multi-tenant. Lower cost. Less control.'),
        ('Private Cloud',  'Dedicated to one org. On-premises or hosted. Higher cost. More control.'),
        ('Hybrid Cloud',   'Mix of public + private. Data sovereignty + burst to public for scale.'),
        ('Community Cloud','Shared by orgs with common concerns (e.g. government, healthcare).'),
    ],
    col_widths=[3.5, 13.5]
)

# ══════════════════════════════════════════════════════════════════════════
H('12. MLlib / Machine Learning Terms')
# ══════════════════════════════════════════════════════════════════════════
table2(
    ['Term', 'Definition'],
    [
        ('Feature',          'Input variable to a ML model. Must be numeric. Categorical features need StringIndexer + OneHotEncoder.'),
        ('Label',            'Target variable (what the model predicts). For classification: integer class index. For regression: float.'),
        ('Pipeline',         'Spark ML Pipeline chains Transformers and Estimators. fit() trains; transform() applies. Reproducible workflow.'),
        ('Transformer',      'ML Pipeline stage: takes DataFrame, produces DataFrame. Examples: StringIndexer, VectorAssembler, trained model.'),
        ('Estimator',        'ML Pipeline stage that is fit to data: produces a Transformer. Examples: LogisticRegression, GBTClassifier.'),
        ('VectorAssembler',  'Transformer that combines multiple columns into a single "features" vector column required by MLlib algorithms.'),
        ('StringIndexer',    'Encodes string column to numeric index (0, 1, 2…). Must precede OneHotEncoder for categorical features.'),
        ('OneHotEncoder',    'Converts numeric index to binary vector (sparse). Prevents model treating categories as ordered.'),
        ('GBT',              'Gradient Boosted Trees. Ensemble of decision trees trained sequentially; each tree corrects errors of previous.'),
        ('Random Forest',    'Ensemble of independently trained decision trees. Aggregate by majority vote (classification) or average (regression).'),
        ('Logistic Regression', 'Classification algorithm. Models P(class) using sigmoid function. Binary or multinomial.'),
        ('Linear Regression', 'Regression algorithm. Fits linear relationship between features and continuous target.'),
        ('ALS',              'Alternating Least Squares. MLlib collaborative filtering. Factorises user-item matrix. Used for recommendations.'),
        ('K-Means',          'Clustering algorithm. Assigns points to k clusters by nearest centroid. MLlib: KMeans().setK(k).'),
        ('SVD',              'Singular Value Decomposition. Dimensionality reduction. A = U·Σ·Vᵀ. MLlib RowMatrix.computeSVD().'),
        ('Cross-Validation', 'Technique to evaluate model. k-Fold: split data k ways, train on k-1 folds, test on 1, repeat.'),
        ('Train/Test Split', 'Divide data into training set (fit model) and test set (evaluate). Always time-based for time-series data.'),
        ('Overfitting',      'Model learns training noise; poor generalisation. Signs: high train accuracy, low test accuracy.'),
        ('Hyperparameter',   'Parameter set before training (not learned). E.g. maxDepth, numTrees, stepSize. Tune with CrossValidator.'),
        ('F1 Score',         'Harmonic mean of precision and recall: 2·P·R/(P+R). Weighted F1 accounts for class imbalance.'),
        ('Precision',        'Of all predicted positives, fraction truly positive. TP/(TP+FP).'),
        ('Recall',           'Of all actual positives, fraction correctly predicted. TP/(TP+FN).'),
        ('RMSE',             'Root Mean Square Error. sqrt(mean((ŷ−y)²)). Lower = better. Common regression metric.'),
        ('AUC-ROC',          'Area under ROC curve. 0.5 = random, 1.0 = perfect. Model-wide classification quality metric.'),
    ],
    col_widths=[4.0, 13.0]
)

doc.save('glossary.docx')
print("glossary.docx written.")
