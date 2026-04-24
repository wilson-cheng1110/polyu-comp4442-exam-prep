"""
Generate flask_syntax.docx — Flask documentation in W3Schools style.
Each topic: Definition → Syntax → Parameter table → Example(s) → Notes.
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()
for sec in doc.sections:
    sec.top_margin    = Cm(1.8)
    sec.bottom_margin = Cm(1.8)
    sec.left_margin   = Cm(2.0)
    sec.right_margin  = Cm(2.0)

doc.styles['Normal'].font.name = 'Calibri'
doc.styles['Normal'].font.size = Pt(10.5)

PAGE_W = 17.0   # usable cm after margins

# ─── colour palette (W3Schools-ish) ───────────────────────────────────────
GREEN  = RGBColor(0x04, 0xAA, 0x6D)   # W3Schools green for headings
DKGRAY = RGBColor(0x28, 0x28, 0x28)   # near-black body text
BLUE   = RGBColor(0x00, 0x59, 0xB3)   # function / keyword colour
GRAY_FILL  = "F1F1F1"   # code block background
GREEN_FILL = "D4EDDA"   # example block background
YELL_FILL  = "FFF9C4"   # note block background


# ─── XML helpers ──────────────────────────────────────────────────────────
def _shade(p, fill):
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  fill)
    pPr.append(shd)


def _border_box(p, color="AAAAAA"):
    """Add a thin border on all four sides of a paragraph."""
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    for side in ('top', 'left', 'bottom', 'right'):
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'),   'single')
        el.set(qn('w:sz'),    '4')
        el.set(qn('w:space'), '4')
        el.set(qn('w:color'), color)
        pBdr.append(el)
    pPr.append(pBdr)


def _cell_shade(cell, fill):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  fill)
    tcPr.append(shd)


# ─── building blocks ──────────────────────────────────────────────────────
def topic(title):
    """Green H1 title like W3Schools page heading."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(title)
    run.bold        = True
    run.font.size   = Pt(20)
    run.font.color.rgb = GREEN
    # bottom rule
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot = OxmlElement('w:bottom')
    bot.set(qn('w:val'),   'single')
    bot.set(qn('w:sz'),    '6')
    bot.set(qn('w:space'), '1')
    bot.set(qn('w:color'), '04AA6D')
    pBdr.append(bot)
    pPr.append(pBdr)


def sub(title):
    """Dark bold sub-heading (H2 equivalent)."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(3)
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = DKGRAY


def body(text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(4)
    for r in p.runs:
        r.font.size = Pt(10.5)
        r.font.color.rgb = DKGRAY
    return p


def note(text):
    """Yellow-tinted note box."""
    p = doc.add_paragraph()
    _shade(p, YELL_FILL)
    _border_box(p, "F0C040")
    p.paragraph_format.left_indent  = Cm(0.3)
    p.paragraph_format.right_indent = Cm(0.3)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run("Note:  ")
    run.bold = True
    run.font.size = Pt(10)
    run2 = p.add_run(text)
    run2.font.size = Pt(10)


def syntax_block(text, fill=GRAY_FILL):
    """
    Render a multi-line code block with shaded background + border.
    Each line becomes its own paragraph so shading tiles correctly.
    """
    lines = text.split('\n')
    for i, line in enumerate(lines):
        p = doc.add_paragraph()
        _shade(p, fill)
        p.paragraph_format.left_indent  = Cm(0.5)
        p.paragraph_format.right_indent = Cm(0.3)
        p.paragraph_format.space_before = Pt(0) if i > 0 else Pt(4)
        p.paragraph_format.space_after  = Pt(0) if i < len(lines)-1 else Pt(6)
        run = p.add_run(line if line else " ")
        run.font.name  = 'Consolas'
        run.font.size  = Pt(9.5)
        run.font.color.rgb = RGBColor(0x19, 0x19, 0x70)   # dark blue
    # blank spacer
    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Pt(2)


def example_block(text):
    """Green-tinted example code block."""
    sub("Example")
    syntax_block(text, fill=GREEN_FILL)


def param_table(headers, rows, col_widths=None):
    """Clean bordered parameter/attribute reference table."""
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Light Grid Accent 1'

    # header row
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        _cell_shade(cell, "04AA6D")
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    # data rows
    for ri, row in enumerate(rows, 1):
        fill = "FFFFFF" if ri % 2 == 1 else "F9F9F9"
        for ci, val in enumerate(row):
            cell = t.rows[ri].cells[ci]
            _cell_shade(cell, fill)
            cell.text = val
            for run in cell.paragraphs[0].runs:
                run.font.size = Pt(10)

    if col_widths:
        for row in t.rows:
            for ci, w in enumerate(col_widths):
                row.cells[ci].width = Cm(w)

    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Pt(4)


def divider():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot = OxmlElement('w:bottom')
    bot.set(qn('w:val'),   'single')
    bot.set(qn('w:sz'),    '4')
    bot.set(qn('w:space'), '1')
    bot.set(qn('w:color'), 'CCCCCC')
    pBdr.append(bot)
    pPr.append(pBdr)


# ══════════════════════════════════════════════════════════════════════════
#  COVER
# ══════════════════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Python Flask — Syntax Reference')
r.bold = True
r.font.size = Pt(22)
r.font.color.rgb = GREEN

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = p2.add_run('COMP4442 · W3Schools Style · Complete object and method reference')
r2.font.size = Pt(11)
r2.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════
#  1. Flask App — __init__
# ══════════════════════════════════════════════════════════════════════════
topic("Flask App Setup")
body("The Flask class creates your web application. Import it and call it with __name__ "
     "so Flask knows where to find templates and static files.")

sub("Syntax")
syntax_block(
    "from flask import Flask\n"
    "\n"
    "app = Flask(__name__)\n"
    "app.secret_key = 'your-secret-key'   # required for sessions\n"
    "\n"
    "if __name__ == '__main__':\n"
    "    app.run(host='0.0.0.0', port=5000, debug=True)"
)

sub("Flask() Parameters")
param_table(
    ["Parameter", "Type", "Default", "Description"],
    [
        ("import_name",      "str",  "required", "Pass __name__ — sets the root path for templates/static"),
        ("template_folder",  "str",  "'templates'","Folder name for Jinja2 templates"),
        ("static_folder",    "str",  "'static'",  "Folder name for static files (CSS, JS, images)"),
        ("static_url_path",  "str",  "'/static'", "URL prefix for static files"),
    ],
    col_widths=[3.5, 2.0, 3.0, 8.5]
)

sub("app.run() Parameters")
param_table(
    ["Parameter", "Type", "Default", "Description"],
    [
        ("host",    "str",  "'127.0.0.1'", "Interface to listen on. '0.0.0.0' = all interfaces (needed in Docker/EC2)"),
        ("port",    "int",  "5000",        "TCP port number"),
        ("debug",   "bool", "False",       "Auto-reload on code change; show debugger. NEVER True in production"),
        ("threaded","bool", "True",        "Handle each request in its own thread"),
    ],
    col_widths=[3.0, 2.0, 3.0, 9.0]
)

sub("app.config — Common Keys")
param_table(
    ["Key", "Type", "Description"],
    [
        ("SECRET_KEY",                    "str",  "Signs session cookies with HMAC. Must be set before using session."),
        ("DEBUG",                         "bool", "Enable debug mode. Equivalent to app.run(debug=True)."),
        ("TESTING",                       "bool", "Enable testing mode. Propagates exceptions instead of handling them."),
        ("SQLALCHEMY_DATABASE_URI",       "str",  "Database connection string for Flask-SQLAlchemy."),
        ("SQLALCHEMY_TRACK_MODIFICATIONS","bool", "Set False to suppress deprecation warning."),
        ("MAX_CONTENT_LENGTH",            "int",  "Max upload size in bytes. e.g. 16 * 1024 * 1024 = 16 MB."),
        ("JSON_SORT_KEYS",                "bool", "Sort keys in jsonify() output. Default True."),
    ],
    col_widths=[5.5, 2.0, 9.5]
)

example_block(
    "from flask import Flask\n"
    "\n"
    "app = Flask(__name__)\n"
    "app.config['SECRET_KEY'] = 'change-me-in-prod'\n"
    "app.config['DEBUG'] = True\n"
    "app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024   # 16 MB max upload\n"
    "\n"
    "# Elastic Beanstalk requires the variable to be named 'application'\n"
    "application = app\n"
    "\n"
    "if __name__ == '__main__':\n"
    "    app.run(host='0.0.0.0', port=5000, debug=True)"
)

note("In production use a real WSGI server (gunicorn, uWSGI). app.run() is for development only.")
divider()

# ══════════════════════════════════════════════════════════════════════════
#  2. @app.route
# ══════════════════════════════════════════════════════════════════════════
topic("@app.route()  —  Routing")
body("The @app.route() decorator binds a URL pattern to a Python function (called a view function). "
     "When Flask receives a request matching that URL it calls the function and returns its result as the HTTP response.")

sub("Syntax")
syntax_block(
    "@app.route(rule, methods=['GET'], **options)\n"
    "def view_function():\n"
    "    return response"
)

sub("Parameters")
param_table(
    ["Parameter", "Type", "Default", "Description"],
    [
        ("rule",          "str",        "required",   "URL pattern. e.g. '/users' or '/users/<int:id>'"),
        ("methods",       "list[str]",  "['GET']",    "HTTP methods allowed. e.g. ['GET','POST'] or ['GET','PUT','DELETE']"),
        ("strict_slashes", "bool",      "True",       "If False, both /path and /path/ match. Default: /path/ redirects to /path"),
        ("endpoint",      "str",        "function name","Name used in url_for(). Defaults to the function name."),
    ],
    col_widths=[3.5, 3.0, 3.5, 7.0]
)

sub("URL Variable Types")
param_table(
    ["Converter", "Syntax", "Matches", "Python Type"],
    [
        ("string", "<name>",        "Any text without a slash (default)", "str"),
        ("int",    "<int:name>",    "Positive integers",                  "int"),
        ("float",  "<float:name>",  "Positive floating point numbers",    "float"),
        ("path",   "<path:name>",   "Like string but also accepts slashes","str"),
        ("uuid",   "<uuid:name>",   "UUID strings",                       "uuid.UUID"),
    ],
    col_widths=[3.0, 4.0, 6.0, 4.0]
)

example_block(
    "from flask import Flask, jsonify\n"
    "\n"
    "app = Flask(__name__)\n"
    "\n"
    "# ── GET only (default) ──\n"
    "@app.route('/')\n"
    "def index():\n"
    "    return 'Hello World'\n"
    "\n"
    "# ── Multiple methods ──\n"
    "@app.route('/users', methods=['GET', 'POST'])\n"
    "def users():\n"
    "    if request.method == 'POST':\n"
    "        return jsonify(created=True), 201\n"
    "    return jsonify(users=[]), 200\n"
    "\n"
    "# ── URL variable ──\n"
    "@app.route('/users/<int:user_id>', methods=['GET', 'PUT', 'DELETE'])\n"
    "def user_detail(user_id):          # user_id is passed as int\n"
    "    return jsonify(id=user_id)\n"
    "\n"
    "# ── Path variable (allows slashes) ──\n"
    "@app.route('/files/<path:filename>')\n"
    "def serve_file(filename):          # filename may contain '/'\n"
    "    return f'Serving {filename}'"
)
divider()

# ══════════════════════════════════════════════════════════════════════════
#  3. url_for / redirect / abort
# ══════════════════════════════════════════════════════════════════════════
topic("url_for()  —  Reverse URL Building")
body("url_for() builds a URL from the name of a view function instead of hardcoding strings. "
     "This ensures links stay correct if you ever change the route pattern.")

sub("Syntax")
syntax_block("url_for(endpoint, **values)")

sub("Parameters")
param_table(
    ["Parameter", "Type", "Description"],
    [
        ("endpoint", "str",    "Name of the view function (or 'blueprint.function' for blueprints)"),
        ("**values", "kwargs", "URL variables to fill in. Extra kwargs become query string params."),
        ("_external","bool",   "If True, return absolute URL with scheme and host (for emails, redirects)."),
        ("_anchor",  "str",    "Add a fragment: url_for('index', _anchor='top') → '/#top'"),
    ],
    col_widths=[3.0, 2.5, 11.5]
)

example_block(
    "from flask import url_for, redirect\n"
    "\n"
    "# Inside a request context:\n"
    "url_for('index')                        # → '/'\n"
    "url_for('user_detail', user_id=42)      # → '/users/42'\n"
    "url_for('users', page=2)                # → '/users?page=2'  (extra = query string)\n"
    "url_for('static', filename='app.js')    # → '/static/app.js'\n"
    "url_for('index', _external=True)        # → 'http://localhost:5000/'\n"
    "\n"
    "# Blueprint route:\n"
    "url_for('auth.login')                   # → '/auth/login'"
)

divider()

topic("redirect()  and  abort()")
body("redirect() returns an HTTP redirect response. abort() immediately stops the request and "
     "returns an error response — no code after abort() runs.")

sub("Syntax")
syntax_block(
    "redirect(location, code=302)\n"
    "\n"
    "abort(status_code)"
)

sub("redirect() Parameters")
param_table(
    ["Parameter", "Type", "Default", "Description"],
    [
        ("location", "str",  "required", "The URL to redirect to. Use url_for() to build it."),
        ("code",     "int",  "302",      "HTTP status code. 301 = permanent, 302 = temporary (default), 303 = see other."),
    ],
    col_widths=[3.0, 2.0, 2.5, 9.5]
)

sub("abort() — Common Status Codes")
param_table(
    ["Code", "Meaning",          "When to use"],
    [
        ("400", "Bad Request",    "Request body is malformed or missing required fields"),
        ("401", "Unauthorized",   "User is not logged in / no valid token"),
        ("403", "Forbidden",      "User is logged in but lacks permission"),
        ("404", "Not Found",      "Resource does not exist"),
        ("405", "Method Not Allowed", "Route exists but not for this HTTP method"),
        ("409", "Conflict",       "Duplicate resource (e.g. username already taken)"),
        ("422", "Unprocessable",  "Input is syntactically valid but semantically wrong"),
        ("500", "Internal Error", "Unexpected server-side failure"),
    ],
    col_widths=[2.0, 4.0, 11.0]
)

example_block(
    "from flask import redirect, url_for, abort, jsonify\n"
    "\n"
    "@app.route('/dashboard')\n"
    "def dashboard():\n"
    "    if 'user' not in session:\n"
    "        return redirect(url_for('login'))       # 302 redirect to /login\n"
    "    return jsonify(welcome=session['user'])\n"
    "\n"
    "@app.route('/users/<int:uid>')\n"
    "def get_user(uid):\n"
    "    user = db.get(uid)\n"
    "    if user is None:\n"
    "        abort(404)                              # stops here, returns 404\n"
    "    if user['role'] != 'admin' and uid != session.get('uid'):\n"
    "        abort(403)                              # forbidden\n"
    "    return jsonify(user)"
)
divider()

# ══════════════════════════════════════════════════════════════════════════
#  4. request
# ══════════════════════════════════════════════════════════════════════════
topic("The  request  Object")
body("The request object gives you access to all data the client sent with the HTTP request. "
     "It is available inside any view function automatically — no need to pass it as a parameter. "
     "Import it from flask at the top of your file.")

sub("Import")
syntax_block("from flask import request")

sub("Attribute Reference")
param_table(
    ["Attribute", "Type", "Description"],
    [
        ("request.method",      "str",               "HTTP verb: 'GET', 'POST', 'PUT', 'DELETE', 'PATCH'"),
        ("request.args",        "ImmutableMultiDict", "Query string parameters  (?key=value after the URL)"),
        ("request.form",        "ImmutableMultiDict", "Form body data  (POST with Content-Type: application/x-www-form-urlencoded or multipart/form-data)"),
        ("request.json",        "dict | None",        "Parsed JSON body. None if Content-Type is not application/json."),
        ("request.data",        "bytes",              "Raw request body as bytes"),
        ("request.headers",     "Headers",            "HTTP request headers  (dict-like, case-insensitive)"),
        ("request.cookies",     "dict",               "Cookies sent by the browser"),
        ("request.files",       "ImmutableMultiDict", "Uploaded files  (multipart/form-data)"),
        ("request.path",        "str",                "URL path without query string.  e.g.  '/users/42'"),
        ("request.url",         "str",                "Full request URL including query string"),
        ("request.base_url",    "str",                "URL without query string"),
        ("request.host",        "str",                "Hostname and port.  e.g.  'localhost:5000'"),
        ("request.remote_addr", "str",                "Client IP address"),
        ("request.is_json",     "bool",               "True when Content-Type is application/json"),
    ],
    col_widths=[4.0, 3.5, 9.5]
)

sub("Method Reference")
param_table(
    ["Method", "Returns", "Description"],
    [
        ("request.get_json(force=False, silent=False, cache=True)",
         "dict | None",
         "Parse JSON body with options.  force=True ignores Content-Type.  silent=True returns None on bad JSON instead of raising 400."),
        ("request.args.get(key, default=None)",
         "str | None",
         "Safely read a query-string param. Returns default if the key is absent."),
        ("request.args.getlist(key)",
         "list[str]",
         "Return all values for a repeated query param.  e.g. ?tag=a&tag=b → ['a','b']"),
        ("request.form.get(key, default=None)",
         "str | None",
         "Safely read a form field. Returns default if absent."),
        ("request.form.getlist(key)",
         "list[str]",
         "Return all values for a repeated form field (e.g. multi-select checkbox)."),
        ("request.headers.get(key, default=None)",
         "str | None",
         "Read an HTTP header.  e.g. request.headers.get('Authorization')"),
        ("request.cookies.get(key, default=None)",
         "str | None",
         "Safely read a cookie value."),
    ],
    col_widths=[6.5, 2.5, 8.0]
)

example_block(
    "from flask import Flask, request, jsonify\n"
    "\n"
    "app = Flask(__name__)\n"
    "\n"
    "@app.route('/search')                        # GET /search?q=flask&page=2\n"
    "def search():\n"
    "    q    = request.args.get('q', '')         # 'flask'\n"
    "    page = request.args.get('page', 1, type=int)  # 2  (cast to int)\n"
    "    return jsonify(query=q, page=page)\n"
    "\n"
    "\n"
    "@app.route('/login', methods=['POST'])       # POST with JSON body\n"
    "def login():\n"
    "    data = request.get_json(silent=True)     # None if bad/missing JSON\n"
    "    if not data or 'username' not in data:\n"
    "        return jsonify(error='bad request'), 400\n"
    "    return jsonify(user=data['username']), 200\n"
    "\n"
    "\n"
    "@app.route('/upload', methods=['POST'])      # file upload\n"
    "def upload():\n"
    "    f = request.files.get('photo')           # None if not sent\n"
    "    if not f:\n"
    "        return jsonify(error='no file'), 400\n"
    "    f.save('/uploads/' + f.filename)\n"
    "    return jsonify(saved=f.filename), 201"
)

note("Use request.get_json(silent=True) instead of request.json — silent=True "
     "returns None on malformed JSON instead of a 400 error, so you can send your own error message.")
divider()

# ══════════════════════════════════════════════════════════════════════════
#  5. session
# ══════════════════════════════════════════════════════════════════════════
topic("The  session  Object")
body("The session object stores data across multiple requests for the same user. "
     "By default, Flask stores the entire session as a signed cookie on the client "
     "(not on the server). The data is base64-encoded (visible) but tamper-proof "
     "(signed with HMAC using app.secret_key).")

sub("Import")
syntax_block("from flask import session")

sub("Requirement")
syntax_block("app.secret_key = 'your-secret-key'   # MUST be set before using session")

sub("Syntax — All Operations")
syntax_block(
    "# Write\n"
    "session['key'] = value\n"
    "\n"
    "# Read (safe — returns None if key is missing)\n"
    "session.get('key')\n"
    "session.get('key', default_value)\n"
    "\n"
    "# Read (unsafe — raises KeyError if missing)\n"
    "session['key']\n"
    "\n"
    "# Check if key exists\n"
    "'key' in session\n"
    "\n"
    "# Remove one key (safe — None prevents KeyError)\n"
    "session.pop('key', None)\n"
    "\n"
    "# Remove one key (unsafe — raises KeyError if missing)\n"
    "del session['key']\n"
    "\n"
    "# Remove ALL keys  (logout)\n"
    "session.clear()\n"
    "\n"
    "# Make session persist after browser close\n"
    "session.permanent = True"
)

sub("Method / Attribute Reference")
param_table(
    ["Operation",              "Code",                        "Description"],
    [
        ("Set value",          "session['key'] = value",      "Store any JSON-serialisable value (str, int, list, dict)"),
        ("Get — safe",         "session.get('key')",          "Returns None if key does not exist. Always prefer over []"),
        ("Get — with default", "session.get('key', default)", "Returns default if key does not exist"),
        ("Get — unsafe",       "session['key']",              "Returns value. Raises KeyError if key missing. Avoid in production"),
        ("Check key exists",   "'key' in session",            "Returns True / False"),
        ("Remove one — safe",  "session.pop('key', None)",    "Removes key. The None argument suppresses KeyError if absent"),
        ("Remove one — unsafe","del session['key']",          "Removes key. Raises KeyError if absent"),
        ("Remove all",         "session.clear()",             "Wipes the entire session — standard logout approach"),
        ("Make permanent",     "session.permanent = True",    "Session survives browser close. Set app.permanent_session_lifetime for expiry duration"),
    ],
    col_widths=[4.0, 5.0, 8.0]
)

example_block(
    "from flask import Flask, request, jsonify, session\n"
    "\n"
    "app = Flask(__name__)\n"
    "app.secret_key = 'super-secret'\n"
    "\n"
    "\n"
    "# ── Login: write to session ──\n"
    "@app.route('/login', methods=['POST'])\n"
    "def login():\n"
    "    data = request.get_json(silent=True) or {}\n"
    "    if data.get('password') != 'correct':\n"
    "        return jsonify(error='bad credentials'), 401\n"
    "    session['username'] = data['username']   # store in session\n"
    "    session['role']     = 'admin'\n"
    "    return jsonify(message='logged in'), 200\n"
    "\n"
    "\n"
    "# ── Protected route: read from session ──\n"
    "@app.route('/profile')\n"
    "def profile():\n"
    "    if 'username' not in session:            # check key exists\n"
    "        return jsonify(error='not logged in'), 401\n"
    "    return jsonify(\n"
    "        user=session.get('username'),        # safe read\n"
    "        role=session.get('role', 'user')     # safe read with default\n"
    "    ), 200\n"
    "\n"
    "\n"
    "# ── Logout: clear session ──\n"
    "@app.route('/logout', methods=['POST'])\n"
    "def logout():\n"
    "    session.clear()                          # remove all keys\n"
    "    return jsonify(message='logged out'), 200\n"
    "\n"
    "\n"
    "# ── Remove one key only ──\n"
    "@app.route('/forget-role', methods=['POST'])\n"
    "def forget_role():\n"
    "    session.pop('role', None)                # remove 'role', safe if absent\n"
    "    return jsonify(message='role removed'), 200"
)

note("Flask's default session is CLIENT-SIDE. The full dict is stored in the browser cookie — "
     "not on the server. It is signed (HMAC) but NOT encrypted, so users can read the data. "
     "Never store passwords or sensitive info in session. Use Flask-Session + Redis for server-side storage.")
divider()

# ══════════════════════════════════════════════════════════════════════════
#  6. Cookies
# ══════════════════════════════════════════════════════════════════════════
topic("Cookies  —  set_cookie()  /  delete_cookie()")
body("Unlike session (which Flask manages automatically), cookies must be set and deleted "
     "manually on a Response object. Use make_response() to get a Response you can attach cookies to.")

sub("Syntax")
syntax_block(
    "# Read a cookie from the incoming request\n"
    "value = request.cookies.get('cookie_name')\n"
    "\n"
    "# Set a cookie on the outgoing response\n"
    "resp = make_response(jsonify(ok=True))\n"
    "resp.set_cookie('cookie_name', value, **options)\n"
    "return resp\n"
    "\n"
    "# Delete a cookie\n"
    "resp.delete_cookie('cookie_name')"
)

sub("set_cookie() Parameters")
param_table(
    ["Parameter", "Type", "Default", "Description"],
    [
        ("key",      "str",           "required", "Cookie name"),
        ("value",    "str",           "''",       "Cookie value (must be a string)"),
        ("max_age",  "int | None",    "None",     "Seconds until expiry. None = session cookie (deleted when browser closes)"),
        ("expires",  "datetime|None", "None",     "Exact expiry as a datetime object. max_age takes precedence if both set"),
        ("path",     "str",           "'/'",      "URL path scope. '/' = sent with every request to this domain"),
        ("domain",   "str | None",    "None",     "Cookie domain. None = current domain only"),
        ("secure",   "bool",          "False",    "True = only sent over HTTPS. Always True in production"),
        ("httponly", "bool",          "False",    "True = JavaScript cannot read this cookie (prevents XSS theft)"),
        ("samesite", "str | None",    "None",     "'Strict' | 'Lax' | 'None'. Controls cross-site sending (CSRF protection)"),
    ],
    col_widths=[3.0, 3.0, 2.5, 8.5]
)

example_block(
    "from flask import request, make_response, jsonify\n"
    "\n"
    "\n"
    "# ── Read a cookie ──\n"
    "@app.route('/whoami')\n"
    "def whoami():\n"
    "    token = request.cookies.get('auth_token')   # None if absent\n"
    "    if not token:\n"
    "        return jsonify(error='no cookie'), 401\n"
    "    return jsonify(token=token), 200\n"
    "\n"
    "\n"
    "# ── Set a cookie ──\n"
    "@app.route('/set-cookie')\n"
    "def set_cookie_view():\n"
    "    resp = make_response(jsonify(message='cookie set'))\n"
    "    resp.set_cookie(\n"
    "        'auth_token',\n"
    "        value='abc123xyz',\n"
    "        max_age=3600,          # expires in 1 hour\n"
    "        secure=True,           # HTTPS only\n"
    "        httponly=True,         # not accessible via JavaScript\n"
    "        samesite='Lax'         # CSRF protection\n"
    "    )\n"
    "    return resp\n"
    "\n"
    "\n"
    "# ── Delete a cookie ──\n"
    "@app.route('/clear-cookie')\n"
    "def clear_cookie_view():\n"
    "    resp = make_response(jsonify(message='cookie cleared'))\n"
    "    resp.delete_cookie('auth_token')\n"
    "    return resp"
)
divider()

# ══════════════════════════════════════════════════════════════════════════
#  7. jsonify / make_response
# ══════════════════════════════════════════════════════════════════════════
topic("jsonify()  and  make_response()")
body("jsonify() converts a Python dict or list into a Flask Response with Content-Type: application/json. "
     "make_response() gives you a Response object you can customise before returning.")

sub("Syntax")
syntax_block(
    "jsonify(*args, **kwargs)\n"
    "\n"
    "make_response(response_body, status_code, headers)"
)

sub("Return Value Shorthand")
param_table(
    ["What you return",                 "Flask interprets it as"],
    [
        ("return 'text'",               "200 OK, Content-Type: text/html, body = 'text'"),
        ("return 'text', 404",          "404 response with body 'text'"),
        ("return 'text', 201, {'X':'v'}","201 response with extra header X: v"),
        ("return jsonify(…)",           "200 OK, Content-Type: application/json"),
        ("return jsonify(…), 201",      "201 Created, Content-Type: application/json"),
        ("return '', 204",              "204 No Content — empty body (use for DELETE)"),
    ],
    col_widths=[5.5, 11.5]
)

example_block(
    "from flask import jsonify, make_response\n"
    "\n"
    "\n"
    "# ── jsonify with dict ──\n"
    "@app.route('/user')\n"
    "def user():\n"
    "    return jsonify({'name': 'Alice', 'age': 30}), 200\n"
    "\n"
    "\n"
    "# ── jsonify with keyword args (equivalent) ──\n"
    "@app.route('/status')\n"
    "def status():\n"
    "    return jsonify(status='ok', version='1.0'), 200\n"
    "\n"
    "\n"
    "# ── jsonify a list ──\n"
    "@app.route('/items')\n"
    "def items():\n"
    "    return jsonify([{'id': 1}, {'id': 2}]), 200\n"
    "\n"
    "\n"
    "# ── make_response for custom headers / cookies ──\n"
    "@app.route('/custom')\n"
    "def custom():\n"
    "    resp = make_response(jsonify(message='hello'), 200)\n"
    "    resp.headers['X-Request-Id'] = 'abc-123'\n"
    "    resp.headers['Cache-Control'] = 'no-cache'\n"
    "    resp.set_cookie('visited', '1', max_age=86400)\n"
    "    return resp\n"
    "\n"
    "\n"
    "# ── 204 No Content (DELETE with no body) ──\n"
    "@app.route('/items/<int:id>', methods=['DELETE'])\n"
    "def delete_item(id):\n"
    "    # ... delete from DB ...\n"
    "    return '', 204"
)
divider()

# ══════════════════════════════════════════════════════════════════════════
#  8. render_template
# ══════════════════════════════════════════════════════════════════════════
topic("render_template()  —  HTML Templates")
body("render_template() renders a Jinja2 HTML template file from the templates/ folder "
     "and returns it as a string. Variables passed as keyword arguments are available inside the template.")

sub("Syntax")
syntax_block(
    "render_template(template_name_or_list, **context)\n"
    "\n"
    "# Render from a string (no file needed)\n"
    "render_template_string(source_string, **context)"
)

sub("Parameters")
param_table(
    ["Parameter",              "Type",       "Description"],
    [
        ("template_name_or_list","str | list","Filename relative to templates/.  e.g. 'index.html' or 'admin/panel.html'"),
        ("**context",           "kwargs",     "Variables injected into the template. Accessible as {{ name }} in Jinja2."),
    ],
    col_widths=[4.5, 3.0, 9.5]
)

example_block(
    "# ── app.py ──\n"
    "from flask import render_template\n"
    "\n"
    "@app.route('/hello/<name>')\n"
    "def hello(name):\n"
    "    return render_template('hello.html', username=name, items=['a','b','c'])"
)

sub("Jinja2 Template Syntax  (hello.html)")
syntax_block(
    "{# comment — not rendered to HTML #}\n"
    "\n"
    "<!-- Output a variable -->\n"
    "<h1>Hello, {{ username }}</h1>\n"
    "\n"
    "<!-- Attribute access / index -->\n"
    "<p>{{ user.email }}   or   {{ user['email'] }}</p>\n"
    "\n"
    "<!-- Filters  (pipe character) -->\n"
    "{{ username | upper }}          converts to UPPERCASE\n"
    "{{ username | lower }}          lowercase\n"
    "{{ username | title }}          Title Case\n"
    "{{ username | trim }}           strip whitespace\n"
    "{{ price    | round(2) }}       round to 2 decimal places\n"
    "{{ items    | length }}         count items\n"
    "{{ name     | default('guest') }}   fallback if None/undefined\n"
    "{{ html     | safe }}           mark as safe — disables auto-escaping (XSS risk!)\n"
    "\n"
    "<!-- if / elif / else -->\n"
    "{% if user.role == 'admin' %}\n"
    "  <a href='/admin'>Admin</a>\n"
    "{% elif user.role == 'mod' %}\n"
    "  <a href='/mod'>Moderator</a>\n"
    "{% else %}\n"
    "  <p>Welcome, guest</p>\n"
    "{% endif %}\n"
    "\n"
    "<!-- for loop -->\n"
    "{% for item in items %}\n"
    "  <li>{{ loop.index }}. {{ item }}</li>\n"
    "{% else %}\n"
    "  <li>No items found.</li>     {# rendered only when list is empty #}\n"
    "{% endfor %}\n"
    "\n"
    "<!-- Loop helper variables -->\n"
    "loop.index      current iteration (1-based)\n"
    "loop.index0     current iteration (0-based)\n"
    "loop.first      True on first iteration\n"
    "loop.last       True on last iteration\n"
    "loop.length     total count\n"
    "\n"
    "<!-- Set a variable -->\n"
    "{% set total = price * qty %}\n"
    "\n"
    "<!-- Template inheritance -->\n"
    "{# base.html #}\n"
    "<!DOCTYPE html><html><body>\n"
    "  {% block content %}{% endblock %}\n"
    "</body></html>\n"
    "\n"
    "{# child.html #}\n"
    "{% extends 'base.html' %}\n"
    "{% block content %}\n"
    "  <h1>Child content here</h1>\n"
    "{% endblock %}\n"
    "\n"
    "<!-- Include a partial template -->\n"
    "{% include 'partials/nav.html' %}"
)
divider()

# ══════════════════════════════════════════════════════════════════════════
#  9. Error Handlers
# ══════════════════════════════════════════════════════════════════════════
topic("@app.errorhandler()  —  Custom Error Pages")
body("@app.errorhandler() registers a function to call when Flask would return a specific HTTP error code. "
     "Without it, Flask returns a plain HTML error page. With it, you can return JSON (for APIs) or a custom template.")

sub("Syntax")
syntax_block(
    "@app.errorhandler(status_code)\n"
    "def handler_name(error):\n"
    "    return response, status_code"
)

example_block(
    "from flask import jsonify\n"
    "\n"
    "@app.errorhandler(400)\n"
    "def bad_request(error):\n"
    "    return jsonify(error='bad request',  message=str(error)), 400\n"
    "\n"
    "@app.errorhandler(401)\n"
    "def unauthorized(error):\n"
    "    return jsonify(error='unauthorized'), 401\n"
    "\n"
    "@app.errorhandler(403)\n"
    "def forbidden(error):\n"
    "    return jsonify(error='forbidden'), 403\n"
    "\n"
    "@app.errorhandler(404)\n"
    "def not_found(error):\n"
    "    return jsonify(error='not found',    path=request.path), 404\n"
    "\n"
    "@app.errorhandler(500)\n"
    "def server_error(error):\n"
    "    return jsonify(error='server error'), 500\n"
    "\n"
    "# Catch ALL unhandled exceptions\n"
    "@app.errorhandler(Exception)\n"
    "def handle_all(e):\n"
    "    return jsonify(error=str(e)), 500"
)
divider()

# ══════════════════════════════════════════════════════════════════════════
#  10. before_request / after_request / g
# ══════════════════════════════════════════════════════════════════════════
topic("@before_request  /  @after_request  /  g")
body("These hooks let you run code before or after every request without repeating it in each view. "
     "g is a per-request storage object — data put into g lasts until the response is sent, "
     "then it is discarded.")

sub("Syntax")
syntax_block(
    "@app.before_request\n"
    "def function_name():\n"
    "    # runs before every request\n"
    "    # return a response here to block the request (short-circuit)\n"
    "\n"
    "@app.after_request\n"
    "def function_name(response):\n"
    "    # runs after every successful request\n"
    "    return response       # MUST return the response object\n"
    "\n"
    "@app.teardown_request\n"
    "def function_name(exception=None):\n"
    "    # runs after every request, even if an exception occurred\n"
    "    # used to close DB connections, clean up resources"
)

sub("g Object — Attribute Reference")
param_table(
    ["Operation",          "Code",                  "Description"],
    [
        ("Set value",      "g.user = 'alice'",       "Store anything on g for use later in the same request"),
        ("Read value",     "g.user",                 "AttributeError if not set — use g.get() to be safe"),
        ("Read — safe",    "g.get('user')",           "Returns None if attribute was never set"),
        ("Check exists",   "hasattr(g, 'user')",      "Returns True / False"),
        ("Remove",         "g.pop('user', None)",     "Remove attribute safely"),
    ],
    col_widths=[3.5, 4.5, 9.0]
)

example_block(
    "from flask import g, request, jsonify\n"
    "\n"
    "\n"
    "# ── before_request: authenticate every request ──\n"
    "@app.before_request\n"
    "def authenticate():\n"
    "    # Skip auth for public routes\n"
    "    if request.path in ('/login', '/health'):\n"
    "        return\n"
    "    token = request.headers.get('Authorization', '').replace('Bearer ', '')\n"
    "    if not token:\n"
    "        return jsonify(error='missing token'), 401   # return = block request\n"
    "    g.user = verify_token(token)                     # store on g for view to use\n"
    "\n"
    "\n"
    "# ── view uses g.user (set by before_request) ──\n"
    "@app.route('/profile')\n"
    "def profile():\n"
    "    return jsonify(user=g.user)\n"
    "\n"
    "\n"
    "# ── after_request: add CORS headers to every response ──\n"
    "@app.after_request\n"
    "def add_cors_headers(response):\n"
    "    response.headers['Access-Control-Allow-Origin']  = '*'\n"
    "    response.headers['Access-Control-Allow-Headers'] = 'Content-Type,Authorization'\n"
    "    return response      # must return the response\n"
    "\n"
    "\n"
    "# ── teardown_request: close DB connection ──\n"
    "@app.teardown_request\n"
    "def close_connection(exception=None):\n"
    "    conn = g.pop('db_conn', None)\n"
    "    if conn is not None:\n"
    "        conn.close()"
)
divider()

# ══════════════════════════════════════════════════════════════════════════
#  11. Blueprint
# ══════════════════════════════════════════════════════════════════════════
topic("Blueprint  —  Modular Routing")
body("A Blueprint groups related routes, error handlers, and template filters into a reusable module. "
     "Each blueprint is registered on the app with a URL prefix, keeping large apps organised.")

sub("Syntax")
syntax_block(
    "# ── define.py ──\n"
    "from flask import Blueprint\n"
    "\n"
    "bp = Blueprint('blueprint_name', __name__, url_prefix='/prefix')\n"
    "\n"
    "@bp.route('/path')\n"
    "def view(): ...\n"
    "\n"
    "\n"
    "# ── app.py ──\n"
    "from define import bp\n"
    "app.register_blueprint(bp)"
)

sub("Blueprint() Parameters")
param_table(
    ["Parameter",        "Type",       "Default", "Description"],
    [
        ("name",         "str",        "required","Unique name. Used as prefix in url_for(): url_for('name.view')"),
        ("import_name",  "str",        "required","Pass __name__ — same as Flask(__name__)"),
        ("url_prefix",   "str | None", "None",    "Prepended to all blueprint routes. e.g. '/api/v1'"),
        ("template_folder","str|None", "None",    "Folder for blueprint-specific templates"),
        ("static_folder","str | None", "None",    "Folder for blueprint-specific static files"),
    ],
    col_widths=[3.5, 2.5, 2.5, 8.5]
)

example_block(
    "# ── auth.py ──\n"
    "from flask import Blueprint, request, jsonify, session\n"
    "\n"
    "auth = Blueprint('auth', __name__, url_prefix='/auth')\n"
    "\n"
    "@auth.route('/login', methods=['POST'])    # → POST /auth/login\n"
    "def login():\n"
    "    session['user'] = request.get_json()['username']\n"
    "    return jsonify(ok=True), 200\n"
    "\n"
    "@auth.route('/logout', methods=['POST'])   # → POST /auth/logout\n"
    "def logout():\n"
    "    session.clear()\n"
    "    return jsonify(ok=True), 200\n"
    "\n"
    "\n"
    "# ── app.py ──\n"
    "from flask import Flask\n"
    "from auth import auth\n"
    "\n"
    "app = Flask(__name__)\n"
    "app.secret_key = 'secret'\n"
    "app.register_blueprint(auth)\n"
    "\n"
    "# url_for with blueprint:\n"
    "# url_for('auth.login')   → '/auth/login'\n"
    "# url_for('auth.logout')  → '/auth/logout'"
)
divider()

# ══════════════════════════════════════════════════════════════════════════
#  12. boto3 DynamoDB
# ══════════════════════════════════════════════════════════════════════════
topic("boto3  —  DynamoDB  (AWS SDK for Python)")
body("boto3 is the AWS SDK for Python. Use it inside Flask or Lambda to talk to DynamoDB, S3, SQS, etc. "
     "Create the client/resource at module scope so it is reused across warm Lambda invocations.")

sub("Setup")
syntax_block(
    "import boto3\n"
    "from boto3.dynamodb.conditions import Key, Attr\n"
    "\n"
    "dynamodb = boto3.resource('dynamodb', region_name='ap-east-1')\n"
    "table    = dynamodb.Table('TableName')   # module scope — reused each invocation"
)

sub("Method Reference")
param_table(
    ["Method",               "Description"],
    [
        ("table.put_item(Item={…})",
         "Create or fully overwrite an item. Pass the full item dict including primary key."),
        ("table.get_item(Key={…})",
         "Read one item by primary key. Returns {'Item': {…}} or {} if not found. Access with .get('Item')."),
        ("table.update_item(Key, UpdateExpression, ExpressionAttributeValues)",
         "Partially update fields. Use SET to add/change, REMOVE to delete fields. Does NOT overwrite whole item."),
        ("table.delete_item(Key={…})",
         "Delete one item by primary key."),
        ("table.query(KeyConditionExpression=…)",
         "Efficient read using a key condition. Always prefer over scan. Returns {'Items': […]}."),
        ("table.scan(FilterExpression=…)",
         "Read entire table then filter. Expensive — avoid in production for large tables."),
    ],
    col_widths=[6.5, 10.5]
)

example_block(
    "from boto3.dynamodb.conditions import Key, Attr\n"
    "\n"
    "\n"
    "# ── Create / Overwrite ──\n"
    "table.put_item(Item={\n"
    "    'user_id': 'u001',\n"
    "    'name':    'Alice',\n"
    "    'age':     30\n"
    "})\n"
    "\n"
    "\n"
    "# ── Read one item ──\n"
    "resp = table.get_item(Key={'user_id': 'u001'})\n"
    "item = resp.get('Item')         # None if not found\n"
    "\n"
    "\n"
    "# ── Partial update (SET two fields) ──\n"
    "table.update_item(\n"
    "    Key={'user_id': 'u001'},\n"
    "    UpdateExpression='SET #n = :name, age = :age',\n"
    "    ExpressionAttributeNames ={'#n': 'name'},    # #n avoids reserved word\n"
    "    ExpressionAttributeValues={':name': 'Bob', ':age': 31}\n"
    ")\n"
    "\n"
    "\n"
    "# ── Delete ──\n"
    "table.delete_item(Key={'user_id': 'u001'})\n"
    "\n"
    "\n"
    "# ── Query by partition key (fast — uses index) ──\n"
    "resp  = table.query(\n"
    "    KeyConditionExpression=Key('user_id').eq('u001')\n"
    ")\n"
    "items = resp['Items']\n"
    "\n"
    "\n"
    "# ── Query with filter on non-key attribute ──\n"
    "resp = table.query(\n"
    "    KeyConditionExpression=Key('user_id').eq('u001'),\n"
    "    FilterExpression=Attr('age').gt(25)\n"
    ")\n"
    "\n"
    "\n"
    "# ── Scan (avoid — reads whole table) ──\n"
    "resp  = table.scan(FilterExpression=Attr('role').eq('admin'))\n"
    "items = resp['Items']"
)

note("Always query a GSI (Global Secondary Index) instead of scanning. "
     "A Scan reads every item in the table — slow and expensive at scale.")
divider()

# ══════════════════════════════════════════════════════════════════════════
#  13. SQLAlchemy
# ══════════════════════════════════════════════════════════════════════════
topic("Flask-SQLAlchemy  —  Relational Database (MySQL / SQLite)")
body("Flask-SQLAlchemy is an ORM (Object-Relational Mapper) that lets you define database tables as Python "
     "classes and query them with Python methods instead of writing raw SQL.")

sub("Setup")
syntax_block(
    "from flask import Flask\n"
    "from flask_sqlalchemy import SQLAlchemy\n"
    "\n"
    "app = Flask(__name__)\n"
    "app.config['SQLALCHEMY_DATABASE_URI'] = 'mysql+pymysql://user:pass@host:3306/dbname'\n"
    "app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False\n"
    "\n"
    "db = SQLAlchemy(app)"
)

sub("Defining a Model")
syntax_block(
    "class User(db.Model):\n"
    "    __tablename__ = 'users'   # optional — defaults to class name lowercase\n"
    "\n"
    "    id       = db.Column(db.Integer,     primary_key=True, autoincrement=True)\n"
    "    username = db.Column(db.String(80),  unique=True,  nullable=False)\n"
    "    email    = db.Column(db.String(120), unique=True,  nullable=True)\n"
    "    role     = db.Column(db.String(20),  default='user')\n"
    "    created  = db.Column(db.DateTime,    default=db.func.now())\n"
    "\n"
    "    def to_dict(self):          # helper for jsonify\n"
    "        return {'id': self.id, 'username': self.username, 'role': self.role}"
)

sub("Column Types")
param_table(
    ["Type",              "SQL equivalent",  "Python type"],
    [
        ("db.Integer",    "INT",             "int"),
        ("db.String(n)",  "VARCHAR(n)",      "str"),
        ("db.Text",       "TEXT",            "str"),
        ("db.Float",      "FLOAT",           "float"),
        ("db.Boolean",    "BOOLEAN",         "bool"),
        ("db.DateTime",   "DATETIME",        "datetime"),
        ("db.Date",       "DATE",            "date"),
        ("db.JSON",       "JSON",            "dict / list"),
    ],
    col_widths=[3.5, 3.5, 3.0]
)

sub("CRUD Operations")
syntax_block(
    "# ── Create tables (run once) ──\n"
    "with app.app_context():\n"
    "    db.create_all()\n"
    "\n"
    "\n"
    "# ── Create (INSERT) ──\n"
    "user = User(username='alice', email='alice@example.com')\n"
    "db.session.add(user)\n"
    "db.session.commit()              # writes to DB; user.id is now set\n"
    "\n"
    "\n"
    "# ── Read (SELECT) ──\n"
    "user  = User.query.get(1)                              # by primary key\n"
    "user  = User.query.filter_by(username='alice').first() # first match or None\n"
    "users = User.query.all()                               # all rows\n"
    "users = User.query.filter(User.role == 'admin').all()  # with condition\n"
    "users = User.query.order_by(User.username).limit(10).all()\n"
    "count = User.query.filter_by(role='admin').count()     # count rows\n"
    "\n"
    "\n"
    "# ── Update ──\n"
    "user = User.query.get(1)\n"
    "user.email = 'new@example.com'\n"
    "db.session.commit()\n"
    "\n"
    "\n"
    "# ── Delete ──\n"
    "user = User.query.get(1)\n"
    "db.session.delete(user)\n"
    "db.session.commit()\n"
    "\n"
    "\n"
    "# ── Rollback on error ──\n"
    "try:\n"
    "    db.session.add(user)\n"
    "    db.session.commit()\n"
    "except Exception:\n"
    "    db.session.rollback()        # undo changes on failure\n"
    "    raise"
)
divider()

# ══════════════════════════════════════════════════════════════════════════
#  14. Full CRUD REST example
# ══════════════════════════════════════════════════════════════════════════
topic("Complete Example  —  Full CRUD REST API")
body("A minimal but complete REST API with all five endpoints, "
     "input validation, proper status codes, and JSON responses.")

syntax_block(
    "from flask import Flask, request, jsonify, abort\n"
    "\n"
    "app   = Flask(__name__)\n"
    "store = {}   # in-memory dict  (replace with DB in production)\n"
    "\n"
    "\n"
    "# ── GET all ──────────────────────────────────────────────────────\n"
    "@app.route('/items', methods=['GET'])\n"
    "def list_items():\n"
    "    return jsonify(list(store.values())), 200\n"
    "\n"
    "\n"
    "# ── GET one ──────────────────────────────────────────────────────\n"
    "@app.route('/items/<int:item_id>', methods=['GET'])\n"
    "def get_item(item_id):\n"
    "    item = store.get(item_id)\n"
    "    if item is None:\n"
    "        abort(404)\n"
    "    return jsonify(item), 200\n"
    "\n"
    "\n"
    "# ── POST create ──────────────────────────────────────────────────\n"
    "@app.route('/items', methods=['POST'])\n"
    "def create_item():\n"
    "    data = request.get_json(silent=True)\n"
    "    if not data or 'name' not in data:\n"
    "        return jsonify(error='name is required'), 400\n"
    "    new_id = max(store.keys(), default=0) + 1\n"
    "    store[new_id] = {'id': new_id, 'name': data['name']}\n"
    "    return jsonify(store[new_id]), 201             # 201 Created\n"
    "\n"
    "\n"
    "# ── PUT update ───────────────────────────────────────────────────\n"
    "@app.route('/items/<int:item_id>', methods=['PUT'])\n"
    "def update_item(item_id):\n"
    "    if item_id not in store:\n"
    "        abort(404)\n"
    "    data = request.get_json(silent=True) or {}\n"
    "    store[item_id].update(data)\n"
    "    return jsonify(store[item_id]), 200\n"
    "\n"
    "\n"
    "# ── DELETE ───────────────────────────────────────────────────────\n"
    "@app.route('/items/<int:item_id>', methods=['DELETE'])\n"
    "def delete_item(item_id):\n"
    "    if item_id not in store:\n"
    "        abort(404)\n"
    "    store.pop(item_id)\n"
    "    return '', 204                                 # 204 No Content\n"
    "\n"
    "\n"
    "# ── Error handlers ───────────────────────────────────────────────\n"
    "@app.errorhandler(404)\n"
    "def not_found(e):\n"
    "    return jsonify(error='not found'), 404\n"
    "\n"
    "@app.errorhandler(400)\n"
    "def bad_request(e):\n"
    "    return jsonify(error='bad request'), 400\n"
    "\n"
    "\n"
    "if __name__ == '__main__':\n"
    "    app.run(debug=True)"
)

doc.save('flask_syntax.docx')
print("flask_syntax.docx written.")
