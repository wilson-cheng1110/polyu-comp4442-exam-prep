"""
Generate pastpapers_qa.docx — COMP4442 past-paper questions + model answers.
Printable, open-book reference. Three papers: 2023, 2024, 2025.
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_cell_bg(cell, color_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), color_hex)
    tc_pr.append(shd)


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    return p


def add_q(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(11)
    return p


def add_a(doc, text):
    p = doc.add_paragraph(text)
    for r in p.runs:
        r.font.size = Pt(10.5)
    return p


def add_code(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9.5)
    # keep together
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_after = Pt(4)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(text, style='List Bullet')
    for r in p.runs:
        r.font.size = Pt(10.5)
    return p


doc = Document()

# Page margins — compact for printability
for section in doc.sections:
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)

# Default font
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(10.5)

# ===================== TITLE =====================
title = doc.add_heading('COMP4442 Service and Cloud Computing', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub = doc.add_paragraph('Past-Paper Questions & Model Answers (3 papers)')
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.runs[0].bold = True
sub.runs[0].font.size = Pt(13)
meta = doc.add_paragraph('Papers: 2023 (4 May) · 2024 (10 May) · 2025 (13 May) — 100 marks each, 2 hours, open book')
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.runs[0].italic = True

doc.add_paragraph()

# ===================== PAPER 1: 2023 =====================
add_heading(doc, 'PAPER 1 — 4 May 2023 (Session 2022/2023 Sem II)', level=1)
doc.add_paragraph('Time: 12:30–14:30. Total: 100 marks. Q1 intentionally omitted in question paper.').runs[0].italic = True

# ---------- Q2 ----------
add_heading(doc, 'Question 2 (15 marks) — Short Questions', level=2)

add_q(doc, 'Q2(1) What is the difference between MapReduce and Spark in Hadoop? (3 marks)')
add_a(doc,
      'MapReduce is a disk-based, two-stage (map→reduce) batch processing model: intermediate output is always '
      'written to HDFS, so each job pays heavy disk I/O. Spark is an in-memory, DAG-based engine: RDDs keep '
      'data in RAM across stages, supports many operators (map, filter, join, reduceByKey, groupBy…), and is '
      '10–100× faster for iterative and interactive workloads. MapReduce is limited to batch; Spark supports '
      'batch, SQL, streaming, ML (MLlib), and graph (GraphX) in one unified engine. MapReduce is part of '
      'Hadoop core; Spark is a separate engine that can run on YARN and read/write HDFS.')

add_q(doc, 'Q2(2) What is the Hadoop NameNode and what happens if it fails? (3 marks)')
add_a(doc,
      'The NameNode is the HDFS master daemon. It stores the filesystem namespace (directory tree) and all '
      'block-to-DataNode mappings entirely in memory; the on-disk state is fsimage (snapshot) + edits log '
      '(journal). It does NOT store the actual file data — DataNodes do. If it fails, the cluster loses '
      'all metadata lookups and HDFS becomes unavailable (single point of failure in Hadoop 1). Recovery: '
      '(i) Secondary NameNode periodically merges fsimage+edits so the edits log does not grow unbounded — '
      'it is a checkpointer, NOT a hot standby; (ii) in Hadoop 2 HA mode a Standby NameNode shares edits '
      'via a Quorum Journal Manager (or shared NFS) and ZooKeeper failover controllers promote it '
      'automatically, eliminating the SPOF.')

add_q(doc, 'Q2(3) What is the difference between a "transformation" and an "action" in Spark? (3 marks)')
add_a(doc,
      'A transformation produces a new RDD/DataFrame from an existing one and is LAZY — Spark only records '
      'the operation in the lineage DAG, nothing executes. Examples: map, filter, flatMap, reduceByKey, '
      'groupByKey, join, select, withColumn, groupBy. An action forces computation: it materialises a '
      'result and either returns it to the driver or writes to storage. Examples: collect, count, take, '
      'first, reduce, foreach, saveAsTextFile, show, write. Lazy evaluation lets Spark fuse pipelines and '
      'optimise the DAG (Catalyst for DataFrames) before executing.')

add_q(doc, 'Q2(4) What is the role of Spark Executors in a Spark application and how do they work? (3 marks)')
add_a(doc,
      'Executors are JVM processes launched on cluster worker nodes for the lifetime of a Spark '
      'application. They are requested by the Driver from the Cluster Manager (YARN/K8s/Standalone). '
      'Each Executor: (1) runs Tasks assigned by the Driver (one Task per RDD partition per Stage); '
      '(2) caches RDD/DataFrame partitions in its block manager so subsequent stages can reuse them '
      'without recomputation; (3) reports heartbeats and task metrics back to the driver. Executors '
      'process data in parallel — one application typically has dozens or hundreds of Executors, each '
      'with multiple CPU cores.')

add_q(doc, 'Q2(5) Please describe the relationship between Hadoop and Spark. (3 marks)')
add_a(doc,
      'Hadoop and Spark are complementary, not competitors. Hadoop provides: (i) HDFS — a fault-tolerant '
      'distributed filesystem, (ii) YARN — a cluster resource manager, (iii) MapReduce — a batch execution '
      'engine. Spark is primarily a faster execution engine that replaces MapReduce for most workloads: '
      'it reads/writes HDFS, can be scheduled by YARN (spark-submit --master yarn), and reuses the Hadoop '
      'storage/security ecosystem (HCatalog, Hive, Kerberos). In practice, companies keep Hadoop for '
      'storage (HDFS) and resource management (YARN) but run Spark as the processing engine. Spark can '
      'also run fully standalone or on Kubernetes without Hadoop.')

# ---------- Q3 ----------
add_heading(doc, 'Question 3 (15 marks) — Social Media Application on AWS', level=2)

add_q(doc, 'Q3(a) Draw the SOA design diagram of your social media application and explain how your design '
      'follows SOA principles. (8 marks)')
add_a(doc, 'SOA design — layered, loosely-coupled services communicating via well-defined contracts (REST/JSON '
      'over HTTPS). The key services:')
add_code(doc, """[ Web / Mobile / IoT clients ]
          |
          v   (HTTPS, JWT)
+-----------------------------+
|    API Gateway (BFF)        |   <- single entry, routing, auth, rate-limit
+-----------------------------+
          |
  +-------+-------+-------+--------+--------+
  v               v       v       v        v
[ Auth ]   [ User  ]  [ Post  ] [Media  ] [Feed  ]
[Service]  [Profile]  [Service] [Service] [Service]
  |          |          |         |         |
  v          v          v         v         v
[Cognito] [User DB]  [Post DB]  [S3]     [Cache
                                          /Index]""")
add_a(doc, 'SOA principles demonstrated:')
add_bullet(doc, 'Service loose coupling — each service owns its own data and API; others call it only through '
      'the published contract, never by touching its DB.')
add_bullet(doc, 'Service abstraction — clients see only the API Gateway; internal implementation details '
      '(DynamoDB vs RDS, Lambda vs Fargate) are hidden.')
add_bullet(doc, 'Service reusability — Auth service is used by every other service and by external partners.')
add_bullet(doc, 'Service autonomy / statelessness — each service is independently deployable and horizontally '
      'scalable; session state lives in JWT or a shared cache (Redis/DynamoDB), not in the service.')
add_bullet(doc, 'Service composability — the Feed service orchestrates User + Post + Media to build a timeline; '
      'orchestration is explicit, not hidden in shared code.')
add_bullet(doc, 'Service discoverability — services registered behind API Gateway / service mesh (App Mesh) '
      'with OpenAPI/Swagger definitions.')

add_q(doc, 'Q3(b) Describe the AWS services you would like to use and explain how these services can be '
      'orchestrated to implement the functions. (7 marks)')
add_a(doc, 'Concrete AWS mapping:')
add_bullet(doc, 'Amazon Route 53 + CloudFront — DNS and CDN for the static front-end (S3-hosted SPA) so photos and '
      'video thumbnails are served close to the user.')
add_bullet(doc, 'Amazon S3 — durable object store for uploaded photos, videos, blogs. Pre-signed URLs let clients '
      'upload/download directly without going through the API, saving bandwidth.')
add_bullet(doc, 'Amazon Cognito — user sign-up, sign-in, federated identity (Google / Facebook), issues JWT '
      'tokens used by every other service.')
add_bullet(doc, 'Amazon API Gateway — REST/HTTP front door for all read/write operations; validates JWT via a '
      'Cognito authoriser; routes to Lambda.')
add_bullet(doc, 'AWS Lambda (Python) — stateless compute for each micro-service (post_create, feed_fetch, '
      'comment_add…). Auto-scales from 0 to thousands of concurrent invocations.')
add_bullet(doc, 'Amazon DynamoDB — NoSQL store for users, posts, comments. PAY_PER_REQUEST billing; GSIs for '
      'querying by hashtag / user; TTL for ephemeral data (stories).')
add_bullet(doc, 'Amazon SNS / SQS — fan-out notifications (new post → fan-out to followers\' feeds) and decoupling '
      '(like/comment events → analytics pipeline).')
add_bullet(doc, 'Amazon Kinesis / MSK — event stream for the activity log, consumed by analytics / ML services.')
add_bullet(doc, 'Amazon Rekognition / Comprehend — detect NSFW images and moderate text; triggered by S3 '
      'ObjectCreated events.')
add_bullet(doc, 'Amazon CloudWatch + X-Ray — logs, metrics, distributed tracing across services.')
add_a(doc, 'Orchestration flow (post creation): Client → CloudFront → API Gateway → Cognito JWT check → '
      'Post-Create Lambda → S3 presigned URL for media + DynamoDB write → SNS "new_post" → Fan-out Lambda → '
      'DynamoDB feeds + Rekognition async moderation → CloudWatch.')

# ---------- Q4 ----------
add_heading(doc, 'Question 4 (20 marks) — Hypervisors and OS-Level Virtualization', level=2)

add_q(doc, 'Q4(a) According to the description, which type of virtualization does Xen belong to, Type-1 or '
      'Type-2? (4 marks)')
add_a(doc, 'Xen is a TYPE-1 (bare-metal / native) hypervisor. Justification: the description explicitly states '
      '"runs directly on the host machine\'s hardware" — a Type-1 hypervisor sits directly on the physical '
      'hardware with no underlying host OS; the hypervisor itself is the kernel. Guest OSes (DomUs) run on '
      'top of it and a privileged control domain (Dom0) handles device drivers and management. By contrast, '
      'Type-2 hypervisors (VMware Workstation, VirtualBox) run as an application on top of a host OS. So '
      'Xen = Type-1.')

add_q(doc, 'Q4(b) Please provide the main difference between Type-1 and Type-2 hypervisors, along with their '
      'strengths and weaknesses. (8 marks)')

# Table for T1 vs T2
t = doc.add_table(rows=5, cols=3)
t.style = 'Light Grid Accent 1'
hdr = t.rows[0].cells
hdr[0].text = 'Aspect'
hdr[1].text = 'Type-1 (bare-metal)'
hdr[2].text = 'Type-2 (hosted)'
rows = [
    ('Runs on', 'Directly on hardware (the hypervisor IS the kernel)',
     'On top of a host OS as an application'),
    ('Examples', 'Xen, VMware ESXi, Microsoft Hyper-V, KVM*', 'VMware Workstation/Player, VirtualBox, Parallels'),
    ('Strengths', 'Lower overhead; better performance; better isolation & security (smaller TCB); scalable to '
                  'hundreds of VMs; used by public clouds.',
     'Easy to install on any desktop; coexists with host user apps; great for dev/test/labs; no reboot '
     'required to start.'),
    ('Weaknesses', 'Needs dedicated hardware; harder to install; requires driver support for the hypervisor; '
                   'management complexity.',
     'Extra layer of host OS — worse performance; host OS crash brings down all VMs; host OS consumes CPU/RAM; '
     'not suitable for datacenter-scale.'),
]
for i, (a, b, c) in enumerate(rows, start=1):
    t.rows[i].cells[0].text = a
    t.rows[i].cells[1].text = b
    t.rows[i].cells[2].text = c
add_a(doc, '*KVM is strictly a Linux kernel module, so it is sometimes called "Type-1.5" or hybrid.')

add_q(doc, 'Q4(c) Please explain what is "OS-Level" virtualization, along with its advantage compared to other '
      'virtualization techniques. (8 marks)')
add_a(doc, 'OS-level virtualization (a.k.a. containerization) partitions a single Linux kernel into multiple '
      'isolated user-space instances called CONTAINERS. Each container sees its own filesystem, process '
      'tree, network stack, UIDs and hostname, but they all SHARE the host kernel. The isolation is built '
      'from two Linux kernel features:')
add_bullet(doc, 'Namespaces — give each container its own view of system resources: pid (processes), net '
      '(network interfaces), mnt (mounts), uts (hostname), ipc (System V IPC), user (UID/GID), cgroup.')
add_bullet(doc, 'cgroups (control groups) — enforce resource limits and accounting on CPU, memory, block I/O, '
      'network bandwidth, device access.')
add_a(doc, 'Popular implementations: Docker, LXC, Podman, containerd, CRI-O. Kubernetes orchestrates them.')
add_a(doc, 'Advantages compared to hypervisor-based VMs:')
add_bullet(doc, 'Much lighter — a container starts in <1 s and typically uses tens of MB; a VM starts in '
      '30–60 s and uses GBs because it boots a full guest OS.')
add_bullet(doc, 'Higher density — hundreds of containers on one host vs dozens of VMs.')
add_bullet(doc, 'Near-native performance — no hardware emulation, no guest kernel.')
add_bullet(doc, 'Portable images — Docker image bundles app + deps; "build once, run anywhere".')
add_bullet(doc, 'Better suited to microservices and CI/CD pipelines.')
add_a(doc, 'Trade-offs: weaker isolation than VMs (shared kernel — a kernel exploit escapes all containers); '
      'cannot run a different OS family (no Windows containers on Linux kernel); not all kernel features '
      'are namespaced.')

# ---------- Q5 ----------
add_heading(doc, 'Question 5 (20 marks) — Healthcare IoT on AWS', level=2)

add_q(doc, 'Q5(a) Analyze what AWS services can be used and how the services can be orchestrated to '
      'implement such an application. Draw a diagram and provide a description. (10 marks)')
add_a(doc, 'Architecture diagram (data flow left→right):')
add_code(doc, """[Wearables / BP / ECG / Pulse-Ox sensors]
          |  (MQTT over TLS, X.509 mutual auth)
          v
  +-----------------------+
  |   AWS IoT Core        |    <- device registry, MQTT broker, rules engine
  +-----------------------+
     |              |              |
     | Rule 1       | Rule 2       | Rule 3
     v              v              v
 [Kinesis      [DynamoDB       [SNS alert
  Data          device_state]   to clinician]
  Firehose]
     |
     v
 [S3 "raw/" ]  -->  [Glue/Athena analytics]
     |
     v
 [AWS Lambda (featurise)] -> [DynamoDB "patient_dashboard"]
     |
     v
 [Amazon QuickSight / Mobile app (via API Gateway + Cognito)]""")
add_bullet(doc, 'AWS IoT Core — secure MQTT ingestion with per-device X.509 certificates; Device Shadow stores '
      'the last-known state even when device is offline.')
add_bullet(doc, 'IoT Rules Engine — SQL-like rules route telemetry to downstream services: high-priority vitals '
      '→ SNS alert; time-series → Kinesis Firehose; snapshots → DynamoDB.')
add_bullet(doc, 'Amazon Kinesis Data Firehose → S3 — durable "raw" data lake partitioned by date, used later '
      'for analytics and ML training.')
add_bullet(doc, 'Amazon DynamoDB — low-latency store for the patient\'s current vitals dashboard; TTL on '
      'ephemeral records.')
add_bullet(doc, 'Amazon SNS / SES — push notifications and emails to clinicians; also integrates with '
      'on-call paging (PagerDuty).')
add_bullet(doc, 'AWS Lambda — stateless functions to featurise streams, derive anomaly flags, and write to '
      'patient dashboard tables.')
add_bullet(doc, 'Amazon Cognito + API Gateway — authenticates patients/clinicians, exposes REST API for the '
      'mobile app.')
add_bullet(doc, 'Amazon CloudWatch — health metrics on the pipeline; alarms on lag and error rates.')
add_bullet(doc, 'AWS KMS / Secrets Manager — encryption keys and any API keys; HIPAA-compliant configuration.')
add_a(doc, 'Orchestration: device publishes telemetry → IoT Core authenticates and applies rule → rule forks '
      'data to Firehose (S3) + DynamoDB + optional SNS alert → Lambda featurises → patient app queries '
      'via API Gateway.')

add_q(doc, 'Q5(b) Suppose you want to extend the system so that it can use a machine-learning algorithm to '
      'analyze your health status and provide a report. Explain how the designed system in (a) can be '
      'adapted. Draw the new system architecture and provide descriptions. (10 marks)')
add_a(doc, 'Add an ML plane alongside the operational plane from (a). New diagram:')
add_code(doc, """[Sensors]--MQTT-->[IoT Core]--Firehose-->[S3 raw/]
                                         |
                                         v
                 +------ training pipeline (offline, batch) ------+
                 | [S3 raw/] -> [AWS Glue ETL (PySpark)]          |
                 |   -> [S3 processed/feature_store/]             |
                 |   -> [SageMaker Training: XGBoost /            |
                 |       anomaly-detection / LSTM]                |
                 |   -> [S3 models/]                              |
                 +-------------------------------------------------+
                                         |
                                         v   (register & deploy)
                              +----------------------+
                              | SageMaker Model      |
                              | Registry             |
                              +----------------------+
                                         |
                                         v
                              +----------------------+
                              | SageMaker endpoint   |   <- real-time inference
                              | (or serverless)      |
                              +----------------------+
                                         ^
  [IoT Core rule / Lambda "scorer"] -----+ (invokes endpoint on each new sample)
                 |
                 v
  [DynamoDB risk_score] + [SNS alert if risk>threshold]
                 |
                 v
  [API Gateway GET /report] -> [Lambda] -> [PDF via S3 presigned URL]""")
add_a(doc, 'Adaptations made to the original design:')
add_bullet(doc, 'Add AWS Glue ETL or EMR (Spark) to turn the S3 "raw" lake into a feature store (per-patient '
      'windowed statistics: rolling mean HR, HRV, SpO2 deltas, sleep metrics).')
add_bullet(doc, 'Add Amazon SageMaker: Jupyter notebooks for experimentation → Training jobs → Model Registry '
      '→ Endpoint (real-time) or Batch Transform (nightly reports).')
add_bullet(doc, 'Add an inference Lambda ("scorer") triggered by IoT Rule: pulls last N samples from DynamoDB, '
      'invokes SageMaker endpoint, writes risk_score back.')
add_bullet(doc, 'Add a Report-Generation Lambda: on demand or scheduled via EventBridge, runs SageMaker '
      'Batch Transform over the patient\'s history, templates a PDF with matplotlib/reportlab, stores '
      'in S3, surfaces a pre-signed download URL via API Gateway.')
add_bullet(doc, 'MLOps: CodePipeline for model CI/CD, SageMaker Model Monitor for drift detection, '
      'CloudWatch metrics on inference latency and accuracy.')
add_bullet(doc, 'Privacy/compliance: enable KMS-CMK for S3/DynamoDB; VPC endpoints for SageMaker; IAM '
      'least-privilege for the scorer Lambda.')
add_a(doc, 'Because the original architecture is loosely coupled via S3 and IoT rules, the ML plane attaches '
      'without modifying the operational plane — new services are added, none removed.')

# ---------- Q6 ----------
add_heading(doc, 'Question 6 (20 marks) — PySpark RDD: MaxTempByCity', level=2)
add_a(doc, 'Given code:').runs[0].italic = True
add_code(doc, """from pyspark import SparkContext
sc = SparkContext("local", "MaxTempByCity")
data = sc.textFile("path/to/temperature_data.csv").map(lambda line: line.split(","))
data = data.filter(lambda x: x[1] != "null")
data = data.map(lambda x: (x[0], float(x[1])))
max_temps = data.reduceByKey(lambda x, y: max(x, y))
for city, max_temp in max_temps.collect():
    print(f"{city}: {max_temp}")
sc.stop()""")

add_q(doc, 'Q6(a) Please provide a description for each line of the code based on Spark and RDD knowledge. (6 marks)')
add_a(doc, 'Line-by-line:')
add_bullet(doc, '`from pyspark import SparkContext` — import the SparkContext class, the entry point for '
      'RDD-based PySpark jobs.')
add_bullet(doc, '`sc = SparkContext("local", "MaxTempByCity")` — create a SparkContext in local mode '
      '(runs on the driver machine; all cores), application name "MaxTempByCity".')
add_bullet(doc, '`data = sc.textFile(...).map(lambda line: line.split(","))` — read the CSV as an RDD of '
      'strings (one per line) and TRANSFORM each line into a list of fields split on commas. Lazy.')
add_bullet(doc, '`data = data.filter(lambda x: x[1] != "null")` — TRANSFORMATION: drop rows whose Date '
      'field is "null". NOTE: in the given data there is no null Date; if the intent was to drop '
      'missing temperatures the filter should test x[2], so this line is a subtle bug — the '
      'examiner may want you to spot that date and temperature are both checked against position 1.')
add_bullet(doc, '`data = data.map(lambda x: (x[0], float(x[1])))` — TRANSFORMATION: turn each row into a '
      'key-value pair (city, temperature). Again note it uses x[1] where the real CSV has temperature '
      'in x[2] (columns: City, Date, Temperature). On the stated dataset City,Date,Temperature this '
      'would compare dates, not temperatures — mention this in the answer.')
add_bullet(doc, '`max_temps = data.reduceByKey(lambda x, y: max(x, y))` — TRANSFORMATION: for each key '
      '(city) combine all values pairwise with max(), producing the maximum temperature per city. '
      'reduceByKey uses a map-side combiner so traffic across the network is minimised.')
add_bullet(doc, '`for city, max_temp in max_temps.collect():` — ACTION: collect() brings the pair-RDD back '
      'to the driver as a Python list; we iterate and print.')
add_bullet(doc, '`sc.stop()` — gracefully stop the SparkContext, releasing cluster resources.')

add_q(doc, 'Q6(b) What is the difference between `map` and `reduceByKey` in Spark RDDs? How are they used '
      'in the context of finding the maximum temperature for each city? (7 marks)')
add_a(doc, '`map(f)` — element-wise TRANSFORMATION. Takes an RDD[T] and a function f: T→U; produces an '
      'RDD[U] with the same number of elements (one output per input). It is narrow — no shuffle; each '
      'partition is transformed independently. In the code, map turns lines into tuples (city, temp).')
add_a(doc, '`reduceByKey(f)` — pair-RDD AGGREGATION. Takes an RDD[(K, V)] and a COMMUTATIVE + ASSOCIATIVE '
      'function f: (V, V) → V; produces RDD[(K, V)] with one entry per key. It is WIDE — requires a '
      'shuffle so values for the same key land on the same partition, but it runs a map-side combine '
      'first, greatly reducing network traffic.')
add_a(doc, 'In this job: `map` prepares (city, temperature) pairs; `reduceByKey(max)` then folds all '
      'temperatures for each city pairwise with the max function, giving one maximum per city. '
      'Map does not aggregate; reduceByKey does. Together they express the SQL `SELECT city, MAX(temp) '
      'FROM t GROUP BY city`.')
add_a(doc, 'Aside — for this problem `reduceByKey(max)` is better than `groupByKey().mapValues(max)` '
      'because the latter shuffles every value over the network; reduceByKey pre-aggregates on each '
      'partition.')

add_q(doc, 'Q6(c) What are some advantages and disadvantages of using RDDs compared to DataFrames? When might '
      'you choose RDDs instead? (7 marks)')
add_a(doc, 'Advantages of RDDs:')
add_bullet(doc, 'Full control over execution — you write the exact map/filter/reduce, no optimiser second-guesses you.')
add_bullet(doc, 'Work naturally with unstructured / arbitrary Python objects — images, bytes, graph nodes.')
add_bullet(doc, 'Low-level transformations available (mapPartitions, zip, sample, glom) that DataFrame API '
      'does not expose cleanly.')
add_bullet(doc, 'Compile-time type safety in Scala/Java; easier when data has no fixed schema.')
add_a(doc, 'Disadvantages of RDDs:')
add_bullet(doc, 'No Catalyst optimiser or Tungsten physical optimisations — slower than DataFrames for '
      'relational queries.')
add_bullet(doc, 'No schema → no column-level optimisations (predicate pushdown, projection pruning, '
      'partition elimination).')
add_bullet(doc, 'All data serialised as Python objects in PySpark — higher memory and GC overhead than '
      'DataFrames which use off-heap Tungsten binary format.')
add_bullet(doc, 'More code / lower productivity — you write lambdas instead of declarative SQL.')
add_a(doc, 'When to choose RDD:')
add_bullet(doc, 'Data has no schema (free text, raw bytes, images) or schema varies per record.')
add_bullet(doc, 'You need fine-grained control over partitioning, ordering, or custom serialisation.')
add_bullet(doc, 'You implement a low-level algorithm (e.g. custom graph traversal, custom ML iteration) '
      'that does not map to SQL operators.')
add_bullet(doc, 'Integration with a legacy RDD-only library.')
add_a(doc, 'For the temperature data — which has a fixed schema (City, Date, Temperature) — a DataFrame '
      'would be strictly better: `df.groupBy("city").agg(max("temperature"))`. Use RDD here only if you '
      'deliberately want practice with low-level APIs.')

doc.add_page_break()

# ===================== PAPER 2: 2024 =====================
add_heading(doc, 'PAPER 2 — 10 May 2024 (Session 2023/2024 Sem II)', level=1)
doc.add_paragraph('Time: 12:30–14:30. Total: 100 marks. Q1 intentionally omitted.').runs[0].italic = True

# ---------- Q2 ----------
add_heading(doc, 'Question 2 (15 marks) — Short Questions', level=2)

add_q(doc, 'Q2(1) What is virtualization technique? (3 marks)')
add_a(doc, 'Virtualization is the abstraction of physical computing resources (CPU, memory, disk, network) '
      'into logical resources so multiple virtual workloads can share a single physical machine while '
      'remaining isolated. The abstraction layer may be a hypervisor (hardware virtualization — VMs), '
      'the operating-system kernel (OS-level virtualization — containers), the language runtime (JVM / '
      'CLR), or the network stack (VLANs, overlays, VXLAN). Benefits: higher utilisation, isolation '
      'between tenants, easier provisioning and migration, and reduced cost. Virtualization is the '
      'foundation of modern cloud computing.')

add_q(doc, 'Q2(2) List 2 virtualization techniques and describe their features (e.g. Xen, Docker). (3 marks)')
add_bullet(doc, 'XEN — hypervisor-based (Type-1 bare-metal) virtualization. Runs directly on hardware; '
      'supports paravirtualization (PV) and hardware-assisted full virtualization (HVM); uses a '
      'privileged Dom0 for device drivers and unprivileged DomU guests. Strong isolation, can run '
      'different OS kernels (Linux, Windows, BSD); heavier — each VM boots its own OS. Used by AWS EC2 '
      '(classic) and Citrix XenServer.')
add_bullet(doc, 'DOCKER — OS-level (containerization) virtualization. Shares one Linux kernel across '
      'containers; uses cgroups for resource limits and namespaces (pid/net/mnt/uts/ipc/user) for '
      'isolation. Packages applications as layered images distributed via registries; starts in '
      'milliseconds and uses megabytes rather than gigabytes. Cannot run a different OS kernel but '
      'offers near-native performance and is the basis of modern microservices.')

add_q(doc, 'Q2(3) List 2 Data Center Network architectures and describe their features (e.g. Fat-Tree, DCell). (3 marks)')
add_bullet(doc, 'FAT-TREE — a switch-centric multi-rooted tree built from identical k-port commodity '
      'switches arranged in three layers: edge, aggregation, core. With k-port switches it supports '
      'k³/4 servers, is non-blocking, and provides full bisection bandwidth using equal-cost multi-path '
      '(ECMP) routing. Advantages: uniform, easy to build from cheap switches, high bisection, '
      'multi-path fault tolerance. Weakness: wiring complexity grows with k; core layer is a potential '
      'hotspot; scaling requires replacing switches when k changes.')
add_bullet(doc, 'DCell — a server-centric recursive architecture (DCell₀ has n servers connected to one '
      'mini-switch; DCell_k is built by recursively combining DCell_{k-1} cells, each server using an '
      'extra NIC per level). With n=6 ports, DCell can scale to millions of servers with only 4 levels. '
      'Servers forward traffic (multi-port NIC + software routing). Advantages: extreme scalability, '
      'high fault tolerance (many disjoint paths), low cost. Weakness: servers do packet forwarding → '
      'CPU overhead and software complexity; longer paths than Fat-Tree.')

add_q(doc, 'Q2(4) Please describe the relationship between Hadoop and Spark. (3 marks)')
add_a(doc, '(See Paper 1 Q2(5) — same answer.) Hadoop supplies storage (HDFS) and resource management '
      '(YARN); Spark is a faster in-memory execution engine that can replace MapReduce. Spark reads and '
      'writes HDFS, can be scheduled by YARN, and reuses Hadoop ecosystem (Hive, Kerberos, HCatalog). '
      'They are complementary — Hadoop for data, Spark for compute. Spark can also run standalone or on '
      'Kubernetes without Hadoop.')

add_q(doc, 'Q2(5) What is the difference between a transformation and an action in Spark? (3 marks)')
add_a(doc, '(See Paper 1 Q2(3).) Transformations are LAZY — they build the lineage DAG (map, filter, '
      'groupBy, join, reduceByKey, select, withColumn…). Actions trigger execution and return a value to '
      'the driver or write output (count, collect, reduce, take, first, foreach, show, saveAsTextFile, '
      'write). Lazy evaluation lets Spark optimise the whole pipeline (Catalyst) before running.')

# ---------- Q3 ----------
add_heading(doc, 'Question 3 (15 marks) — Network Virtualization (Xen)', level=2)

add_q(doc, 'Q3(a) For network virtualization in the Xen virtualization technique, describe how it works '
      '(hint: network interfaces, etc.). (8 marks)')
add_a(doc, 'Xen implements network virtualization via a split driver model between Domain 0 (Dom0 — the '
      'privileged control domain) and Domain U (DomU — unprivileged guest VMs):')
add_bullet(doc, 'The physical NIC is owned by Dom0, which runs the real network driver. Dom0 sees it as '
      '`eth0`.')
add_bullet(doc, 'Dom0 runs a software bridge, typically `xenbr0`. The physical NIC `eth0` (after being '
      'renamed to `peth0` when bridged) becomes a port on `xenbr0`.')
add_bullet(doc, 'Each guest VM (DomU) gets a paravirtualised front-end network driver exposing an '
      'interface the guest calls `eth0` internally.')
add_bullet(doc, 'Corresponding to each guest\'s `eth0` there is a back-end virtual interface in Dom0 '
      'called `vifX.Y` (X = DomU id, Y = virtual NIC index). These `vifX.Y` interfaces are also plugged '
      'into the bridge `xenbr0`.')
add_bullet(doc, 'Packets travel over shared-memory ring buffers (event channels + grant tables) between '
      'front-end and back-end — a Xen-specific zero-copy mechanism.')
add_bullet(doc, 'The bridge `xenbr0` is a layer-2 software switch — it learns MACs, forwards frames by '
      'destination MAC, and broadcasts ARP. All VMs on the same bridge see each other directly.')
add_bullet(doc, 'Because xenbr0 is connected to the physical NIC, VMs can also talk to the outside world. '
      'For isolation you can build separate bridges (xenbr1, xenbr2) and connect different vifs.')
add_a(doc, 'In hardware-assisted (HVM) mode Xen can also emulate a real NIC (QEMU-emulated e1000/rtl8139), '
      'but PV drivers are preferred for performance.')

add_q(doc, 'Q3(b) Figure out the data flow when VM A wants to send data to VM B in the following figure. (7 marks)')
add_a(doc, 'Given the figure (VM A with eth0, Dom0 with xenbr0 + vif1.0 and vif2.0, VM B with eth0):')
add_code(doc, """VM A application writes to socket
      |
      v
VM A kernel TCP/IP stack produces Ethernet frame (dst MAC = VM B)
      |
      v
VM A eth0 (Xen PV front-end) — frame placed in shared-memory tx ring; event channel notifies Dom0
      |
      v
Dom0 back-end picks up frame at vif1.0 (the back-end paired with VM A's eth0)
      |
      v
vif1.0 hands the frame to the bridge xenbr0
      |
      v
xenbr0 looks up destination MAC in its forwarding table; forwards to vif2.0
      |
      v
vif2.0 places frame in shared-memory rx ring of VM B's back-end; notifies VM B
      |
      v
VM B eth0 (Xen PV front-end) receives frame
      |
      v
VM B kernel TCP/IP stack -> VM B application""")
add_a(doc, 'Key point: the frame never leaves the host — Xen bridges VM A and VM B entirely in software '
      'through Dom0. If VM B were on a different host, xenbr0 would forward the frame out peth0 onto '
      'the physical wire.')

# ---------- Q4 Hypervisor (same as Paper 1) ----------
add_heading(doc, 'Question 4 (20 marks) — Hypervisor and OS-Level Virtualization', level=2)
add_a(doc, '(Identical to Paper 1 Question 4 — reuse the same answer.)').runs[0].italic = True

add_q(doc, 'Q4(a) Xen Type-1 or Type-2? (4 marks)')
add_a(doc, 'Xen is Type-1 (bare-metal / native) — it runs directly on hardware with no underlying host OS. '
      'See Paper 1 Q4(a).')

add_q(doc, 'Q4(b) Main difference between Type-1 and Type-2 hypervisors + strengths/weaknesses. (8 marks)')
add_a(doc, '(See Paper 1 Q4(b) — table of Aspect / Type-1 / Type-2.)')

add_q(doc, 'Q4(c) What is OS-Level virtualization and its advantages? (8 marks)')
add_a(doc, '(See Paper 1 Q4(c).) Shared-kernel partitioning via namespaces + cgroups; containers (Docker, '
      'LXC). Lighter, faster startup, higher density, portable images, near-native performance. Weaker '
      'isolation than VMs; single-OS-family limit.')

# ---------- Q5 ----------
add_heading(doc, 'Question 5 (20 marks) — Kubernetes', level=2)

add_q(doc, 'Q5(a) For Docker, describe what Cgroups and NameSpaces are. (5 marks)')
add_a(doc, 'Cgroups (Control Groups) — a Linux kernel feature that groups processes and enforces and '
      'accounts for resource usage on each group. Subsystems: cpu (shares + CFS quota/period), cpuset '
      '(CPU affinity), memory (limit + OOM), blkio (disk IOPS/bandwidth), net_cls/net_prio (network '
      'tagging), devices (allow/deny), pids (max process count), freezer. Docker uses cgroups to enforce '
      '`--cpus`, `--memory`, `--pids-limit` etc.')
add_a(doc, 'NameSpaces — a Linux kernel feature that gives each process its own VIEW of a system resource, '
      'providing isolation. Docker creates these namespaces per container:')
add_bullet(doc, 'PID — separate process tree; container sees its own PID 1.')
add_bullet(doc, 'NET — separate network stack, interfaces, routing table, iptables.')
add_bullet(doc, 'MNT — separate mount table, so the container has its own root FS.')
add_bullet(doc, 'UTS — separate hostname & domain name.')
add_bullet(doc, 'IPC — separate System V IPC (semaphores, shared memory, message queues).')
add_bullet(doc, 'USER — separate UID/GID mapping; container root ≠ host root.')
add_bullet(doc, 'CGROUP — hides the cgroup hierarchy from the container.')
add_a(doc, 'Together: cgroups = "how much" (resource limits), namespaces = "what you can see" (isolation).')

add_q(doc, 'Q5(b) Explain the architecture of Kubernetes. (7 marks)')
add_a(doc, 'Kubernetes (K8s) is a cluster of machines split into two planes:')
add_bullet(doc, 'CONTROL PLANE (Master) — brains:')
add_bullet(doc, '  kube-apiserver — the single front-door REST API (all components talk to it); '
      'validates and persists desired state.')
add_bullet(doc, '  etcd — a distributed, consistent KV store holding the authoritative cluster state '
      '(desired state, pod specs, secrets); Raft consensus.')
add_bullet(doc, '  kube-scheduler — watches unscheduled pods, picks the best node based on resource '
      'requests, affinity/anti-affinity, taints/tolerations, topology spread.')
add_bullet(doc, '  kube-controller-manager — runs reconciliation loops (ReplicaSet, Deployment, '
      'Node, Endpoint, Namespace, ServiceAccount, …) that continuously drive actual → desired state.')
add_bullet(doc, '  cloud-controller-manager (optional) — integrates with cloud LBs, nodes, volumes.')
add_bullet(doc, 'DATA PLANE (Worker Nodes) — brawn:')
add_bullet(doc, '  kubelet — node agent; pulls pod specs from the API server and tells the container '
      'runtime to create/start containers; reports health back.')
add_bullet(doc, '  container runtime — containerd / CRI-O / Docker-shim; actually runs containers.')
add_bullet(doc, '  kube-proxy — programs iptables/IPVS rules to implement Services (virtual IP → pod IPs).')
add_bullet(doc, '  CNI plugin (Flannel, Calico, Weave) — pod networking.')
add_a(doc, 'Communication flow: user → kubectl → kube-apiserver → etcd. Controllers watch the apiserver '
      '(list/watch) and make the cluster match desired state. kubelets run on every node.')

add_q(doc, 'Q5(c) Describe the main workflow of Kubernetes and draw a figure to show it. (8 marks)')
add_a(doc, 'Kubernetes workflow for deploying a pod (kubectl apply → running container):')
add_code(doc, """  [User]
     |  1. kubectl apply -f deployment.yaml
     v
  [kube-apiserver] ---2. validate & write desired state--> [etcd]
     ^                                                   |
     |                                   3. controllers watch apiserver
     |                                                   v
     |            [Deployment controller] -> creates ReplicaSet
     |            [ReplicaSet controller] -> creates Pod objects (status Pending)
     |
     |  4. scheduler watches Pending pods
     |            [kube-scheduler] -> picks a Node -> writes nodeName back to apiserver
     |
     |  5. target node's kubelet watches apiserver for pods assigned to it
     v
  [kubelet @ Node]
     |  6. pulls pod spec from apiserver
     |  7. invokes CRI (containerd) to pull image from registry
     |  8. starts container via runc; sets up cgroups + namespaces
     |  9. CNI plugin wires pod network
     |  10. kubelet runs liveness/readiness probes
     |  11. kubelet updates pod status to Running -> apiserver
     v
  [kube-proxy] programs iptables/IPVS for any Service that selects this pod

  [User] kubectl get pods -> apiserver -> returns "Running" status""")
add_a(doc, 'Key reconciliation loop: every controller continuously compares desired state (in etcd) with '
      'actual state (reported by kubelets) and issues corrective actions. Self-healing follows '
      'automatically — if a pod crashes, the ReplicaSet controller notices replica count dropped and '
      'creates a replacement pod.')

# ---------- Q6 ----------
add_heading(doc, 'Question 6 (20 marks) — PySpark DataFrame: Sales Data', level=2)
add_a(doc, 'Given code:').runs[0].italic = True
add_code(doc, """from pyspark.sql import SparkSession
from pyspark.sql.functions import col, sum, to_date

spark = SparkSession.builder.appName("Sales Data Analysis").getOrCreate()

sales_df = spark.read.csv("sales_data.csv", header=True, inferSchema=True)

sales_df = sales_df.withColumn("date", to_date(col("date"), "yyyy-MM-dd"))

sales_2024_df = sales_df.filter(col("date").between("2024-01-01", "2024-4-18"))
revenue_2024_df = sales_2024_df.groupBy("product_id") \\
    .agg(sum(col("quantity") * col("price")).alias("total_revenue")) \\
    .orderBy(col("total_revenue").desc())

revenue_2024_df.show()

top_3_products_2024 = revenue_2024_df.limit(3)
top_3_products_2024.show()

spark.stop()""")

add_q(doc, 'Q6(a) Please provide a description for each line of code. (5 marks)')
add_bullet(doc, '`from pyspark.sql import SparkSession` — entry point for DataFrame API.')
add_bullet(doc, '`from pyspark.sql.functions import col, sum, to_date` — column ref, aggregate, date parser.')
add_bullet(doc, '`spark = SparkSession.builder.appName("Sales Data Analysis").getOrCreate()` — create or '
      'reuse a SparkSession named "Sales Data Analysis"; session owns underlying SparkContext.')
add_bullet(doc, '`sales_df = spark.read.csv("sales_data.csv", header=True, inferSchema=True)` — read CSV '
      'into DataFrame; first row = column names; types inferred by sampling.')
add_bullet(doc, '`sales_df = sales_df.withColumn("date", to_date(col("date"), "yyyy-MM-dd"))` — '
      'TRANSFORMATION: cast the `date` string column to a proper DateType using the format yyyy-MM-dd.')
add_bullet(doc, '`sales_2024_df = sales_df.filter(col("date").between("2024-01-01", "2024-4-18"))` — '
      'TRANSFORMATION: keep only rows in the range 1 Jan 2024 to 18 Apr 2024.')
add_bullet(doc, '`revenue_2024_df = ...groupBy("product_id").agg(sum(col("quantity")*col("price"))'
      '.alias("total_revenue")).orderBy(col("total_revenue").desc())` — group by product; per group sum '
      'the line total (quantity × price) aliased as `total_revenue`; sort descending. All '
      'TRANSFORMATIONS (lazy).')
add_bullet(doc, '`revenue_2024_df.show()` — ACTION: materialises up to 20 rows on the driver and prints.')
add_bullet(doc, '`top_3_products_2024 = revenue_2024_df.limit(3)` — TRANSFORMATION: keep only the first 3 rows.')
add_bullet(doc, '`top_3_products_2024.show()` — ACTION: prints the top-3 products.')
add_bullet(doc, '`spark.stop()` — terminate the SparkSession and release resources.')

add_q(doc, 'Q6(b) What is "Transformation" in Spark? Describe and give an example from the above code. (5 marks)')
add_a(doc, 'A Transformation is a lazy operation that defines a NEW Dataset (RDD/DataFrame) from an '
      'existing one. It returns another Dataset but does NOT execute — Spark only records it in the '
      'logical plan (lineage / DAG). Execution is triggered later by an Action. Transformations can be '
      'NARROW (each output partition depends on a single input partition — map, filter, select, '
      'withColumn, drop) or WIDE (requires a shuffle — groupBy, reduceByKey, join, distinct, orderBy).')
add_a(doc, 'Lazy evaluation enables pipeline optimisation — Catalyst fuses multiple transformations into '
      'one whole-stage code-gen execution and pushes predicates down to the source.')
add_a(doc, 'Examples from the code:')
add_bullet(doc, '`sales_df.withColumn("date", to_date(col("date"), "yyyy-MM-dd"))` — narrow.')
add_bullet(doc, '`sales_df.filter(col("date").between("2024-01-01","2024-4-18"))` — narrow.')
add_bullet(doc, '`sales_2024_df.groupBy("product_id").agg(sum(...).alias("total_revenue"))` — wide.')
add_bullet(doc, '`orderBy(col("total_revenue").desc())` — wide (shuffle to a single range).')
add_bullet(doc, '`limit(3)` — narrow per-partition then gather.')

add_q(doc, 'Q6(c) What is "Action" in Spark? Describe and give an example from the above code. (5 marks)')
add_a(doc, 'An Action is an operation that triggers execution of the lineage DAG of accumulated '
      'transformations and either returns a value to the Driver program or writes data to an external '
      'sink. Because Spark is lazy, NOTHING runs on the cluster until an Action is called. Actions '
      'submit a Spark Job — the scheduler splits the DAG into Stages (by shuffle boundaries), each '
      'Stage into Tasks (one per partition), and dispatches them to Executors.')
add_a(doc, 'Common actions: count, collect, take, first, reduce, foreach, show, toPandas, '
      'saveAsTextFile, write.parquet, write.format(...).save().')
add_a(doc, 'Examples from the code:')
add_bullet(doc, '`revenue_2024_df.show()` — prints the first 20 rows; each call triggers a job.')
add_bullet(doc, '`top_3_products_2024.show()` — prints the top-3 rows; because we called show() on a '
      'filtered descendant, the whole lineage (read → withColumn → filter → groupBy → agg → orderBy → '
      'limit → show) is executed.')

add_q(doc, 'Q6(d) What are the main components (architecture) of a Spark application? Explain each. (5 marks)')
add_bullet(doc, 'DRIVER PROGRAM — the JVM process that runs `main()` (the SparkSession). Holds the '
      'SparkContext, builds the DAG from transformations, converts it to stages+tasks, and schedules '
      'those to Executors. Also collects results (from `collect`, `show`). Single point per application '
      '— if it dies the app dies.')
add_bullet(doc, 'CLUSTER MANAGER — external component that hands out resources: Standalone, YARN, Mesos, '
      'Kubernetes. The Driver requests Executors ("containers") from the Cluster Manager.')
add_bullet(doc, 'WORKER NODE — machine in the cluster on which Executors are launched.')
add_bullet(doc, 'EXECUTOR — JVM process launched on a Worker; lives for the life of the application. '
      'Runs tasks in thread pools, caches partitions in its BlockManager, ships metrics back to Driver.')
add_bullet(doc, 'TASK — the smallest unit of work; one per partition per stage; assigned by Driver to '
      'Executor.')
add_bullet(doc, 'JOB & STAGE — an Action triggers a Job; the DAG Scheduler splits the Job into Stages '
      'by shuffle boundaries; each Stage has many Tasks run in parallel.')
add_a(doc, 'Catalyst + Tungsten (for DataFrame/SQL) sit inside the Driver: Catalyst builds and '
      'optimises the logical plan (predicate pushdown, constant folding, join reordering); Tungsten '
      'generates whole-stage bytecode and uses off-heap binary memory layout.')

doc.add_page_break()

# ===================== PAPER 3: 2025 =====================
add_heading(doc, 'PAPER 3 — 13 May 2025 (Session 2024/2025 Sem II)', level=1)
doc.add_paragraph('Time: 12:30–14:30. Total: 100 marks. Q1 intentionally omitted.').runs[0].italic = True

# ---------- Q2 ----------
add_heading(doc, 'Question 2 (15 marks) — Short Questions', level=2)

add_q(doc, 'Q2(1) What is the primary difference between IaaS and PaaS in cloud computing? (3 marks)')
add_a(doc, 'IaaS (Infrastructure as a Service) gives the customer virtual hardware — VMs, block storage, '
      'virtual networks — on which the customer installs OS, middleware, runtime, and application code. '
      'Examples: AWS EC2, Azure Virtual Machines, Google Compute Engine. The customer manages everything '
      'from the OS upward.')
add_a(doc, 'PaaS (Platform as a Service) gives the customer a managed application platform — runtime, '
      'database, load balancer, auto-scaling — so they only supply application code and configuration. '
      'Examples: AWS Elastic Beanstalk, Google App Engine, Heroku, Azure App Service. The provider '
      'manages OS patching, capacity, and most middleware.')
add_a(doc, 'Primary difference: IaaS exposes infrastructure primitives (more control, more responsibility); '
      'PaaS exposes an application runtime (less control, faster to ship).')

add_q(doc, 'Q2(2) What is the purpose of a hypervisor in cloud computing? (3 marks)')
add_a(doc, 'A hypervisor (Virtual Machine Monitor, VMM) is the software layer that lets multiple isolated '
      'virtual machines share one physical host. Its purposes:')
add_bullet(doc, 'CPU / memory / device virtualization — emulates or arbitrates access to hardware.')
add_bullet(doc, 'Isolation — each guest OS is sandboxed; a crash or compromise in one guest cannot reach '
      'another.')
add_bullet(doc, 'Resource multiplexing — runs many VMs on one host, improving utilisation.')
add_bullet(doc, 'Live migration, snapshots, cloning — enables elasticity, HA, and disaster recovery.')
add_bullet(doc, 'Security boundary — provides the strong isolation needed for multi-tenant cloud.')
add_a(doc, 'Two flavours: Type-1 (bare-metal — Xen, ESXi, Hyper-V, KVM) runs directly on hardware; '
      'Type-2 (hosted — VMware Workstation, VirtualBox) runs as an app on a host OS. Public clouds '
      'rely on Type-1 hypervisors.')

add_q(doc, 'Q2(3) Introduce the main components of Kubernetes and their roles and functions. (3 marks)')
add_a(doc, 'Control plane (master) — kube-apiserver (REST front-door), etcd (consistent state store), '
      'kube-scheduler (picks a node for each pod), kube-controller-manager (runs reconciliation loops '
      'for Deployment/ReplicaSet/Node/etc.), cloud-controller-manager (cloud integration).')
add_a(doc, 'Data plane (worker node) — kubelet (node agent, starts pods via CRI), container runtime '
      '(containerd/CRI-O), kube-proxy (Service networking via iptables/IPVS), CNI plugin (pod '
      'networking).')
add_a(doc, 'Objects — Pod (smallest deployable unit, 1+ containers sharing netns/volumes), Deployment '
      '(rolling-update wrapper over ReplicaSet), Service (stable virtual IP / DNS for pods), '
      'ConfigMap / Secret (config and sensitive data), Ingress (L7 HTTP routing), PersistentVolume / '
      'Claim (storage).')

add_q(doc, 'Q2(4) Please describe the relationship between Hadoop and Spark. (3 marks)')
add_a(doc, '(See Paper 1 Q2(5) and Paper 2 Q2(4) — same answer.)')

add_q(doc, 'Q2(5) What is the difference between a transformation and an action in Spark? (3 marks)')
add_a(doc, '(See Paper 1 Q2(3) and Paper 2 Q2(5) — same answer.)')

# ---------- Q3 ----------
add_heading(doc, 'Question 3 (15 marks) — Docker / Kubernetes Resource Scheduling', level=2)

add_q(doc, 'Q3(a) Given the figure showing CPU/Memory requests & usage by Pods A, B, C on a node, can Pod '
      'D be deployed? Explain in detail. (8 marks)')
add_a(doc, 'Kubernetes scheduling is based on RESOURCE REQUESTS, not current usage. For each resource '
      '(CPU, memory), the scheduler computes:')
add_code(doc, """available_for_new_pod = node_capacity - sum(requests of pods already on node)""")
add_a(doc, 'Reading the figure:')
add_bullet(doc, 'CPU REQUESTS row: Pods A+B+C reserve most of the bar; a small "Unallocated" slice remains.')
add_bullet(doc, 'CPU USAGE row: Pods A+B+C use less than their requests — a "Currently Unused" portion '
      'of reserved CPU is idle.')
add_bullet(doc, 'MEMORY REQUESTS row: Pods A+B+C reserve most, with a small "Unallocated" slice.')
add_bullet(doc, 'MEMORY USAGE row: Pods A+B+C are under their memory requests — "Currently Unused".')
add_a(doc, 'Pod D has its own CPU and memory REQUESTS. The scheduler only looks at the Unallocated slice '
      'of each REQUEST bar — NOT at the Currently-Unused portion of the USAGE bar.')
add_a(doc, 'Decision:')
add_bullet(doc, 'If Pod D\'s CPU-request ≤ CPU-unallocated AND Pod D\'s Memory-request ≤ Memory-unallocated '
      '→ POD D CAN BE SCHEDULED on this node.')
add_bullet(doc, 'If either request exceeds its unallocated slice → Pod D CANNOT BE SCHEDULED on this '
      'node, even though the current USAGE shows free space. Reason: scheduling guarantees are given '
      'against REQUESTS, so the scheduler honours the worst-case case that existing pods will use '
      'everything they requested.')
add_a(doc, 'In the figure the "Unallocated" slice is small; judging by eye Pod D\'s requests exceed '
      'the remaining CPU and Memory request slices, so Pod D CANNOT be deployed on this node. '
      'The scheduler will either place D on another node or leave D Pending until capacity becomes '
      'available. Important nuance: one could reduce Pod D\'s requests (or set limits instead) to '
      'fit, but at the cost of potential CPU throttling or OOM-kill if existing pods ramp to their '
      'requests.')

add_q(doc, 'Q3(b) We want to deploy 4 Docker containers on a node with 3 CPUs. Container weights are '
      '1024, 256, 512, 128. Provide the calculation formula and compute how many CPUs each container '
      'can use. (7 marks)')
add_a(doc, 'Docker CPU shares (via the cgroup `cpu.shares`) are RELATIVE weights, applied only when CPUs '
      'are contended. When all containers are fully busy, each gets a fraction of the CPU time '
      'proportional to its weight. Formula:')
add_code(doc, """CPU_i = (weight_i / Σ weight_j) × total_CPUs""")
add_a(doc, 'Total of weights:')
add_code(doc, """Σ = 1024 + 256 + 512 + 128 = 1920""")
add_a(doc, 'Plug in total_CPUs = 3:')
add_bullet(doc, 'Container 1 (w = 1024): (1024/1920) × 3 = 1.6 CPUs')
add_bullet(doc, 'Container 2 (w = 256):  (256/1920) × 3 = 0.4 CPUs')
add_bullet(doc, 'Container 3 (w = 512):  (512/1920) × 3 = 0.8 CPUs')
add_bullet(doc, 'Container 4 (w = 128):  (128/1920) × 3 = 0.2 CPUs')
add_a(doc, 'Check: 1.6 + 0.4 + 0.8 + 0.2 = 3.0 CPUs ✓')
add_a(doc, 'Note: these are the shares ONLY when all 4 containers are simultaneously CPU-bound. If some '
      'are idle, the remaining containers can use more than their share up to the hard limit (if any). '
      'Docker `--cpus N` sets a hard cap via CFS quota/period; `--cpu-shares W` sets relative weight.')

# ---------- Q4 Hadoop ----------
add_heading(doc, 'Question 4 (20 marks) — Hadoop Fundamentals', level=2)

add_q(doc, 'Q4(a) Provide the workflow for how Hadoop works for a MapReduce application. (8 marks)')
add_a(doc, 'Classic MapReduce job flow on YARN:')
add_code(doc, """1. CLIENT submits the job to the ResourceManager
      (jar, input path, output path, config).

2. ResourceManager allocates a container and starts
   the ApplicationMaster (AM) for this job on some Node.

3. AM asks HDFS NameNode for the input splits
   (typically one split per HDFS block).

4. AM requests Containers from RM for map & reduce tasks.
   RM + NodeManagers allocate containers; AM launches
   MapTask JVMs (prefer data-local: place Map on the
   same node as the input block's DataNode).

5. MAP PHASE
   - RecordReader reads key/value pairs from the split.
   - map() emits intermediate (K,V).
   - Each map buffers output in a circular memory buffer;
     when buffer is 80 % full a background thread sorts
     by partition+key, optionally runs a Combiner, and
     spills to local disk.
   - When the map finishes, all spill files are merged
     into one sorted partitioned file per map task.

6. SHUFFLE & SORT
   - Reducers fetch their partitions from every map's
     local disk (over HTTP via NM).
   - Fetched segments are merged on the reducer side,
     producing one sorted stream per reducer.

7. REDUCE PHASE
   - reduce(K, values_iterator) is called once per key;
     emits output (K',V').
   - OutputFormat writes the result back to HDFS via
     DataNodes (with replication).

8. AM reports success to RM; RM tells the client;
   AM container is released.""")
add_a(doc, 'Fault tolerance: failed map/reduce tasks are re-scheduled on another node (input splits are '
      'on HDFS). AM failure causes the whole job to restart (unless AM-HA is enabled).')

add_q(doc, 'Q4(b) Explain the three main components of Hadoop (HDFS, MapReduce, Yarn). (6 marks)')
add_bullet(doc, 'HDFS (Hadoop Distributed File System) — append-only distributed filesystem optimised '
      'for large files. Master-slave: a single NameNode holds metadata (namespace + block→DataNode '
      'mapping) in memory; multiple DataNodes store fixed-size blocks (default 128 MB) with replication '
      '(default 3). Rack-aware placement for fault tolerance (2 in one rack + 1 in another). '
      'Strengths: linear scale-out, commodity hardware, fault-tolerant. Weakness: single NameNode is a '
      'SPOF in Hadoop 1 — solved by HDFS HA in Hadoop 2.')
add_bullet(doc, 'MapReduce — a programming model and execution engine for batch processing of large '
      'datasets. Developer writes `map(k,v)` and `reduce(k, values)` functions; framework handles '
      'split, shuffle, sort, fault tolerance. Suited to embarrassingly parallel batch (ETL, indexing, '
      'log analysis). Disk-heavy; largely superseded by Spark for interactive/iterative workloads.')
add_bullet(doc, 'YARN (Yet Another Resource Negotiator) — cluster resource manager introduced in '
      'Hadoop 2 that decoupled resource management from MapReduce. Central ResourceManager allocates '
      'containers; NodeManagers run on every slave; per-application ApplicationMaster negotiates '
      'resources and drives its own execution (map/reduce for MR jobs, driver/executor for Spark, '
      'etc.). Enables multiple frameworks (MR, Spark, Flink, Tez) to share one cluster.')

add_q(doc, 'Q4(c) Provide a description of Application Master, Resource Manager, Node Manager, NameNode, '
      'and DataNode — their roles and functions. (6 marks)')
add_bullet(doc, 'RESOURCE MANAGER (RM) — the master daemon of YARN. Has two internal components: a '
      'Scheduler (allocates containers based on policy — Capacity/Fair) and the ApplicationsManager '
      '(accepts job submissions, negotiates the first container to run an AM, restarts AM on failure). '
      'One per cluster (HA pair in production).')
add_bullet(doc, 'NODE MANAGER (NM) — YARN slave daemon, one per worker machine. Launches and monitors '
      'containers on behalf of the RM; reports node health and resource usage; enforces per-container '
      'limits via cgroups; runs auxiliary services such as the MapReduce shuffle handler.')
add_bullet(doc, 'APPLICATION MASTER (AM) — a per-application process. For MapReduce: MRAppMaster. It '
      'negotiates container resources from the RM, launches map & reduce tasks on NMs, retries failed '
      'tasks, and tells the RM when the job is done. Lives only for the duration of the job.')
add_bullet(doc, 'NAMENODE — HDFS master. Stores the filesystem namespace (inodes + block IDs per file) '
      'and the block-to-DataNode map entirely in memory; persistent form is fsimage + edits log. '
      'Handles client metadata operations (open, create, delete, rename) but NEVER the data itself. '
      'SPOF in Hadoop 1; HA (active/standby + QJM) in Hadoop 2.')
add_bullet(doc, 'DATANODE — HDFS slave. Stores blocks on local disk; serves block read/write to '
      'clients; reports block list to NameNode via heartbeat (every 3 s) and block report (every hour). '
      'Participates in replication pipeline when writing: client → DN1 → DN2 → DN3.')

# ---------- Q5 Kubernetes (same as Paper 2 Q5) ----------
add_heading(doc, 'Question 5 (20 marks) — Kubernetes', level=2)
add_a(doc, '(Identical to Paper 2 Question 5 — reuse the answers.)').runs[0].italic = True

add_q(doc, 'Q5(a) Cgroups and NameSpaces for Docker. (5 marks)')
add_a(doc, '(See Paper 2 Q5(a).) cgroups enforce resource limits/accounting (cpu, memory, blkio, pids, '
      'net_cls, devices). namespaces provide isolation (pid, net, mnt, uts, ipc, user, cgroup).')

add_q(doc, 'Q5(b) Explain the architecture of Kubernetes. (7 marks)')
add_a(doc, '(See Paper 2 Q5(b).) Control plane: kube-apiserver, etcd, scheduler, controller-manager. '
      'Data plane: kubelet, container runtime, kube-proxy, CNI.')

add_q(doc, 'Q5(c) Describe the main workflow of Kubernetes and draw a figure to show it. (8 marks)')
add_a(doc, '(See Paper 2 Q5(c) — kubectl → apiserver → etcd; controllers reconcile; scheduler picks a '
      'node; kubelet launches container via CRI; kube-proxy programs service rules.)')

# ---------- Q6 ----------
add_heading(doc, 'Question 6 (20 marks) — PySpark DataFrame: E-commerce Transactions', level=2)
add_a(doc, 'Given code:').runs[0].italic = True
add_code(doc, """from pyspark.sql import SparkSession
from pyspark.sql.functions import col, sum as _sum, avg

spark = SparkSession.builder.appName("ECommerceAnalysis").getOrCreate()

df = spark.read.json("path/to/transactions.json")

revenue_per_category = (df
    .groupBy("category")
    .agg(_sum("amount").alias("total_revenue")))

top_3_categories = (revenue_per_category
    .orderBy(col("total_revenue").desc())
    .limit(3))

top_category_transactions = df.join(top_3_categories, "category")
average_transaction_amount = (top_category_transactions
    .groupBy("category")
    .agg(avg("amount").alias("average_transaction_amount")))

top_3_categories.show()
average_transaction_amount.show()

spark.stop()""")

add_q(doc, 'Q6(a) Describe each line of the code. (5 marks)')
add_bullet(doc, '`from pyspark.sql import SparkSession` — DataFrame entry point.')
add_bullet(doc, '`from pyspark.sql.functions import col, sum as _sum, avg` — helper funcs; `sum` is '
      'renamed to `_sum` so as not to shadow Python\'s built-in `sum`.')
add_bullet(doc, '`spark = SparkSession.builder.appName("ECommerceAnalysis").getOrCreate()` — get/create '
      'a session called ECommerceAnalysis.')
add_bullet(doc, '`df = spark.read.json("path/to/transactions.json")` — read JSON lines into a DataFrame; '
      'schema is inferred.')
add_bullet(doc, '`revenue_per_category = df.groupBy("category").agg(_sum("amount").alias("total_revenue"))` '
      '— TRANSFORMATION: group by category, sum `amount` → `total_revenue`.')
add_bullet(doc, '`top_3_categories = revenue_per_category.orderBy(col("total_revenue").desc()).limit(3)` '
      '— TRANSFORMATION: sort descending by revenue; keep the top 3.')
add_bullet(doc, '`top_category_transactions = df.join(top_3_categories, "category")` — '
      'TRANSFORMATION: inner join original transactions against the top-3 categories (filtering df '
      'to only those three categories).')
add_bullet(doc, '`average_transaction_amount = top_category_transactions.groupBy("category")'
      '.agg(avg("amount").alias("average_transaction_amount"))` — TRANSFORMATION: per kept category, '
      'compute mean transaction amount.')
add_bullet(doc, '`top_3_categories.show()` — ACTION: prints top-3 categories.')
add_bullet(doc, '`average_transaction_amount.show()` — ACTION: prints averages.')
add_bullet(doc, '`spark.stop()` — shut the session down.')

add_q(doc, 'Q6(b) What is "Transformation" in Spark? Describe and give an example from the above code. (5 marks)')
add_a(doc, '(See Paper 2 Q6(b) for the definition.) A transformation is a lazy operation returning a new '
      'Dataset; narrow (map/filter/select/withColumn) or wide (groupBy/join/orderBy/distinct).')
add_a(doc, 'Examples from this code:')
add_bullet(doc, '`df.groupBy("category").agg(_sum("amount").alias("total_revenue"))` — wide '
      '(shuffle by category).')
add_bullet(doc, '`revenue_per_category.orderBy(col("total_revenue").desc())` — wide (range partitioning).')
add_bullet(doc, '`df.join(top_3_categories, "category")` — wide (shuffle by join key).')
add_bullet(doc, '`top_category_transactions.groupBy("category").agg(avg("amount").alias(...))` — wide.')

add_q(doc, 'Q6(c) What is "Action" in Spark? Describe and give an example from the above code. (5 marks)')
add_a(doc, '(See Paper 2 Q6(c) for the definition.) Actions trigger execution and return a value to the '
      'driver / write to external sink. Each Action = one Spark Job.')
add_a(doc, 'Examples from this code:')
add_bullet(doc, '`top_3_categories.show()` — prints up to 20 rows; triggers the whole lineage (read → '
      'groupBy → orderBy → limit).')
add_bullet(doc, '`average_transaction_amount.show()` — triggers: read → groupBy total_revenue → orderBy '
      '→ limit → join → groupBy → avg → show. Spark re-computes the lineage unless we `.cache()` an '
      'intermediate DataFrame such as top_3_categories.')

add_q(doc, 'Q6(d) What are the main components of Spark? (5 marks)')
add_a(doc, '(See Paper 2 Q6(d).) Driver, Cluster Manager, Worker, Executor, Task. Plus Catalyst + '
      'Tungsten for SQL/DataFrame. Important refinement: Executors live for the life of the '
      'application; Tasks are one per partition per stage; Jobs are triggered by Actions and split into '
      'stages by shuffle boundaries.')

# ===================== END =====================
doc.add_page_break()
end = doc.add_heading('End of past-paper Q&A', level=1)
end.alignment = WD_ALIGN_PARAGRAPH.CENTER
closing = doc.add_paragraph('Prepared for open-book COMP4442 exam. Bring this + exam_notes.docx + coding_questions.docx.')
closing.alignment = WD_ALIGN_PARAGRAPH.CENTER
closing.runs[0].italic = True

out = r"C:\Users\User\Downloads\4442\exam_output\pastpapers_qa.docx"
doc.save(out)
print(f"Saved {out}")
